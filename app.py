import os
import csv
import streamlit as st
import re
import json
from datetime import datetime
from crossref.restful import Works
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from tqdm import tqdm
from docx.oxml import OxmlElement
import base64
import html
import concurrent.futures
from typing import List, Dict, Tuple, Set, Any, Optional
import hashlib
import time
from collections import Counter
import functools
import logging
from pathlib import Path
import sqlite3
from contextlib import contextmanager
import requests
import pandas as pd
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
import nltk
from nltk.corpus import stopwords, wordnet
from nltk.stem import WordNetLemmatizer
from nltk.tokenize import word_tokenize
import spacy
from sentence_transformers import SentenceTransformer, util
from gensim.models import Phrases
from gensim.models.phrases import Phraser
from typing import Optional
from PIL import Image as PILImage
from PIL import Image

# Download NLTK data - do it immediately and not quietly to see errors
import nltk
try:
    nltk.download('punkt')
    nltk.download('stopwords')
    nltk.download('wordnet')
    nltk.download('averaged_perceptron_tagger')
    nltk.download('punkt_tab')  # Add this for sentence tokenization
    print("NLTK data downloaded successfully")
except Exception as e:
    print(f"Error downloading NLTK data: {e}")
    # Try to download punkt resources specifically
    try:
        nltk.download('punkt_tab')
        print("punkt_tab downloaded")
    except:
        pass

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('citation_processor.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Configuration
class Config:
    """Application configuration constants"""
    # File paths
    DB_PATH = "doi_cache.db"
    LTWA_CSV_PATH = "ltwa.csv"
    USER_PREFS_DB = "user_preferences.db"
    
    # API settings
    CROSSREF_WORKERS = 3
    CROSSREF_RETRY_WORKERS = 2
    REQUEST_TIMEOUT = 30
    
    # Caching
    CACHE_TTL_HOURS = 24 * 7  # 1 week
    
    # Validation
    MIN_REFERENCES_FOR_STATS = 5
    MAX_REFERENCES = 1000
    MIN_REFERENCES_FOR_RECOMMENDATIONS = 10
    
    # Retry failed DOI
    MAX_RETRY_ATTEMPTS = 2
    RETRY_DELAY_SECONDS = 1

    # OpenAlex API settings
    OPENALEX_MAX_WORKERS = 10
    OPENALEX_PER_PAGE = 200  # Увеличили с 100
    OPENALEX_MAX_PAGES = 5   # Увеличили с 3
    OPENALEX_REQUEST_TIMEOUT = 25
    OPENALEX_MAX_WORKS_PER_TOPIC = 500  # Увеличили с 200
    OPENALEX_MAX_TOTAL_WORKS = 1000
    OPENALEX_CACHE_TTL_MINUTES = 60  # Кэширование тем
    
    # Styles
    NUMBERING_STYLES = ["No numbering", "1", "1.", "1)", "(1)", "[1]"]
    AUTHOR_FORMATS = ["AA Smith", "A.A. Smith", "Smith AA", "Smith A.A", "Smith, A.A.",
                     "A A Smith", "A. A. Smith", "Smith A A", "Smith A. A.", "Smith A. A", "Smith, A. A."]
    PAGE_FORMATS = ["122 - 128", "122-128", "122 – 128", "122–128", "122–8", "122"]
    DOI_FORMATS = ["10.10/xxx", "doi:10.10/xxx", "DOI:10.10/xxx", "https://doi.org/10.10/xxx"]
    JOURNAL_STYLES = ["{Full Journal Name}", "{J. Abbr.}", "{J Abbr}"]
    AVAILABLE_ELEMENTS = ["", "Authors", "Title", "Journal", "Year", "Volume", "Issue", "Pages", "DOI"]
    
    # Progress bar colors
    PROGRESS_COLORS = {
        'start': '#FF6B6B',
        'middle': '#4ECDC4', 
        'end': '#45B7D1'
    }
    
    # Themes
    THEMES = {
        'light': {
            'name': 'Light',
            'primary': '#1f77b4',
            'secondary': '#2ca02c',
            'accent': '#ff7f0e',
            'background': '#f8f9fa',
            'secondaryBackground': '#ffffff',
            'text': '#212529',
            'font': "'Segoe UI', 'Helvetica Neue', sans-serif",
            'border': '#dee2e6',
            'cardBackground': '#ffffff',
            'buttonStyle': 'rounded',
            'shadow': '0 2px 4px rgba(0,0,0,0.1)'
        },
        'dark': {
            'name': 'Dark',
            'primary': '#4ECDC4',
            'secondary': '#FF6B6B',
            'accent': '#45B7D1',
            'background': '#1a1d23',
            'secondaryBackground': '#2d323d',
            'text': '#e9ecef',
            'font': "'Inter', 'Roboto', sans-serif",
            'border': '#495057',
            'cardBackground': '#2d323d',
            'buttonStyle': 'rounded',
            'shadow': '0 2px 8px rgba(0,0,0,0.3)'
        },
        'library': {
            'name': 'Library',
            'primary': '#8B4513',
            'secondary': '#654321',
            'accent': '#D2691E',
            'background': '#F5F5DC',
            'secondaryBackground': '#FAF0E6',
            'text': '#2F4F4F',
            'font': "'Georgia', 'Times New Roman', serif",
            'border': '#DEB887',
            'cardBackground': '#FFF8DC',
            'buttonStyle': 'classic',
            'shadow': '0 2px 6px rgba(139,69,19,0.2)'
        },
        'barbie': {
            'name': 'Barbie-style',
            'primary': '#FF69B4',
            'secondary': '#FF1493',
            'accent': '#FFB6C1',
            'background': '#FFF0F5',
            'secondaryBackground': '#FFE4E1',
            'text': '#8B008B',
            'font': "'Comic Sans MS', cursive, sans-serif",
            'border': '#FFB6C1',
            'cardBackground': '#FFF0F5',
            'buttonStyle': 'rounded-full',
            'shadow': '0 4px 12px rgba(255,105,180,0.3)'
        },
        'newspaper': {
            'name': 'Newspaper',
            'primary': '#1E1A4F',
            'secondary': '3C2AA8',
            'accent': '#E5C9A8',
            'background': '#EEEDF7',
            'secondaryBackground': '#FFFFFF',
            'text': '#4A3728',
            'font': "'Playfair Display', 'Georgia', serif",
            'border': '#D4B996',
            'cardBackground': '#FFFFFF',
            'buttonStyle': 'classic',
            'shadow': '0 4px 12px rgba(193, 154, 107, 0.1)'
        }
    }
    
    # Application stages
    STAGES = {
        'start': 'Start',
        'select': 'Select',
        'create': 'Create',
        'io': 'Input/Output',
        'results': 'Results'
    }
    
    # Statistics settings
    DISPLAY_STATISTICS = True
    
    # Article recommendations settings
    RECOMMENDATION_EMAIL = "citation.style.constructor@gmail.com"
    MAX_RECOMMENDATIONS = 20
    RECOMMENDATION_YEARS_BACK = 5
    MIN_SIMILARITY_SCORE = 0.1

# Translations
TRANSLATIONS = {
    'en': {
        'header': '🎨 Citation Style Constructor',
        'general_settings': '⚙️ General Settings',
        'element_config': '📑 Element Configuration',
        'style_preview': '👀 Style Preview',
        'data_input': '📁 Data Input',
        'data_output': '📤 Data Output',
        'numbering_style': 'Numbering:',
        'author_format': 'Authors:',
        'author_separator': 'Separator:',
        'et_al_limit': 'Et al after:',
        'use_and': "'and'",
        'use_ampersand': "'&'",
        'doi_format': 'DOI format:',
        'doi_hyperlink': 'DOI as hyperlink',
        'page_format': 'Pages:',
        'final_punctuation': 'Final punctuation:',
        'element': 'Element',
        'italic': 'Italic',
        'bold': 'Bold',
        'parentheses': 'Parentheses',
        'separator': 'Separator',
        'input_method': 'Input:',
        'output_method': 'Output:',
        'select_docx': 'Select DOCX',
        'enter_references': 'Enter DOI/DOIs (one per line)',
        'references': 'References:',
        'results': 'Results:',
        'process': '🚀 Process',
        'example': 'Example:',
        'error_select_element': 'Select at least one element!',
        'processing': '⏳ Processing...',
        'upload_file': 'Upload a file!',
        'enter_references_error': 'Enter references!',
        'select_docx_output': 'Select DOCX output to download!',
        'doi_txt': '📄 DOI (TXT)',
        'references_docx': '📋 References (DOCX)',
        'found_references': 'Found {} references.',
        'found_references_text': 'Found {} references in text.',
        'statistics': 'Statistics: {} DOI found, {} not found.',
        'language': 'Language:',
        'gost_style': 'Apply GOST Style',
        'export_style': '📤 Export Style',
        'import_style': '📥 Import Style',
        'export_file_name': 'File name:',
        'import_file': 'Select style file:',
        'export_success': 'Style exported successfully!',
        'import_success': 'Style imported successfully!',
        'import_error': 'Error importing style file!',
        'processing_status': 'Processing references...',
        'current_reference': 'Current: {}',
        'processed_stats': 'Processed: {}/{} | Found: {} | Errors: {}',
        'time_remaining': 'Estimated time remaining: {}',
        'duplicate_reference': '🔄 Repeated Reference (See #{})',
        'batch_processing': 'Batch processing DOI...',
        'extracting_metadata': 'Extracting metadata...',
        'checking_duplicates': 'Checking for duplicates...',
        'retrying_failed': 'Retrying failed DOI requests...',
        'bibliographic_search': 'Searching by bibliographic data...',
        'style_presets': 'Style Presets',
        'gost_button': 'GOST',
        'acs_button': 'ACS (MDPI)',
        'rsc_button': 'RSC',
        'cta_button': 'CTA',
        'style_preset_tooltip': 'Here are some styles maintained by individual publishers. For major publishers (Elsevier, Springer Nature, Wiley), styles vary from journal to journal. To create (or reformat) references for a specific journal, use the Citation Style Constructor.',
        'journal_style': 'Journal style:',
        'full_journal_name': 'Full Journal Name',
        'journal_abbr_with_dots': 'J. Abbr.',
        'journal_abbr_no_dots': 'J Abbr',
        'short_guide_title': 'A short guide for the conversion of doi-based references',
        'step_1': '❶ Select a ready reference style (ACS(MDPI), RSC, or CTA), or create your own style by selecting the sequence, design, and punctuation of the element configurations',
        'step_1_note': '(!) The punctuation boxes enable various items to be included between element configurations (simple punctuation, Vol., Issue…)',
        'step_2': '❷ Then, use the Style Presets to change certain element configurations for each reformatted reference.',
        'step_3': '❸ The Style Preview function enables users to visualize the final form of their reference style',
        'step_4': '❹ If the final style is appropriate, select the Docx or Text option in the Data Input section and upload the corresponding information (reference list). Then, in the Data Output section, select the required options and press "Process" to initiate reformatting.',
        'step_5': '❺ After processing is complete, download the reformatted references in your preferred format.',
        'step_5_note': '(!) Outputting the Docx file is recommended, as it preserves formatting (e.g., bold, italic, and hyperlinks) and includes additional stats at the end of the document.',
        'step_6': '❻ After creating your final version of the style, save it so that you can upload it again in the next session. Use the Style Management section for this purpose.',
        'validation_error_no_elements': 'Please configure at least one element or select a preset style!',
        'validation_error_too_many_references': 'Too many references (maximum {} allowed)',
        'validation_warning_few_references': 'Few references for meaningful statistics',
        'cache_initialized': 'Cache initialized successfully',
        'cache_cleared': 'Cache cleared successfully',
        'theme_selector': 'Theme:',
        'light_theme': 'Light',
        'dark_theme': 'Dark',
        'library_theme': 'Library',
        'barbie_theme': 'Barbie',
        'newspaper_theme': 'Newspaper',
        'mobile_view': 'Mobile View',
        'desktop_view': 'Desktop View',
        'clear_button': '🗑️ Clear',
        'back_button': '↩️ Back',
        'stage_start': 'Start',
        'stage_select': 'Select',
        'stage_create': 'Create',
        'stage_io': 'Input/Output',
        'stage_results': 'Results',
        'start_title': 'Welcome to DOI-based Citation Style Constructor',
        'start_ready_presets': '📋 Ready Style Presets',
        'start_create_style': '🎨 Create Style',
        'start_load_style': '📂 Load Your Saved Style',
        'start_description': 'Choose how you want to format your references:',
        'select_title': 'Select Style Preset',
        'select_description': 'Choose one of the ready-made citation styles:',
        'create_title': 'Create Custom Style',
        'create_description': 'Configure your custom citation style',
        'io_title': 'Input and Output',
        'io_description': 'Provide your references and choose output format',
        'results_title': 'Results',
        'results_description': 'Processing complete! Download your formatted references',
        'export_style_button': '💾 Export Style',
        'proceed_to_io': '➡️ Proceed to Input/Output',
        'back_to_start': '⬅️ Back to Start',
        'clear_all': '🗑️ Clear All',
        'choose_theme': 'Choose Theme:',
        'choose_language': 'Choose Language:',
        'stage_indicator': 'Stage:',
        'loading': 'Loading...',
        'no_file_selected': 'No file selected',
        'style_loaded': 'Style loaded successfully!',
        'ready_styles': 'Ready Styles',
        'custom_style': 'Custom Style',
        'load_style': 'Load Style',
        'next_step': 'Next Step',
        'prev_step': 'Previous Step',
        'process_references': 'Process References',
        'download_results': 'Download Results',
        'view_statistics': 'View Statistics',
        'statistics_title': 'Processing Statistics',
        'total_references': 'Total References:',
        'doi_found': 'DOI Found:',
        'doi_not_found': 'DOI Not Found:',
        'duplicates_found': 'Duplicates Found:',
        'processing_time': 'Processing Time:',
        'download_txt': 'Download TXT',
        'download_docx': 'Download DOCX',
        'try_again': 'Try Again',
        'new_session': 'New Session',
        'recommend_similar_articles': '🔍 Recommend Similar Articles',
        'recommendations_title': 'Article Recommendations',
        'recommendations_description': 'Recommendations are generated based on new or low-citation articles published within the last 3 years',
        'recommendations_loading': '🔍 Searching for recommendations...',
        'recommendations_not_enough': 'At least {} references are required for recommendations.',
        'recommendations_no_results': 'No recommendations found. Try adjusting search parameters.',
        'recommendations_error': 'Error fetching recommendations: {}',
        'recommendations_count': 'Found {} recommendations',
        'recommendation_score': 'Relevance score:',
        'recommendation_year': 'Year:',
        'recommendation_journal': 'Journal:',
        'recommendation_abstract': 'Abstract:',
        'recommendation_show_abstract': 'Show abstract',
        'recommendation_hide_abstract': 'Hide abstract',
        'recommendation_download': '📥 Download Recommendations',
        'recommendation_download_txt': 'Download as TXT',
        'recommendation_download_csv': 'Download as CSV',
        'missing_metadata_warning': '⚠️ Volume/page/article number information is missing. This may indicate a non-journal source (book, chapter, or conference paper) or a journal article with incomplete issue assignment. Please verify the source.',
    },
    'ru': {
        'header': '🎨 Конструктор стилей цитирования',
        'general_settings': '⚙️ Настройки',
        'element_config': '📑 Конфигурация элементов',
        'style_preview': '👀 Предпросмотр',
        'data_input': '📁 Ввод',
        'data_output': '📤 Вывод',
        'numbering_style': 'Нумерация:',
        'author_format': 'Авторы:',
        'author_separator': 'Разделитель:',
        'et_al_limit': 'Et al после:',
        'use_and': "'и'",
        'use_ampersand': "'&'",
        'doi_format': 'Формат DOI:',
        'doi_hyperlink': 'DOI как ссылка',
        'page_format': 'Страницы:',
        'final_punctuation': 'Конечная пунктуация:',
        'element': 'Элемент',
        'italic': 'Курсив',
        'bold': 'Жирный',
        'parentheses': 'Скобки',
        'separator': 'Разделитель',
        'input_method': 'Ввод:',
        'output_method': 'Вывод:',
        'select_docx': 'Выберите DOCX',
        'enter_references': 'Введите DOI/DOIs (по одной на строку)',
        'references': 'Ссылки:',
        'results': 'Результаты:',
        'process': '🚀 Обработать',
        'example': 'Пример:',
        'error_select_element': 'Выберите хотя бы один элемент!',
        'processing': '⏳ Обработка...',
        'upload_file': 'Загрузите файл!',
        'enter_references_error': 'Введите ссылки!',
        'select_docx_output': 'Выберите DOCX для скачивания!',
        'doi_txt': '📄 DOI (TXT)',
        'references_docx': '📋 Ссылки (DOCX)',
        'found_references': 'Найдено {} ссылок.',
        'found_references_text': 'Найдено {} ссылок в тексте.',
        'statistics': 'Статистика: {} DOI найдено, {} не найдено.',
        'language': 'Язык:',
        'gost_style': 'Применить стиль ГОСТ',
        'export_style': '📤 Экспорт стиля',
        'import_style': '📥 Импорт стиля',
        'export_file_name': 'Имя файла:',
        'import_file': 'Выберите файл стиля:',
        'export_success': 'Стиль экспортирован успешно!',
        'import_success': 'Стиль импортирован успешно!',
        'import_error': 'Ошибка импорта файла стиля!',
        'processing_status': 'Обработка ссылок...',
        'current_reference': 'Текущая: {}',
        'processed_stats': 'Обработано: {}/{} | Найдено: {} | Ошибки: {}',
        'time_remaining': 'Примерное время до завершения: {}',
        'duplicate_reference': '🔄 Повторная ссылка (См. #{})',
        'batch_processing': 'Пакетная обработка DOI...',
        'extracting_metadata': 'Извлечение метаданных...',
        'checking_duplicates': 'Проверка на дубликаты...',
        'retrying_failed': 'Повторная попытка для неудачных DOI...',
        'bibliographic_search': 'Поиск по библиографическим данным...',
        'style_presets': 'Готовые стили',
        'gost_button': 'ГОСТ',
        'acs_button': 'ACS (MDPI)',
        'rsc_button': 'RSC',
        'cta_button': 'CTA',
        'style_preset_tooltip': 'Здесь указаны некоторые стили, которые сохраняются в пределах одного издательства. Для ряда крупных издательств (Esevier, Springer Nature, Wiley) стиль отличается от журнала к журналу. Для формирования (или переформатирования) ссылок для конкретного журнала предлагаем воспользоваться конструктором ссылок.',
        'journal_style': 'Стиль журнала:',
        'full_journal_name': 'Полное название журнала',
        'journal_abbr_with_dots': 'J. Abbr.',
        'journal_abbr_no_dots': 'J Abbr',
        'short_guide_title': 'Краткое руководство для конвертации ссылок, имеющих doi',
        'step_1': '❶ Выберите готовый стиль ссылок (ГОСТ, ACS(MDPI), RSC или CTA) или создайте свой собственный стиль, выбрав последовательность, оформление и пунктуацию конфигураций элементов',
        'step_1_note': '(!) Поля пунктуации позволяют включать различные элементы между конфигурациями (простая пунктуация, Том, Выпуск…)',
        'step_2': '❷ Затем используйте готовые стили, чтобы изменить определенные конфигурации элементов для каждой переформатированной ссылки.',
        'step_3': '❸ Функция предпросмотра стиля позволяет визуализировать окончательную форму вашего стиля ссылок',
        'step_4': '❹ Если окончательный стиль подходит, выберите опцию Docx или Текст в разделе ввода данных и загрузите соответствующую информацию (список литературы). Затем в разделе вывода данных выберите нужные опции и нажмите "Обработать" для начала переформатирования.',
        'step_5': '❺ После завершения обработки загрузите переформатированные ссылки в предпочитаемом формате.',
        'step_5_note': '(!) Рекомендуется выводить файл Docx, так как он сохраняет форматирование (например, жирный шрифт, курсив и гиперссылки) и включает дополнительную статистику в конце документа.',
        'step_6': '❻ После создания окончательной версии стиля сохраните его, чтобы можно было снова загрузить в следующей сессии. Для этого используйте раздел Style Management.',
        'validation_error_no_elements': 'Пожалуйста, настройте хотя бы один элемент или выберите готовый стиль!',
        'validation_error_too_many_references': 'Слишком много ссылок (максимум {} разрешено)',
        'validation_warning_few_references': 'Мало ссылок для значимой статистики',
        'cache_initialized': 'Кэш инициализирован успешно',
        'cache_cleared': 'Кэш очищен успешно',
        'theme_selector': 'Тема:',
        'light_theme': 'Светлая',
        'dark_theme': 'Тёмная',
        'library_theme': 'Библиотечная',
        'barbie_theme': 'Барби',
        'newspaper_theme': 'Газетная',
        'mobile_view': 'Мобильный вид',
        'desktop_view': 'Десктопный вид',
        'clear_button': '🗑️ Очистить',
        'back_button': '↩️ Назад',
        'stage_start': 'Старт',
        'stage_select': 'Выбор',
        'stage_create': 'Создание',
        'stage_io': 'Ввод/Вывод',
        'stage_results': 'Результаты',
        'start_title': 'Добро пожаловать в Конструктор стилей цитирования ссылок по их DOI',
        'start_ready_presets': '📋 Готовые стили',
        'start_create_style': '🎨 Создать стиль',
        'start_load_style': '📂 Загрузить сохраненный стиль',
        'start_description': 'Выберите способ форматирования ссылок:',
        'select_title': 'Выбор готового стиля',
        'select_description': 'Выберите один из готовых стилей цитирования:',
        'create_title': 'Создание пользовательского стиля',
        'create_description': 'Настройте свой собственный стиль цитирования',
        'io_title': 'Ввод и вывод данных',
        'io_description': 'Предоставьте ссылки и выберите формат вывода',
        'results_title': 'Результаты обработки',
        'results_description': 'Обработка завершена! Скачайте отформатированные ссылки',
        'export_style_button': '💾 Экспорт стиля',
        'proceed_to_io': '➡️ Перейти к Вводу/Выводу',
        'back_to_start': '⬅️ Вернуться к Старту',
        'clear_all': '🗑️ Очистить всё',
        'choose_theme': 'Выберите тему:',
        'choose_language': 'Выберите язык:',
        'stage_indicator': 'Этап:',
        'loading': 'Загрузка...',
        'no_file_selected': 'Файл не выбран',
        'style_loaded': 'Стиль успешно загружен!',
        'ready_styles': 'Готовые стили',
        'custom_style': 'Пользовательский стиль',
        'load_style': 'Загрузить стиль',
        'next_step': 'Следующий шаг',
        'prev_step': 'Предыдущий шаг',
        'process_references': 'Обработать ссылки',
        'download_results': 'Скачать результаты',
        'view_statistics': 'Просмотр статистики',
        'statistics_title': 'Статистика обработки',
        'total_references': 'Всего ссылок:',
        'doi_found': 'DOI найдено:',
        'doi_not_found': 'DOI не найдено:',
        'duplicates_found': 'Дубликатов найдено:',
        'processing_time': 'Время обработки:',
        'download_txt': 'Скачать TXT',
        'download_docx': 'Скачать DOCX',
        'try_again': 'Попробовать снова',
        'new_session': 'Новая сессия',
        'recommend_similar_articles': '🔍 Рекомендовать похожие статьи',
        'recommendations_title': 'Рекомендации статей',
        'recommendations_description': 'Рекомендации формируются на основе новых (или малоцитируемых) статей, опубликованных за 3 года от времени проведения анализа',
        'recommendations_loading': '🔍 Поиск рекомендаций...',
        'recommendations_not_enough': 'Для рекомендаций требуется не менее {} ссылок.',
        'recommendations_no_results': 'Рекомендации не найдены. Попробуйте изменить параметры поиска.',
        'recommendations_error': 'Ошибка при получении рекомендаций: {}',
        'recommendations_count': 'Найдено {} рекомендаций',
        'recommendation_score': 'Оценка релевантности:',
        'recommendation_year': 'Год:',
        'recommendation_journal': 'Журнал:',
        'recommendation_abstract': 'Аннотация:',
        'recommendation_show_abstract': 'Показать аннотацию',
        'recommendation_hide_abstract': 'Скрыть аннотацию',
        'recommendation_download': '📥 Скачать рекомендации',
        'recommendation_download_txt': 'Скачать как TXT',
        'recommendation_download_csv': 'Скачать как CSV',
        'missing_metadata_warning': '⚠️ В этой ссылке отсутствует информация о томе/страниацах/номере статьи. Это может указывать на нежурнальный источник (книгу, главу, или конференционный тезис). Необходимо уточнение.',
    }
}

# Initialize Crossref
works = Works()

# DOI Cache
class DOICache:
    """Cache for storing DOI metadata"""
    
    def __init__(self, db_path: str = Config.DB_PATH):
        self.db_path = db_path
        self._init_db()
    
    def _init_db(self):
        """Initialize database"""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('''
                CREATE TABLE IF NOT EXISTS doi_cache (
                    doi TEXT PRIMARY KEY,
                    metadata TEXT NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    accessed_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_doi ON doi_cache(doi)')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_accessed_at ON doi_cache(accessed_at)')
    
    def get(self, doi: str) -> Optional[Dict]:
        """Get metadata from cache"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                result = conn.execute(
                    'SELECT metadata FROM doi_cache WHERE doi = ? AND datetime(accessed_at) > datetime("now", ?)',
                    (doi, f"-{Config.CACHE_TTL_HOURS} hours")
                ).fetchone()
                
                if result:
                    conn.execute(
                        'UPDATE doi_cache SET accessed_at = CURRENT_TIMESTAMP WHERE doi = ?',
                        (doi,)
                    )
                    return json.loads(result[0])
        except Exception as e:
            logger.error(f"Cache get error for {doi}: {e}")
        return None
    
    def set(self, doi: str, metadata: Dict):
        """Save metadata to cache"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.execute(
                    'INSERT OR REPLACE INTO doi_cache (doi, metadata) VALUES (?, ?)',
                    (doi, json.dumps(metadata))
                )
        except Exception as e:
            logger.error(f"Cache set error for {doi}: {e}")
    
    def clear_old_entries(self):
        """Clear outdated entries"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.execute(
                    'DELETE FROM doi_cache WHERE datetime(accessed_at) <= datetime("now", ?)',
                    (f"-{Config.CACHE_TTL_HOURS} hours",)
                )
        except Exception as e:
            logger.error(f"Cache cleanup error: {e}")

# Initialize cache
doi_cache = DOICache()

# User Preferences Manager
class UserPreferencesManager:
    """User preferences manager"""
    
    def __init__(self, db_path: str = Config.USER_PREFS_DB):
        self.db_path = db_path
        self._init_db()
    
    def _init_db(self):
        """Initialize preferences database"""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('''
                CREATE TABLE IF NOT EXISTS user_preferences (
                    ip_address TEXT PRIMARY KEY,
                    language TEXT DEFAULT 'ru',
                    theme TEXT DEFAULT 'light',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            conn.execute('CREATE INDEX IF NOT EXISTS idx_ip ON user_preferences(ip_address)')
    
    def get_user_ip(self):
        """Get user IP address"""
        try:
            if hasattr(st, 'experimental_user'):
                return getattr(st.experimental_user, 'ip', 'unknown')
        except:
            pass
        return 'unknown'
    
    def get_preferences(self, ip: str) -> Dict[str, Any]:
        """Get user preferences"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                result = conn.execute(
                    'SELECT language, theme FROM user_preferences WHERE ip_address = ?',
                    (ip,)
                ).fetchone()
                
                if result:
                    return {
                        'language': result[0],
                        'theme': result[1]
                    }
        except Exception as e:
            logger.error(f"Error getting preferences for {ip}: {e}")
        
        return {
            'language': 'en',
            'theme': 'light'
        }
    
    def save_preferences(self, ip: str, preferences: Dict[str, Any]):
        """Save user preferences"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.execute('''
                    INSERT OR REPLACE INTO user_preferences 
                    (ip_address, language, theme, updated_at) 
                    VALUES (?, ?, ?, CURRENT_TIMESTAMP)
                ''', (
                    ip,
                    preferences.get('language', 'ru'),
                    preferences.get('theme', 'light')
                ))
        except Exception as e:
            logger.error(f"Error saving preferences for {ip}: {e}")

# Style Validator
class StyleValidator:
    """Style configuration validator"""
    
    @staticmethod
    def validate_style_config(style_config: Dict) -> Tuple[bool, List[str]]:
        """Validate style configuration"""
        errors = []
        warnings = []
        
        has_elements = bool(style_config.get('elements'))
        has_preset = any([
            style_config.get('gost_style', False),
            style_config.get('acs_style', False), 
            style_config.get('rsc_style', False),
            style_config.get('cta_style', False)
        ])
        
        if not has_elements and not has_preset:
            errors.append(get_text('validation_error_no_elements'))
        
        if has_elements:
            elements = style_config['elements']
            for i, (element, config) in enumerate(elements):
                if not element:
                    errors.append(f"Element {i+1} is empty")
                if not config.get('separator', '').strip() and i < len(elements) - 1:
                    warnings.append(f"Element {i+1} has empty separator")
        
        return len(errors) == 0, errors + warnings
    
    @staticmethod
    def validate_references_count(references: List[str]) -> Tuple[bool, List[str]]:
        """Validate references count"""
        errors = []
        warnings = []
        
        if len(references) > Config.MAX_REFERENCES:
            errors.append(get_text('validation_error_too_many_references').format(Config.MAX_REFERENCES))
        
        if len(references) < Config.MIN_REFERENCES_FOR_STATS:
            warnings.append(get_text('validation_warning_few_references'))
        
        return len(errors) == 0, errors + warnings

# Progress Manager
class ProgressManager:
    """Processing progress manager"""
    
    def __init__(self):
        self.start_time = None
        self.progress_data = {
            'total': 0,
            'processed': 0,
            'found': 0,
            'errors': 0,
            'phase': 'initializing'
        }
    
    def start_processing(self, total: int):
        """Start processing"""
        self.start_time = time.time()
        self.progress_data = {
            'total': total,
            'processed': 0,
            'found': 0,
            'errors': 0,
            'phase': 'processing'
        }

    def update_progress(self, processed: int, found: int, errors: int, phase: str = None):
        """Update progress"""
        self.progress_data.update({
            'processed': processed,
            'found': found,
            'errors': errors
        })
        if phase:
            self.progress_data['phase'] = phase
        
        if self.start_time:
            elapsed = time.time() - self.start_time
            total = self.progress_data['total']
            
            if processed > 0 and total > 0:
                estimated_total = (elapsed / processed) * total
                self.progress_data['time_remaining'] = estimated_total - elapsed
    
    def get_progress_info(self) -> Dict[str, Any]:
        """Get progress information"""
        if not self.start_time:
            return self.progress_data
        
        elapsed = time.time() - self.start_time
        processed = self.progress_data['processed']
        total = self.progress_data['total']
        
        time_remaining = None
        if processed > 0 and total > 0:
            estimated_total = (elapsed / processed) * total
            time_remaining = estimated_total - elapsed
            if time_remaining < 0:
                time_remaining = 0
        
        progress_ratio = processed / total if total > 0 else 0
        
        return {
            **self.progress_data,
            'elapsed_time': elapsed,
            'time_remaining': time_remaining,
            'progress_ratio': progress_ratio
        }
    
    def get_progress_color(self, progress_ratio: float) -> str:
        """Get progress bar color based on progress"""
        if progress_ratio < 0.33:
            return Config.PROGRESS_COLORS['start']
        elif progress_ratio < 0.66:
            return Config.PROGRESS_COLORS['middle']
        else:
            return Config.PROGRESS_COLORS['end']

# Initialize session state
def init_session_state():
    """Initialize session state"""
    defaults = {
        'current_language': 'en',
        'current_theme': 'light',
        'current_stage': 'start',
        'imported_style': None,
        'style_applied': False,
        'apply_imported_style': False,
        'output_text_value': "",
        'show_results': False,
        'download_data': {},
        'use_and_checkbox': False,
        'use_ampersand_checkbox': False,
        'journal_style': '{Full Journal Name}',
        'num': "No numbering",
        'auth': "AA Smith",
        'sep': ", ",
        'etal': 0,
        'doi': "https://doi.org/10.10/xxx",
        'doilink': True,
        'page': "122–128",
        'punct': "",
        'gost_style': False,
        'acs_style': False,
        'rsc_style': False,
        'cta_style': False,
        'last_style_update': 0,
        'cache_initialized': False,
        'user_prefs_loaded': False,
        'file_processing_complete': False,
        'style_import_processed': False,
        'last_imported_file_hash': None,
        'style_management_initialized': False,
        'previous_states': [],
        'max_undo_steps': 10,
        'stage_history': ['start'],
        'selected_preset': None,
        'custom_style_created': False,
        'style_config': None,
        'processing_start_time': None,
        'processing_results': None,
        'input_method': 'DOCX',
        'output_method': 'DOCX',
        'uploaded_file': None,
        'text_input': '',
        'style_export_name': 'my_citation_style',
        'show_statistics': False,
        'processing_complete': False,
        'duplicates_info': {},
        'doi_found_count': 0,
        'doi_not_found_count': 0,
        'formatted_refs': [],
        'txt_buffer': None,
        'docx_buffer': None,
        'formatted_txt_buffer': None,
        'selected_style_preview': None,
        'recommendations': None,
        'recommendations_loading': False,
        'recommendations_generated': False,
        'recommendations_metadata': None,
        'show_recommendations': False,
    }
    
    for key, default in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = default
    
    for i in range(8):
        for prop in ['el', 'it', 'bd', 'pr', 'sp']:
            key = f"{prop}{i}"
            if key not in st.session_state:
                if prop == 'sp':
                    st.session_state[key] = ". "
                elif prop == 'el':
                    st.session_state[key] = ""
                else:
                    st.session_state[key] = False

def get_text(key: str) -> str:
    """Get translation by key"""
    return TRANSLATIONS[st.session_state.current_language].get(key, key)

# Journal Abbreviation System
class JournalAbbreviation:
    def __init__(self):
        self.ltwa_data = {}
        self.load_ltwa_data()
        self.uppercase_abbreviations = {'acs', 'ecs', 'rsc', 'ieee', 'iet', 'acm', 'aims', 'bmc', 'bmj', 'npj'}
        self.special_endings = {'A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 
                               'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z',
                               'I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X'}
    
    def load_ltwa_data(self):
        """Load abbreviation data from ltwa.csv file"""
        try:
            csv_path = Config.LTWA_CSV_PATH
            if os.path.exists(csv_path):
                with open(csv_path, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f, delimiter='\t')
                    next(reader)
                    for row in reader:
                        if len(row) >= 2:
                            word = row[0].strip()
                            abbreviation = row[1].strip() if row[1].strip() else None
                            self.ltwa_data[word] = abbreviation
            else:
                logger.warning(f"File {csv_path} not found, using standard abbreviation")
        except Exception as e:
            logger.error(f"Error loading ltwa.csv: {e}")
    
    def abbreviate_word(self, word: str) -> str:
        """Abbreviate single word based on LTWA data"""
        word_lower = word.lower()
        
        if word_lower in self.ltwa_data:
            abbr = self.ltwa_data[word_lower]
            return abbr if abbr else word
        
        for ltwa_word, abbr in self.ltwa_data.items():
            if ltwa_word.endswith('-') and word_lower.startswith(ltwa_word[:-1]):
                return abbr if abbr else word
        
        return word
    
    def extract_special_endings(self, journal_name: str) -> Tuple[str, str]:
        """Extract special endings (A, B, C, etc.) from journal name"""
        patterns = [
            r'\s+([A-Z])\s*$',
            r'\s+([IVX]+)\s*$',
            r'\s+Part\s+([A-Z0-9]+)\s*$',
            r'\s+([A-Z]):\s+[A-Z]',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, journal_name)
            if match:
                ending = match.group(1)
                if ending in self.special_endings or re.match(r'^[A-Z]$', ending):
                    base_name = journal_name[:match.start()].strip()
                    return base_name, ending
        
        return journal_name, ""
    
    def abbreviate_journal_name(self, journal_name: str, style: str = "{J. Abbr.}") -> str:
        """Abbreviate journal name according to selected style"""
        if not journal_name:
            return ""
        
        base_name, special_ending = self.extract_special_endings(journal_name)
        
        words_to_remove = {'a', 'an', 'the', 'of', 'in', 'and', '&', 'for', 'on', 'with', 'by'}
        words = [word for word in base_name.split() if word.lower() not in words_to_remove]
        words = [word.replace(':', '') for word in words]
        
        if len(words) <= 1:
            result = journal_name
        else:
            abbreviated_words = []
            for i, word in enumerate(words):
                original_first_char = word[0]
                abbreviated = self.abbreviate_word(word.lower())
                
                if abbreviated and original_first_char.isupper():
                    abbreviated = abbreviated[0].upper() + abbreviated[1:]
                
                if i == 0 and abbreviated.lower() in self.uppercase_abbreviations:
                    abbreviated = abbreviated.upper()
                
                abbreviated_words.append(abbreviated)
            
            if style == "{J. Abbr.}":
                result = " ".join(abbreviated_words)
            elif style == "{J Abbr}":
                result = " ".join(abbr.replace('.', '') for abbr in abbreviated_words)
            else:
                result = base_name
        
        if special_ending:
            if ':' in journal_name and special_ending + ':' in journal_name:
                result += f" {special_ending}:"
                after_colon = journal_name.split(special_ending + ':', 1)[1].strip()
                if after_colon:
                    result += f" {after_colon}"
            else:
                result += f" {special_ending}"
        
        result = re.sub(r'\.\.+', '.', result)
        return result

# Initialize abbreviation system
journal_abbrev = JournalAbbreviation()

def clean_double_dots(text: str) -> str:
    """Remove double dots in text"""
    return re.sub(r'\.\.+', '.', text)

# Base Citation Formatter
class BaseCitationFormatter:
    """Base class for citation formatting"""
    
    def __init__(self, style_config: Dict[str, Any]):
        self.style_config = style_config
    
    def format_authors(self, authors: List[Dict[str, str]]) -> str:
        """Format authors list"""
        if not authors:
            return ""
        
        author_format = self.style_config['author_format']
        separator = self.style_config['author_separator']
        et_al_limit = self.style_config['et_al_limit']
        use_and_bool = self.style_config['use_and_bool']
        use_ampersand_bool = self.style_config['use_ampersand_bool']
        
        author_str = ""
        
        if use_and_bool or use_ampersand_bool:
            limit = len(authors)
        else:
            limit = et_al_limit if et_al_limit and et_al_limit > 0 else len(authors)
        
        for i, author in enumerate(authors[:limit]):
            given = author['given']
            family = author['family']
            
            initials = given.split()[:2]
            first_initial = initials[0][0] if initials else ''
            second_initial = initials[1][0].upper() if len(initials) > 1 else ''

            if author_format == "AA Smith":
                formatted_author = f"{first_initial}{second_initial} {family}"
            elif author_format == "A.A. Smith":
                if second_initial:
                    formatted_author = f"{first_initial}.{second_initial}. {family}"
                else:
                    formatted_author = f"{first_initial}. {family}"
            elif author_format == "Smith AA":
                formatted_author = f"{family} {first_initial}{second_initial}"
            elif author_format == "Smith A.A":
                if second_initial:
                    formatted_author = f"{family} {first_initial}.{second_initial}."
                else:
                    formatted_author = f"{family} {first_initial}"
            elif author_format == "Smith, A.A.":
                if second_initial:
                    formatted_author = f"{family}, {first_initial}.{second_initial}."
                else:
                    formatted_author = f"{family}, {first_initial}."
            elif author_format == "A A Smith":
                if second_initial:
                    formatted_author = f"{first_initial} {second_initial} {family}"
                else:
                    formatted_author = f"{first_initial} {family}"
            elif author_format == "A. A. Smith":
                if second_initial:
                    formatted_author = f"{first_initial}. {second_initial}. {family}"
                else:
                    formatted_author = f"{first_initial}. {family}"
            elif author_format == "Smith A A":
                if second_initial:
                    formatted_author = f"{family} {first_initial} {second_initial}"
                else:
                    formatted_author = f"{family} {first_initial}"
            elif author_format == "Smith A. A.":
                if second_initial:
                    formatted_author = f"{family} {first_initial}. {second_initial}."
                else:
                    formatted_author = f"{family} {first_initial}."
            elif author_format == "Smith A. A":
                if second_initial:
                    formatted_author = f"{family} {first_initial}. {second_initial}"
                else:
                    formatted_author = f"{family} {first_initial}"
            elif author_format == "Smith, A. A.":
                if second_initial:
                    formatted_author = f"{family}, {first_initial}. {second_initial}."
                else:
                    formatted_author = f"{family}, {first_initial}."
            else:
                formatted_author = f"{first_initial}. {family}"
            
            author_str += formatted_author
            
            if i < len(authors[:limit]) - 1:
                if i == len(authors[:limit]) - 2 and (use_and_bool or use_ampersand_bool):
                    if use_and_bool:
                        author_str += " and "
                    else:
                        author_str += " & "
                else:
                    author_str += separator
        
        if et_al_limit and len(authors) > et_al_limit and not (use_and_bool or use_ampersand_bool):
            author_str += " et al"
        
        return author_str.strip()
          
    def format_pages(self, pages: str, article_number: str, style_type: str = "default") -> str:
        """Format pages depending on style"""
        page_format = self.style_config['page_format']
        
        if pages:
            if style_type == "rsc":
                if '-' in pages:
                    first_page = pages.split('-')[0].strip()
                    return first_page
                else:
                    return pages.strip()
            elif style_type == "cta":
                if '-' in pages:
                    start, end = pages.split('-')
                    start = start.strip()
                    end = end.strip()
                    
                    if len(start) == len(end) and start[:-1] == end[:-1]:
                        return f"{start}–{end[-1]}"
                    elif len(start) > 1 and len(end) > 1 and start[:-2] == end[:-2]:
                        return f"{start}–{end[-2:]}"
                    else:
                        return f"{start}–{end}"
                else:
                    return pages.strip()
            else:
                if '-' not in pages:
                    if page_format == "122":
                        return pages.strip()
                    return pages.strip()
                
                start, end = pages.split('-')
                start = start.strip()
                end = end.strip()
                
                if page_format == "122 - 128":
                    return f"{start} - {end}"
                elif page_format == "122-128":
                    return f"{start}-{end}"
                elif page_format == "122 – 128":
                    return f"{start} – {end}"
                elif page_format == "122–128":
                    return f"{start}–{end}"
                elif page_format == "122–8":
                    i = 0
                    while i < len(start) and i < len(end) and start[i] == end[i]:
                        i += 1
                    return f"{start}–{end[i:]}"
                elif page_format == "122":
                    return start
        
        return article_number
    
    def format_doi(self, doi: str) -> Tuple[str, str]:
        """Format DOI and return text and URL"""
        doi_format = self.style_config['doi_format']
        
        if doi_format == "10.10/xxx":
            value = doi
        elif doi_format == "doi:10.10/xxx":
            value = f"doi:{doi}"
        elif doi_format == "DOI:10.10/xxx":
            value = f"DOI:{doi}"
        elif doi_format == "https://doi.org/10.10/xxx":
            value = f"https://doi.org/{doi}"
        else:
            value = doi
        
        return value, f"https://doi.org/{doi}"
    
    def format_journal_name(self, journal_name: str) -> str:
        """Format journal name considering selected style"""
        journal_style = self.style_config.get('journal_style', '{Full Journal Name}')
        return journal_abbrev.abbreviate_journal_name(journal_name, journal_style)

# Custom Citation Formatter
class CustomCitationFormatter(BaseCitationFormatter):
    """Formatter for custom styles with improved Issue handling"""
    
    def format_reference(self, metadata: Dict[str, Any], for_preview: bool = False) -> Tuple[Any, bool]:
        if not metadata:
            error_message = "Error: Could not format the reference." if st.session_state.current_language == 'en' else "Ошибка: Не удалось отформатировать ссылку."
            return (error_message, True)
        
        elements = []
        previous_element_was_empty = False
        
        for i, (element, config) in enumerate(self.style_config['elements']):
            value = ""
            doi_value = None
            element_empty = False
            
            if element == "Authors":
                value = self.format_authors(metadata['authors'])
                element_empty = not value
            elif element == "Title":
                value = metadata['title']
                element_empty = not value
            elif element == "Journal":
                value = self.format_journal_name(metadata['journal'])
                element_empty = not value
            elif element == "Year":
                value = str(metadata['year']) if metadata['year'] else ""
                element_empty = not value
            elif element == "Volume":
                value = metadata['volume']
                element_empty = not value
            elif element == "Issue":
                value = metadata['issue']
                element_empty = not value
            elif element == "Pages":
                value = self.format_pages(metadata['pages'], metadata['article_number'])
                element_empty = not value
            elif element == "DOI":
                doi = metadata['doi']
                doi_value = doi
                value, _ = self.format_doi(doi)
                element_empty = not value
            
            if value:
                if config['parentheses'] and value:
                    value = f"({value})"
                
                separator = ""
                if i < len(self.style_config['elements']) - 1:
                    if not element_empty:
                        separator = config['separator']
                    elif previous_element_was_empty:
                        separator = ""
                    else:
                        separator = config['separator']
                
                if for_preview:
                    formatted_value = value
                    if config['italic'] and config['bold']:
                        formatted_value = f"**_{formatted_value}_**"
                    elif config['italic']:
                        formatted_value = f"_{formatted_value}_"
                    elif config['bold']:
                        formatted_value = f"**{formatted_value}**"
                    
                    elements.append((formatted_value, config['italic'], config['bold'], separator, False, None, element_empty))
                else:
                    elements.append((value, config['italic'], config['bold'], separator,
                                   (element == "DOI" and self.style_config['doi_hyperlink']), doi_value, element_empty))
                
                previous_element_was_empty = False
            else:
                previous_element_was_empty = True
        
        cleaned_elements = []
        for i, element_data in enumerate(elements):
            value, italic, bold, separator, is_doi_hyperlink, doi_value, element_empty = element_data
            
            if not element_empty:
                if i == len(elements) - 1:
                    separator = ""
                
                cleaned_elements.append((value, italic, bold, separator, is_doi_hyperlink, doi_value))
        
        if for_preview:
            ref_str = ""
            for i, (value, _, _, separator, _, _) in enumerate(cleaned_elements):
                ref_str += value
                if separator and i < len(cleaned_elements) - 1:
                    ref_str += separator
                elif i == len(cleaned_elements) - 1 and self.style_config['final_punctuation']:
                    ref_str = ref_str.rstrip(',.') + "."
            
            ref_str = re.sub(r'\.\.+', '.', ref_str)
            return ref_str, False
        else:
            return cleaned_elements, False

# GOST Citation Formatter
class GOSTCitationFormatter(BaseCitationFormatter):
    """Formatter for GOST style (updated version)"""
    
    def format_reference(self, metadata: Dict[str, Any], for_preview: bool = False) -> Tuple[Any, bool]:
        if not metadata:
            error_message = "Error: Could not format the reference." if st.session_state.current_language == 'en' else "Ошибка: Не удалось отформатировать ссылку."
            return (error_message, True)
        
        authors_str = ""
        for i, author in enumerate(metadata['authors']):
            given = author['given']
            family = author['family']
            initials = given.split()[:2]
            first_initial = initials[0][0] if initials else ''
            second_initial = initials[1][0].upper() if len(initials) > 1 else ''
            
            if second_initial:
                author_str = f"{family} {first_initial}.{second_initial}."
            else:
                author_str = f"{family} {first_initial}."
            
            authors_str += author_str
            
            if i < len(metadata['authors']) - 1:
                authors_str += ", "
        
        pages = metadata['pages']
        article_number = metadata['article_number']
        
        journal_name = metadata['journal']
        
        doi_url = f"https://doi.org/{metadata['doi']}"
        
        if metadata['issue']:
            gost_ref = f"{authors_str} {metadata['title']} // {journal_name}. – {metadata['year']}. – Vol. {metadata['volume']}, № {metadata['issue']}"
        else:
            gost_ref = f"{authors_str} {metadata['title']} // {journal_name}. – {metadata['year']}. – Vol. {metadata['volume']}"
        
        if article_number and article_number.strip():
            gost_ref += f". – Art. {article_number.strip()}"
        elif pages and pages.strip():
            if '-' in pages:
                start_page, end_page = pages.split('-')
                pages_formatted = f"{start_page.strip()}-{end_page.strip()}"
            else:
                pages_formatted = pages.strip()
            gost_ref += f". – Р. {pages_formatted}"
        else:
            if st.session_state.current_language == 'ru':
                gost_ref += ". – [Без пагинации]"
            else:
                gost_ref += ". – [No pagination]"
        
        gost_ref += f". – {doi_url}"
        
        if for_preview:
            return gost_ref, False
        else:
            elements = []
            text_before_doi = gost_ref.replace(doi_url, "")
            elements.append((text_before_doi, False, False, "", False, None))
            elements.append((doi_url, False, False, "", True, metadata['doi']))
            return elements, False

# ACS Citation Formatter
class ACSCitationFormatter(BaseCitationFormatter):
    """Formatter for ACS (MDPI) style"""
    
    def format_reference(self, metadata: Dict[str, Any], for_preview: bool = False) -> Tuple[Any, bool]:
        if not metadata:
            error_message = "Error: Could not format the reference." if st.session_state.current_language == 'en' else "Ошибка: Не удалось отформатировать ссылку."
            return (error_message, True)
        
        authors_str = ""
        for i, author in enumerate(metadata['authors']):
            given = author['given']
            family = author['family']
            
            initials = given.split()[:2]
            first_initial = initials[0][0] if initials else ''
            second_initial = initials[1][0].upper() if len(initials) > 1 else ''
            
            if second_initial:
                author_str = f"{family}, {first_initial}.{second_initial}."
            else:
                author_str = f"{family}, {first_initial}."
            
            authors_str += author_str
            
            if i < len(metadata['authors']) - 1:
                authors_str += "; "
        
        pages = metadata['pages']
        article_number = metadata['article_number']
        
        if pages:
            if '-' in pages:
                start_page, end_page = pages.split('-')
                start_page = start_page.strip()
                end_page = end_page.strip()
                pages_formatted = f"{start_page}–{end_page}"
            else:
                pages_formatted = pages
        elif article_number:
            pages_formatted = article_number
        else:
            pages_formatted = ""
        
        journal_name = self.format_journal_name(metadata['journal'])
        
        doi_url = f"https://doi.org/{metadata['doi']}"
        
        acs_ref = f"{authors_str} {metadata['title']}. {journal_name} {metadata['year']}, {metadata['volume']}, {pages_formatted}. {doi_url}"
        acs_ref = re.sub(r'\.\.+', '.', acs_ref)
        
        if for_preview:
            return acs_ref, False
        else:
            elements = []
            elements.append((authors_str, False, False, " ", False, None))
            elements.append((metadata['title'], False, False, ". ", False, None))
            elements.append((journal_name, True, False, " ", False, None))
            elements.append((str(metadata['year']), False, True, ", ", False, None))
            elements.append((metadata['volume'], True, False, ", ", False, None))
            elements.append((pages_formatted, False, False, ". ", False, None))
            elements.append((doi_url, False, False, "", True, metadata['doi']))
            return elements, False

# RSC Citation Formatter
class RSCCitationFormatter(BaseCitationFormatter):
    """Formatter for RSC style"""
    
    def format_reference(self, metadata: Dict[str, Any], for_preview: bool = False) -> Tuple[Any, bool]:
        if not metadata:
            error_message = "Error: Could not format the reference." if st.session_state.current_language == 'en' else "Ошибка: Не удалось отформатировать ссылку."
            return (error_message, True)
        
        authors_str = ""
        for i, author in enumerate(metadata['authors']):
            given = author['given']
            family = author['family']
            
            initials = given.split()[:2]
            first_initial = initials[0][0] if initials else ''
            second_initial = initials[1][0].upper() if len(initials) > 1 else ''
            
            if second_initial:
                author_str = f"{first_initial}.{second_initial}. {family}"
            else:
                author_str = f"{first_initial}. {family}"
            
            authors_str += author_str
            
            if i < len(metadata['authors']) - 1:
                if i == len(metadata['authors']) - 2:
                    authors_str += " and "
                else:
                    authors_str += ", "
        
        pages = metadata['pages']
        article_number = metadata['article_number']
        
        if pages:
            if '-' in pages:
                first_page = pages.split('-')[0].strip()
                pages_formatted = first_page
            else:
                pages_formatted = pages.strip()
        elif article_number:
            pages_formatted = article_number
        else:
            pages_formatted = ""
        
        journal_name = self.format_journal_name(metadata['journal'])
        rsc_ref = f"{authors_str}, {journal_name}, {metadata['year']}, {metadata['volume']}, {pages_formatted}."
        rsc_ref = re.sub(r'\.\.+', '.', rsc_ref)
        
        if for_preview:
            return rsc_ref, False
        else:
            elements = []
            elements.append((authors_str, False, False, ", ", False, None))
            elements.append((journal_name, True, False, ", ", False, None))
            elements.append((str(metadata['year']), False, False, ", ", False, None))
            elements.append((metadata['volume'], False, True, ", ", False, None))
            elements.append((pages_formatted, False, False, ".", False, None))
            return elements, False

# CTA Citation Formatter
class CTACitationFormatter(BaseCitationFormatter):
    """Formatter for CTA style"""
    
    def format_reference(self, metadata: Dict[str, Any], for_preview: bool = False) -> Tuple[Any, bool]:
        if not metadata:
            error_message = "Error: Could not format the reference." if st.session_state.current_language == 'en' else "Ошибка: Не удалось отформатировать ссылку."
            return (error_message, True)
        
        authors_str = ""
        for i, author in enumerate(metadata['authors']):
            given = author['given']
            family = author['family']
            
            initials = given.split()[:2]
            first_initial = initials[0][0] if initials else ''
            second_initial = initials[1][0].upper() if len(initials) > 1 else ''
            
            if second_initial:
                author_str = f"{family} {first_initial}{second_initial}"
            else:
                author_str = f"{family} {first_initial}"
            
            authors_str += author_str
            
            if i < len(metadata['authors']) - 1:
                authors_str += ", "
        
        pages = metadata['pages']
        article_number = metadata['article_number']
        pages_formatted = self.format_pages(pages, article_number, "cta")
        journal_name = self.format_journal_name(metadata['journal'])
        issue_part = f"({metadata['issue']})" if metadata['issue'] else ""
        
        cta_ref = f"{authors_str}. {metadata['title']}. {journal_name}. {metadata['year']};{metadata['volume']}{issue_part}:{pages_formatted}. doi:{metadata['doi']}"
        
        if for_preview:
            return cta_ref, False
        else:
            elements = []
            elements.append((authors_str, False, False, ". ", False, None))
            elements.append((metadata['title'], False, False, ". ", False, None))
            elements.append((journal_name, True, False, ". ", False, None))
            elements.append((str(metadata['year']), False, False, ";", False, None))
            elements.append((metadata['volume'], False, False, "", False, None))
            if metadata['issue']:
                elements.append((f"({metadata['issue']})", False, False, ":", False, None))
            else:
                elements.append(("", False, False, ":", False, None))
            elements.append((pages_formatted, False, False, ". ", False, None))
            
            elements.append(("doi:", False, False, "", False, None))
            elements.append((metadata['doi'], False, False, "", True, metadata['doi']))
            
            return elements, False

# Style 5 Formatter
class Style5Formatter(BaseCitationFormatter):
    """Formatter for Style 5"""
    
    def format_reference(self, metadata: Dict[str, Any], for_preview: bool = False) -> Tuple[Any, bool]:
        if not metadata:
            error_message = "Error: Could not format the reference." if st.session_state.current_language == 'en' else "Ошибка: Не удалось отформатировать ссылку."
            return (error_message, True)
        
        authors_str = ""
        for i, author in enumerate(metadata['authors']):
            given = author['given']
            family = author['family']
            
            initials = given.split()[:2]
            first_initial = initials[0][0] if initials else ''
            second_initial = initials[1][0].upper() if len(initials) > 1 else ''
            
            if second_initial:
                author_str = f"{first_initial}.{second_initial}. {family}"
            else:
                author_str = f"{first_initial}. {family}"
            
            authors_str += author_str
            
            if i < len(metadata['authors']) - 1:
                authors_str += ", "
        
        journal_name = self.format_journal_name(metadata['journal'])
        
        doi_url = f"https://doi.org/{metadata['doi']}"
        
        pages = metadata.get('pages', '')
        article_number = metadata.get('article_number', '')
        
        # Сначала проверяем наличие страниц
        if pages and pages.strip():
            if '-' in pages:
                start_page, end_page = pages.split('-')
                pages_formatted = f"{start_page}–{end_page}"
            else:
                pages_formatted = pages
        # Если страниц нет, но есть номер статьи, используем его
        elif article_number and article_number.strip():
            pages_formatted = article_number
        else:
            pages_formatted = ""
        
        style5_ref = f"{authors_str}, {metadata['title']}, {journal_name} {metadata['volume']} ({metadata['year']}) {pages_formatted}. {doi_url}"
        
        if for_preview:
            return style5_ref, False
        else:
            elements = []
            elements.append((authors_str, False, False, ", ", False, None))
            elements.append((metadata['title'], False, False, ", ", False, None))
            elements.append((journal_name, False, False, " ", False, None))
            elements.append((metadata['volume'], False, False, " (", False, None))
            elements.append((str(metadata['year']), False, False, ") ", False, None))
            elements.append((pages_formatted, False, False, " ", False, None))
            elements.append((doi_url, False, False, "", True, metadata['doi']))
            return elements, False

# Style 6 Formatter
class Style6Formatter(BaseCitationFormatter):
    """Formatter for Style 6"""
    
    def format_reference(self, metadata: Dict[str, Any], for_preview: bool = False) -> Tuple[Any, bool]:
        if not metadata:
            error_message = "Error: Could not format the reference." if st.session_state.current_language == 'en' else "Ошибка: Не удалось отформатировать ссылку."
            return (error_message, True)
        
        authors_str = ""
        for i, author in enumerate(metadata['authors']):
            given = author['given']
            family = author['family']
            
            initials = given.split()[:2]
            first_initial = initials[0][0] if initials else ''
            second_initial = initials[1][0].upper() if len(initials) > 1 else ''
            
            if second_initial:
                author_str = f"{family}, {first_initial}.{second_initial}."
            else:
                author_str = f"{family}, {first_initial}."
            
            authors_str += author_str
            
            if i < len(metadata['authors']) - 1:
                authors_str += ", "
        
        journal_name = metadata['journal']
        
        doi_url = f"https://doi.org/{metadata['doi']}"

        pages = metadata['pages']
        article_number = metadata.get('article_number', '')
        if pages:
            if '-' in pages:
                start_page, end_page = pages.split('-')
                pages_formatted = f"{start_page}–{end_page}"
            else:
                pages_formatted = pages
        elif article_number:
            pages_formatted = article_number
        else:
            pages_formatted = ""
        
        style6_ref = f"{authors_str} ({metadata['year']}). {metadata['title']}. {journal_name} {metadata['volume']}, {pages_formatted}. {doi_url}."
        
        if for_preview:
            return style6_ref, False
        else:
            elements = []
            elements.append((authors_str, False, False, " (", False, None))
            elements.append((str(metadata['year']), False, False, "). ", False, None))
            elements.append((metadata['title'], False, False, ". ", False, None))
            elements.append((journal_name, False, False, " ", False, None))
            elements.append((metadata['volume'], True, False, ", ", False, None))
            elements.append((pages_formatted, False, False, ". ", False, None))
            elements.append((doi_url, False, False, "", True, metadata['doi']))
            return elements, False

# Style 7 Formatter
class Style7Formatter(BaseCitationFormatter):
    """Formatter for Style 7"""
    
    def format_reference(self, metadata: Dict[str, Any], for_preview: bool = False) -> Tuple[Any, bool]:
        if not metadata:
            error_message = "Error: Could not format the reference." if st.session_state.current_language == 'en' else "Ошибка: Не удалось отформатировать ссылку."
            return (error_message, True)
        
        authors_str = ""
        for i, author in enumerate(metadata['authors']):
            given = author['given']
            family = author['family']
            
            initials = given.split()[:2]
            first_initial = initials[0][0] if initials else ''
            second_initial = initials[1][0].upper() if len(initials) > 1 else ''
            
            if second_initial:
                author_str = f"{family}, {first_initial}.{second_initial}."
            else:
                author_str = f"{family}, {first_initial}."
            
            authors_str += author_str
            
            if i < len(metadata['authors']) - 1:
                if i == len(metadata['authors']) - 2:
                    authors_str += " & "
                else:
                    authors_str += ", "
        
        journal_name = metadata['journal']
        
        doi_url = f"https://doi.org/{metadata['doi']}"
        
        pages = metadata['pages']
        article_number = metadata.get('article_number', '')
        if pages:
            if '-' in pages:
                start_page, end_page = pages.split('-')
                pages_formatted = f"{start_page}–{end_page}"
            else:
                pages_formatted = pages
        elif article_number:
            pages_formatted = article_number
        else:
            pages_formatted = ""
        
        issue_part = f"({metadata['issue']}), " if metadata['issue'] else ""
        
        style7_ref = f"{authors_str} ({metadata['year']}). {metadata['title']}. {journal_name} {metadata['volume']}{issue_part}{pages_formatted}. {doi_url}."
        
        if for_preview:
            return style7_ref, False
        else:
            elements = []
            elements.append((authors_str, False, False, " (", False, None))
            elements.append((str(metadata['year']), False, False, "). ", False, None))
            elements.append((metadata['title'], False, False, ". ", False, None))
            elements.append((journal_name, True, False, " ", False, None))
            elements.append((metadata['volume'], True, False, "", False, None))
            if metadata['issue']:
                elements.append((f"({metadata['issue']})", False, False, ", ", False, None))
            else:
                elements.append(("", False, False, ", ", False, None))
            elements.append((pages_formatted, False, False, ". ", False, None))
            elements.append((doi_url, False, False, "", True, metadata['doi']))
            return elements, False

# Style 8 Formatter
class Style8Formatter(BaseCitationFormatter):
    """Formatter for Style 8"""
    
    def format_reference(self, metadata: Dict[str, Any], for_preview: bool = False) -> Tuple[Any, bool]:
        if not metadata:
            error_message = "Error: Could not format the reference." if st.session_state.current_language == 'en' else "Ошибка: Не удалось отформатировать ссылку."
            return (error_message, True)
        
        authors_str = ""
        for i, author in enumerate(metadata['authors']):
            given = author['given']
            family = author['family']
            
            initials = given.split()[:2]
            first_initial = initials[0][0] if initials else ''
            second_initial = initials[1][0].upper() if len(initials) > 1 else ''
            
            if second_initial:
                author_str = f"{first_initial}. {second_initial}. {family}"
            else:
                author_str = f"{first_initial}. {family}"
            
            authors_str += author_str
            
            if i < len(metadata['authors']) - 1:
                authors_str += ", "
        
        journal_name = self.format_journal_name(metadata['journal'])

        pages = metadata['pages']
        article_number = metadata.get('article_number', '')
        if pages:
            if '-' in pages:
                first_page = pages.split('-')[0].strip()
                pages_formatted = first_page
            else:
                pages_formatted = pages.strip()
        elif article_number:
            pages_formatted = article_number
        else:
            pages_formatted = ""
        
        style8_ref = f"{authors_str}, {journal_name} {metadata['year']}, {metadata['volume']}, {pages_formatted}."
        
        if for_preview:
            return style8_ref, False
        else:
            elements = []
            elements.append((authors_str, False, False, ", ", False, None))
            elements.append((journal_name, True, False, " ", False, None))
            elements.append((str(metadata['year']), True, False, ", ", False, None))
            elements.append((metadata['volume'], False, True, ", ", False, None))
            elements.append((pages_formatted, False, False, ".", False, None))
            return elements, False

# Style 9 Formatter
class Style9Formatter(BaseCitationFormatter):
    """Formatter for Style 9 (RCR)"""
    
    def format_reference(self, metadata: Dict[str, Any], for_preview: bool = False) -> Tuple[Any, bool]:
        if not metadata:
            error_message = "Error: Could not format the reference." if st.session_state.current_language == 'en' else "Ошибка: Не удалось отформатировать ссылку."
            return (error_message, True)
        
        authors_str = ""
        for i, author in enumerate(metadata['authors']):
            given = author['given']
            family = author['family']
            
            initials = given.split()[:2]
            first_initial = initials[0][0] if initials else ''
            second_initial = initials[1][0].upper() if len(initials) > 1 else ''
            
            if second_initial:
                author_str = f"{first_initial}.{second_initial}.{family}"
            else:
                author_str = f"{first_initial}.{family}"
            
            authors_str += author_str
            
            if i < len(metadata['authors']) - 1:
                authors_str += ", "
        
        journal_name = self.format_journal_name(metadata['journal'])
        
        pages = metadata['pages']
        if pages:
            if '-' in pages:
                first_page = pages.split('-')[0].strip()
                pages_formatted = first_page
            else:
                pages_formatted = pages.strip()
        else:
            pages_formatted = ""
        
        doi_url = f"https://doi.org/{metadata['doi']}"
        
        style9_ref = f"{authors_str}. {journal_name}, {metadata['volume']}, {pages_formatted} ({metadata['year']}); {doi_url}"
        
        if for_preview:
            return style9_ref, False
        else:
            elements = []
            elements.append((authors_str, False, False, ". ", False, None))
            elements.append((journal_name, True, False, ", ", False, None))
            elements.append((metadata['volume'], False, True, ", ", False, None))
            elements.append((pages_formatted, False, False, " (", False, None))
            elements.append((str(metadata['year']), False, False, "); ", False, None))
            elements.append((doi_url, False, False, "", True, metadata['doi']))
            return elements, False

# Style 10 Formatter
class Style10Formatter(BaseCitationFormatter):
    """Formatter for Style 10"""
    
    def format_reference(self, metadata: Dict[str, Any], for_preview: bool = False) -> Tuple[Any, bool]:
        if not metadata:
            error_message = "Error: Could not format the reference." if st.session_state.current_language == 'en' else "Ошибка: Не удалось отформатировать ссылку."
            return (error_message, True)
        
        authors_str = ""
        for i, author in enumerate(metadata['authors']):
            given = author['given']
            family = author['family']
            
            initials = given.split()[:2]
            first_initial = initials[0][0] if initials else ''
            second_initial = initials[1][0].upper() if len(initials) > 1 else ''
            
            if second_initial:
                author_str = f"{family} {first_initial}{second_initial}"
            else:
                author_str = f"{family} {first_initial}"
            
            authors_str += author_str
            
            if i < len(metadata['authors']) - 1:
                authors_str += ", "
        
        journal_name = self.format_journal_name(metadata['journal'])
        
        doi_url = f"https://doi.org/{metadata['doi']}"
        
        pages = metadata['pages']
        if pages:
            if '-' in pages:
                start_page, end_page = pages.split('-')
                pages_formatted = f"{start_page}–{end_page}"
            else:
                pages_formatted = pages
        else:
            pages_formatted = ""
        
        issue_part = f"({metadata['issue']}):" if metadata['issue'] else ""
        
        style10_ref = f"{authors_str} ({metadata['year']}) {metadata['title']}. {journal_name} {metadata['volume']}{issue_part}{pages_formatted}. {doi_url}"
        
        if for_preview:
            return style10_ref, False
        else:
            elements = []
            elements.append((authors_str, False, False, " (", False, None))
            elements.append((str(metadata['year']), False, False, ") ", False, None))
            elements.append((metadata['title'], False, False, ". ", False, None))
            elements.append((journal_name, False, False, " ", False, None))
            elements.append((metadata['volume'], False, False, "", False, None))
            if metadata['issue']:
                elements.append((f"({metadata['issue']})", False, False, ":", False, None))
            else:
                elements.append(("", False, False, ":", False, None))
            elements.append((pages_formatted, False, False, ". ", False, None))
            elements.append((doi_url, False, False, "", True, metadata['doi']))
            return elements, False

# Citation Formatter Factory
class CitationFormatterFactory:
    """Factory for creating citation formatters"""
    
    @staticmethod
    def create_formatter(style_config: Dict[str, Any]) -> BaseCitationFormatter:
        if style_config.get('gost_style', False):
            return GOSTCitationFormatter(style_config)
        elif style_config.get('acs_style', False):
            return ACSCitationFormatter(style_config)
        elif style_config.get('rsc_style', False):
            return RSCCitationFormatter(style_config)
        elif style_config.get('cta_style', False):
            return CTACitationFormatter(style_config)
        elif style_config.get('style5', False):
            return Style5Formatter(style_config)
        elif style_config.get('style6', False):
            return Style6Formatter(style_config)
        elif style_config.get('style7', False):
            return Style7Formatter(style_config)
        elif style_config.get('style8', False):
            return Style8Formatter(style_config)
        elif style_config.get('style9', False):
            return Style9Formatter(style_config)
        elif style_config.get('style10', False):
            return Style10Formatter(style_config)
        else:
            return CustomCitationFormatter(style_config)

class SimpleTopicAnalyzer:
    """Упрощенный анализатор тем по DOI"""
    
    def __init__(self):
        self.session = requests.Session()
        self.session.timeout = 30
        self.headers = {'User-Agent': 'CitationStyleConstructor/1.0'}
        
        # Список стоп-слов
        self.stopwords = {
            'study', 'research', 'paper', 'article', 'review', 'analysis', 'investigation',
            'effect', 'property', 'performance', 'behavior', 'behaviour', 'synthesis',
            'development', 'method', 'approach', 'result', 'discussion', 'conclusion',
            'introduction', 'experimental', 'measurement', 'technique', 'material',
            'system', 'process', 'structure', 'model', 'based', 'using', 'different',
            'various', 'novel', 'new', 'recent', 'potential', 'important', 'significant',
            'high', 'low', 'good', 'better', 'strong', 'weak', 'large', 'small'
        }
        
        # Исключения для множественного числа
        self.plural_exceptions = {
            'analyses': 'analysis', 'properties': 'property', 'materials': 'material',
            'structures': 'structure', 'composites': 'composite', 'nanoparticles': 'nanoparticle',
            'catalysts': 'catalyst', 'electrodes': 'electrode', 'polymers': 'polymer',
            'sensors': 'sensor', 'devices': 'device', 'systems': 'system', 'cells': 'cell',
            'membranes': 'membrane', 'electrolytes': 'electrolyte', 'cathodes': 'cathode',
            'anodes': 'anode', 'activities': 'activity', 'efficiencies': 'efficiency',
            'capacities': 'capacity', 'performances': 'performance', 'stabilities': 'stability'
        }
    
    def normalize_word(self, word):
        """Приводит слово к базовой форме"""
        if not word or len(word) < 4:
            return ''
        
        word_lower = word.lower()
        
        # Проверяем исключения
        if word_lower in self.plural_exceptions:
            return self.plural_exceptions[word_lower]
        
        # Общие правила для английского языка
        if word_lower.endswith('ies'):
            base = word_lower[:-3] + 'y'
            if len(base) >= 4:
                return base
        elif word_lower.endswith('es'):
            # Проверяем особые случаи
            if word_lower.endswith(('ches', 'shes', 'xes', 'zes', 'sses')):
                base = word_lower[:-2]
                if len(base) >= 4:
                    return base
            elif word_lower.endswith('oes'):
                base = word_lower[:-2]
                if len(base) >= 4:
                    return base
        elif word_lower.endswith('s') and not word_lower.endswith(('ss', 'us', 'is', 'ys', 'as')):
            base = word_lower[:-1]
            if len(base) >= 4:
                return base
        
        return word_lower
    
    def extract_keywords_from_title(self, title):
        """Извлекает ключевые слова из заголовка"""
        if not title:
            return []
        
        # Токенизация - ищем слова из 4+ букв
        words = re.findall(r'\b[a-zA-Z]{4,}\b', title)
        
        filtered_words = []
        for word in words:
            word_lower = word.lower()
            
            # Пропускаем стоп-слова
            if word_lower in self.stopwords:
                continue
            
            # Пропускаем числа
            if re.search(r'\d', word_lower):
                continue
            
            # Нормализуем слово
            normalized = self.normalize_word(word_lower)
            
            if normalized and len(normalized) >= 4:
                filtered_words.append(normalized)
        
        return filtered_words
    
    def fetch_work_data(self, doi):
        """Получает данные статьи по DOI из OpenAlex"""
        try:
            # Очищаем DOI
            clean_doi = re.sub(r'^(https?://doi\.org/|doi:|DOI:?\s*)', '', doi.strip(), flags=re.IGNORECASE)
            
            # Пробуем разные форматы
            for fmt in [clean_doi, f"doi:{clean_doi}", f"https://doi.org/{clean_doi}"]:
                try:
                    url = f"https://api.openalex.org/works/{fmt}"
                    response = self.session.get(url, headers=self.headers, timeout=20)
                    
                    if response.status_code == 200:
                        data = response.json()
                        
                        # Извлекаем основную тему
                        primary_topic = None
                        topics = data.get('topics', [])
                        if topics:
                            sorted_topics = sorted(topics, key=lambda x: x.get('score', 0), reverse=True)
                            primary_topic = sorted_topics[0]
                        
                        return {
                            'doi': doi,
                            'success': True,
                            'data': data,
                            'primary_topic': primary_topic
                        }
                        
                except Exception as e:
                    continue
            
            return {
                'doi': doi,
                'success': False,
                'error': "Failed to fetch data",
                'data': None,
                'primary_topic': None
            }
                
        except Exception as e:
            return {
                'doi': doi,
                'success': False,
                'error': str(e),
                'data': None,
                'primary_topic': None
            }
    
    def analyze_dois(self, dois, progress_callback=None):
        """Анализирует список DOI и извлекает темы"""
        if not dois:
            return None
        
        try:
            works_data = []
            all_titles = []
            topic_counter = Counter()
            
            total = len(dois)
            
            # Используем ThreadPoolExecutor для параллельной обработки
            with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
                futures = {executor.submit(self.fetch_work_data, doi): doi for doi in dois}
                
                for i, future in enumerate(concurrent.futures.as_completed(futures), 1):
                    result = future.result()
                    
                    if result['success']:
                        data = result['data']
                        primary_topic = result['primary_topic']
                        
                        if primary_topic:
                            topic_name = primary_topic.get('display_name', 'Unknown')
                            topic_counter[topic_name] += 1
                            
                            topic_id_full = primary_topic.get('id', '')
                            topic_id = topic_id_full.split('/')[-1] if topic_id_full else ''
                            
                            works_data.append({
                                'doi': result['doi'],
                                'title': data.get('title', ''),
                                'primary_topic': topic_name,
                                'topic_id': topic_id,
                                'topic_id_full': topic_id_full,
                                'cited_by_count': data.get('cited_by_count', 0),
                                'publication_date': data.get('publication_date', ''),
                                'publication_year': data.get('publication_year', '')
                            })
                            
                            title = data.get('title')
                            if title:
                                all_titles.append(title)
                    
                    # Обновляем прогресс
                    if progress_callback:
                        progress_val = int((i / total) * 50)
                        progress_callback(progress_val, f"Обработано {i}/{total} DOI")
            
            if not works_data:
                return None
            
            # Анализируем ключевые слова из заголовков
            keyword_counter = Counter()
            for title in all_titles:
                keywords = self.extract_keywords_from_title(title)
                if keywords:
                    keyword_counter.update(keywords)
            
            # Сортируем темы по частоте
            sorted_topics = topic_counter.most_common()
            
            # Считаем низкоцитируемые работы (<11 цитирований)
            low_citation_count = sum(1 for w in works_data if w.get('cited_by_count', 0) < 11)
            
            result = {
                'works_data': works_data,
                'topics': sorted_topics,
                'keywords': keyword_counter.most_common(20),
                'total_works': len(works_data),
                'low_citation_count': low_citation_count,
                'stats': {
                    'successful_doi': len(works_data),
                    'total_topics': len(sorted_topics),
                    'unique_titles': len(all_titles),
                    'avg_citations': sum(w.get('cited_by_count', 0) for w in works_data) / len(works_data) if works_data else 0
                }
            }
            
            return result
            
        except Exception as e:
            logger.error(f"Error in analyze_dois: {e}")
            return None
    
    def analyze_dois_parallel(self, dois, progress_callback=None):
        """Параллельный анализ списка DOI"""
        if not dois:
            return None
        
        try:
            works_data = []
            all_titles = []
            topic_counter = Counter()
            
            total = len(dois)
            
            # ПАРАЛЛЕЛЬНАЯ загрузка данных
            with concurrent.futures.ThreadPoolExecutor(max_workers=Config.OPENALEX_MAX_WORKERS) as executor:
                futures = {executor.submit(self.fetch_work_data, doi): doi for doi in dois}
                
                for i, future in enumerate(concurrent.futures.as_completed(futures), 1):
                    result = future.result()
                    
                    if result['success']:
                        data = result['data']
                        primary_topic = result['primary_topic']
                        
                        if primary_topic:
                            topic_name = primary_topic.get('display_name', 'Unknown')
                            topic_counter[topic_name] += 1
                            
                            topic_id_full = primary_topic.get('id', '')
                            topic_id = topic_id_full.split('/')[-1] if topic_id_full else ''
                            
                            works_data.append({
                                'doi': result['doi'],
                                'title': data.get('title', ''),
                                'primary_topic': topic_name,
                                'topic_id': topic_id,
                                'topic_id_full': topic_id_full,
                                'cited_by_count': data.get('cited_by_count', 0),
                                'publication_date': data.get('publication_date', ''),
                                'publication_year': data.get('publication_year', '')
                            })
                            
                            title = data.get('title')
                            if title:
                                all_titles.append(title)
                    
                    # Обновляем прогресс
                    if progress_callback:
                        progress_val = int((i / total) * 50)
                        progress_callback(progress_val, f"Обработано {i}/{total} DOI")
            
            if not works_data:
                return None
            
            # ПАРАЛЛЕЛЬНЫЙ анализ ключевых слов
            if progress_callback:
                progress_callback(50, "Параллельный анализ ключевых слов...")
            
            keyword_counter = Counter()
            
            with concurrent.futures.ThreadPoolExecutor(max_workers=Config.OPENALEX_MAX_WORKERS) as executor:
                futures = {executor.submit(self.extract_keywords_from_title, title): title for title in all_titles}
                
                for future in concurrent.futures.as_completed(futures):
                    keywords = future.result()
                    if keywords:
                        keyword_counter.update(keywords)
            
            # Сортируем темы по частоте
            sorted_topics = topic_counter.most_common()
            
            # Быстрый подсчет низкоцитируемых работ
            low_citation_count = sum(1 for w in works_data if w.get('cited_by_count', 0) < 11)
            
            result = {
                'works_data': works_data,
                'topics': sorted_topics,
                'keywords': keyword_counter.most_common(15),  # 15 вместо 20
                'total_works': len(works_data),
                'low_citation_count': low_citation_count,
                'stats': {
                    'successful_doi': len(works_data),
                    'total_topics': len(sorted_topics),
                    'unique_titles': len(all_titles),
                    'avg_citations': sum(w.get('cited_by_count', 0) for w in works_data) / len(works_data) if works_data else 0
                }
            }
            
            return result
            
        except Exception as e:
            logger.error(f"Error in analyze_dois_parallel: {e}")
            return None
    
class LowCitationFinder:
    """Оптимизированный поиск низкоцитируемых статей по теме"""
    
    def __init__(self):
        self.session = requests.Session()
        self.session.timeout = Config.OPENALEX_REQUEST_TIMEOUT
        self.headers = {'User-Agent': 'CitationStyleConstructor/1.0'}
        self.topic_cache = {}  # Кэш для данных тем
        self.works_cache = {}  # Кэш для работ по темам
        
    def _make_request(self, url):
        """Оптимизированный HTTP запрос с кэшированием"""
        cache_key = hashlib.md5(url.encode()).hexdigest()
        
        # Проверяем кэш
        if cache_key in self.topic_cache:
            cached_data, timestamp = self.topic_cache[cache_key]
            if time.time() - timestamp < Config.OPENALEX_CACHE_TTL_MINUTES * 60:
                return cached_data
        
        try:
            response = self.session.get(url, headers=self.headers, timeout=Config.OPENALEX_REQUEST_TIMEOUT)
            if response.status_code == 200:
                data = response.json()
                # Сохраняем в кэш
                self.topic_cache[cache_key] = (data, time.time())
                return data
        except Exception as e:
            logger.error(f"Request error for {url}: {e}")
        
        return None
    
    def fetch_works_by_topic_parallel(self, topic_id, max_results=Config.OPENALEX_MAX_WORKS_PER_TOPIC):
        """Параллельная загрузка работ по теме с оптимизированными параметрами"""
        if not topic_id:
            return []
        
        # Проверяем кэш
        cache_key = f"works_{topic_id}_{max_results}"
        if cache_key in self.works_cache:
            cached_data, timestamp = self.works_cache[cache_key]
            if time.time() - timestamp < Config.OPENALEX_CACHE_TTL_MINUTES * 60:
                return cached_data
        
        print(f"  📥 Загружаю работы по теме ID: {topic_id}")
        
        all_works = []
        
        try:
            # Создаем URL для первого запроса с СОРТИРОВКОЙ по дате (КРИТИЧНО!)
            base_url = f"https://api.openalex.org/works?filter=topics.id:{topic_id}"
            base_url += f"&per-page={Config.OPENALEX_PER_PAGE}"
            base_url += "&sort=publication_date:desc"  # ДОБАВЛЕНО: сортировка по дате!
            
            # Параллельно загружаем несколько страниц сразу
            with concurrent.futures.ThreadPoolExecutor(max_workers=Config.OPENALEX_MAX_WORKERS) as executor:
                futures = []
                
                for page in range(1, Config.OPENALEX_MAX_PAGES + 1):
                    url = f"{base_url}&page={page}"
                    future = executor.submit(self._make_request, url)
                    futures.append((future, page))
                
                for future, page in futures:
                    try:
                        data = future.result(timeout=Config.OPENALEX_REQUEST_TIMEOUT)
                        if data and 'results' in data:
                            works = data['results']
                            all_works.extend(works)
                            print(f"    📄 Страница {page}: получено {len(works)} работ")
                            
                            # Если получили меньше работ, чем ожидали, выходим
                            if len(works) < Config.OPENALEX_PER_PAGE:
                                break
                        else:
                            break
                            
                    except Exception as e:
                        print(f"    ⚠️ Ошибка на странице {page}: {e}")
                        continue
                    
                    # Если достигли лимита, выходим
                    if len(all_works) >= max_results:
                        break
            
            print(f"  ✅ Всего загружено {len(all_works)} работ по теме")
            
            # Сохраняем в кэш
            self.works_cache[cache_key] = (all_works, time.time())
            
            return all_works
            
        except Exception as e:
            logger.error(f"Error in fetch_works_by_topic_parallel for topic {topic_id}: {e}")
            return []
    
    def get_topic_info(self, topic_id):
        """Получает информацию о теме с кэшированием"""
        if not topic_id:
            return None
        
        cache_key = f"topic_info_{topic_id}"
        if cache_key in self.topic_cache:
            cached_data, timestamp = self.topic_cache[cache_key]
            if time.time() - timestamp < Config.OPENALEX_CACHE_TTL_MINUTES * 60:
                return cached_data
        
        try:
            url = f"https://api.openalex.org/topics/{topic_id}"
            data = self._make_request(url)
            
            if data:
                # Сохраняем в кэш
                self.topic_cache[cache_key] = (data, time.time())
            
            return data
        except Exception as e:
            logger.error(f"Error getting topic info for {topic_id}: {e}")
            return None
    
    def find_low_citation_works(self, topic_id, keywords, max_citations=10, max_works=Config.OPENALEX_MAX_TOTAL_WORKS):
        """Оптимизированный поиск низкоцитируемых работ по теме"""
        try:
            if not topic_id or not isinstance(topic_id, str):
                return []
            
            print(f"  🎯 Поиск низкоцитируемых работ для темы {topic_id}")
            
            # Получаем работы по теме ПАРАЛЛЕЛЬНО
            works = self.fetch_works_by_topic_parallel(topic_id, max_results=max_works)
            
            if not works:
                print(f"  ⚠️ Не найдено работ по теме {topic_id}")
                return []
            
            print(f"  🔍 Анализ {len(works)} работ на соответствие ключевым словам...")
            
            low_citation_works = []
            keywords_lower = [kw.lower() for kw in keywords] if keywords else []
            
            # Используем comprehension для ускорения фильтрации
            for work in works:
                # Быстрая проверка цитирований
                cited_by_count = work.get('cited_by_count', 0)
                
                if cited_by_count <= max_citations:
                    title = work.get('title', '')
                    if not title:
                        continue
                    
                    title_lower = title.lower()
                    
                    # Проверяем соответствие ключевым словам (если они есть)
                    score = 0
                    matched_keywords = []
                    
                    if keywords_lower:
                        for keyword in keywords_lower:
                            if keyword and keyword in title_lower:
                                score += 1
                                matched_keywords.append(keyword)
                    
                    # Если есть ключевые слова - нужен хотя бы 1 match
                    # Если ключевых слов нет - берем все низкоцитируемые
                    if not keywords_lower or score > 0:
                        # Извлекаем данные
                        authors = []
                        try:
                            authorships = work.get('authorships', [])
                            for authorship in authorships[:3]:  # Первые 3 автора
                                author = authorship.get('author', {})
                                display_name = author.get('display_name', '')
                                if display_name:
                                    authors.append(display_name)
                        except:
                            pass
                        
                        journal = ''
                        try:
                            primary_location = work.get('primary_location', {})
                            if primary_location:
                                source = primary_location.get('source', {})
                                journal = source.get('display_name', '')
                        except:
                            pass
                        
                        doi = work.get('doi', '')
                        if doi and doi.startswith('https://doi.org/'):
                            doi = doi[16:]
                        
                        low_citation_works.append({
                            'title': title,
                            'relevance_score': score,
                            'matched_keywords': matched_keywords,
                            'cited_by_count': cited_by_count,
                            'publication_date': work.get('publication_date', ''),
                            'publication_year': work.get('publication_year', ''),
                            'doi': doi,
                            'authors': authors,
                            'journal': journal,
                            'openalex_id': work.get('id', ''),
                            'is_oa': work.get('open_access', {}).get('is_oa', False)
                        })
            
            # Оптимизированная сортировка
            if keywords_lower:
                low_citation_works.sort(key=lambda x: (x['relevance_score'], -x['cited_by_count']), reverse=True)
            else:
                low_citation_works.sort(key=lambda x: (-x['cited_by_count']), reverse=True)
            
            # Статистика
            zero_citation = sum(1 for w in low_citation_works if w['cited_by_count'] == 0)
            few_citation = sum(1 for w in low_citation_works if 1 <= w['cited_by_count'] <= 5)
            
            print(f"  📊 Найдено низкоцитируемых работ: {len(low_citation_works)}")
            print(f"     • С 0 цитирований: {zero_citation}")
            print(f"     • С 1-5 цитирований: {few_citation}")
            print(f"     • С 6-10 цитирований: {len(low_citation_works) - zero_citation - few_citation}")
            
            # Возвращаем топ-рекомендации
            return low_citation_works[:Config.OPENALEX_MAX_WORKS_PER_TOPIC]
            
        except Exception as e:
            logger.error(f"Error in find_low_citation_works: {e}")
            print(f"  ❌ Ошибка: {str(e)[:100]}")
            return []

class ArticleRecommender:

    @staticmethod
    def generate_recommendations(formatted_refs, progress_callback=None):
        """Оптимизированная генерация рекомендаций"""
        try:
            # Проверяем минимальное количество ссылок
            if len(formatted_refs) < Config.MIN_REFERENCES_FOR_RECOMMENDATIONS:
                if progress_callback:
                    progress_callback(100, "Недостаточно ссылок для рекомендаций")
                return None
            
            if progress_callback:
                progress_callback(5, "Извлекаю DOI из ссылок...")
            
            # Извлекаем DOI (оптимизированная версия)
            dois = []
            for _, is_error, metadata in formatted_refs:
                if not is_error and metadata and metadata.get('doi'):
                    doi = metadata['doi']
                    # Быстрая нормализация DOI
                    if doi.startswith('https://doi.org/'):
                        doi = doi[16:]
                    elif doi.startswith('doi:'):
                        doi = doi[4:]
                    dois.append(doi.strip())
            
            dois = list(set(dois))  # Уникальные DOI
            
            if not dois:
                if progress_callback:
                    progress_callback(100, "Не найдено DOI в ссылках")
                return None
            
            if progress_callback:
                progress_callback(10, f"Найдено {len(dois)} уникальных DOI, начинаю анализ...")
            
            # Создаем анализатор
            analyzer = SimpleTopicAnalyzer()
            
            # ПАРАЛЛЕЛЬНЫЙ анализ DOI
            if progress_callback:
                progress_callback(15, "Параллельный анализ DOI через OpenAlex...")
            
            analysis_result = analyzer.analyze_dois_parallel(dois, progress_callback)
            
            if not analysis_result or not analysis_result.get('topics'):
                if progress_callback:
                    progress_callback(100, "Не удалось проанализировать темы")
                return None
            
            topics = analysis_result['topics']
            
            if progress_callback:
                progress_callback(60, "Анализ тем завершен, ищу низкоцитируемые работы...")
            
            top_topics = topics[:5]
            
            all_recommendations = []
            finder = LowCitationFinder()
            
            if progress_callback:
                progress_val = 60
                progress_callback(progress_val, f"Параллельный поиск по {len(top_topics)} темам...")
            
            keywords = []
            if analysis_result.get('keywords'):
                keywords = [kw.lower() for kw, _ in analysis_result['keywords'][:10]]
            
            with concurrent.futures.ThreadPoolExecutor(max_workers=min(len(top_topics), Config.OPENALEX_MAX_WORKERS)) as executor:
                future_to_topic = {}
                
                for topic_name, topic_count in top_topics:
                    # Находим ID темы
                    topic_id = None
                    for work in analysis_result['works_data']:
                        if work['primary_topic'] == topic_name:
                            topic_id = work['topic_id']
                            break
                    
                    if topic_id:
                        future = executor.submit(
                            finder.find_low_citation_works,
                            topic_id,
                            keywords,
                            10,  # max_citations
                            Config.OPENALEX_MAX_WORKS_PER_TOPIC
                        )
                        future_to_topic[future] = (topic_name, topic_id, topic_count)
                
                # Обрабатываем результаты
                completed = 0
                for future in concurrent.futures.as_completed(future_to_topic):
                    topic_name, topic_id, topic_count = future_to_topic[future]
                    
                    try:
                        works = future.result(timeout=Config.OPENALEX_REQUEST_TIMEOUT * 2)
                        
                        if works:
                            # Добавляем информацию о теме
                            for work in works:
                                work['topic'] = topic_name
                                work['topic_id'] = topic_id
                                work['topic_doi_count'] = topic_count
                            
                            all_recommendations.extend(works)
                            
                            completed += 1
                            if progress_callback:
                                progress_val = 60 + int((completed / len(top_topics)) * 35)
                                progress_callback(progress_val, 
                                                f"Обработано {completed}/{len(top_topics)} тем. Найдено: {len(works)} работ")
                    
                    except Exception as e:
                        logger.error(f"Error processing topic {topic_name}: {e}")
                        completed += 1
            
            if not all_recommendations:
                if progress_callback:
                    progress_callback(100, "Не найдено низкоцитируемых работ")
                return None
            
            if progress_callback:
                progress_callback(95, "Формирую итоговый список...")
            
            # Сортируем все рекомендации
            all_recommendations.sort(key=lambda x: (x['relevance_score'], -x['cited_by_count']), reverse=True)
            
            # Ограничиваем максимум рекомендаций
            all_recommendations = all_recommendations[:50]  # Вместо 100
            
            # Создаем DataFrame
            df = pd.DataFrame(all_recommendations)
            
            # Оптимизированное создание дополнительных колонок
            def create_doi_url(doi):
                if doi and isinstance(doi, str) and doi.strip():
                    clean_doi = re.sub(r'^(https?://doi\.org/|doi:|DOI:?\s*)', '', doi.strip(), flags=re.IGNORECASE)
                    return f"https://doi.org/{clean_doi}"
                return ""
            
            df['doi_url'] = df['doi'].apply(create_doi_url)
            
            def format_authors(authors_list):
                if isinstance(authors_list, list):
                    return ', '.join([str(a) for a in authors_list[:3] if a])
                return "Unknown"
            
            df['authors_formatted'] = df['authors'].apply(format_authors)
            
            def format_date(pub_date):
                if not pub_date:
                    return "Unknown"
                try:
                    if isinstance(pub_date, str):
                        return pub_date[:10]
                    return str(pub_date)
                except:
                    return "Unknown"
            
            df['publication_date_formatted'] = df['publication_date'].apply(format_date)
            
            def format_keywords(keywords_list):
                if isinstance(keywords_list, list):
                    return ', '.join([str(kw) for kw in keywords_list[:3] if kw])  # Только 3 ключевых слова
                return ""
            
            df['keywords_formatted'] = df['matched_keywords'].apply(format_keywords)
            
            if progress_callback:
                progress_callback(100, f"Готово! Найдено {len(df)} рекомендаций")
            
            return df
            
        except Exception as e:
            logger.error(f"Error in generate_recommendations: {e}")
            if progress_callback:
                progress_callback(100, f"Ошибка: {str(e)[:50]}")
            return None
    
    @staticmethod
    def create_recommendations_txt(recommendations_df):
        """Создает TXT файл с рекомендациями"""
        try:
            if recommendations_df is None or recommendations_df.empty:
                return None
            
            output_txt_buffer = io.StringIO()
            output_txt_buffer.write("ARTICLE RECOMMENDATIONS. Citation Style Construction / © IHTE, https://ihte.ru/ © CTA, https://chimicatechnoacta.ru / developed by daM©\n")
            output_txt_buffer.write("=" * 80 + "\n\n")
            output_txt_buffer.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            output_txt_buffer.write(f"Based on topic analysis and keyword matching\n")
            output_txt_buffer.write(f"Showing {len(recommendations_df)} low-citation and fresh articles\n\n")
            
            current_topic = None
            
            for idx, row in recommendations_df.iterrows():
                if row['topic'] != current_topic:
                    current_topic = row['topic']
                    output_txt_buffer.write(f"\n{'='*60}\n")
                    output_txt_buffer.write(f"TOPIC: {current_topic}\n")
                    output_txt_buffer.write(f"DOI in original list: {row.get('topic_doi_count', 'N/A')}\n")
                    output_txt_buffer.write(f"{'='*60}\n\n")
                
                citation_status = "🔴 0 citations" if row['cited_by_count'] == 0 else f"🟡 {row['cited_by_count']} citations"
                
                output_txt_buffer.write(f"{idx+1:2d}. {citation_status} [Relevance: {row['relevance_score']}/10]\n")
                output_txt_buffer.write(f"    Title: {row['title']}\n")
                
                if pd.notna(row.get('authors_formatted')) and row['authors_formatted'] != "Unknown":
                    output_txt_buffer.write(f"    Authors: {row['authors_formatted']}\n")
                
                if pd.notna(row.get('journal')):
                    journal_info = f"Journal: {row['journal']}"
                    if pd.notna(row.get('publication_year')):
                        journal_info += f", Year: {row['publication_year']}"
                    output_txt_buffer.write(f"    {journal_info}\n")
                
                if pd.notna(row.get('publication_date_formatted')) and row['publication_date_formatted'] != "Unknown":
                    output_txt_buffer.write(f"    Publication Date: {row['publication_date_formatted']}\n")
                
                if pd.notna(row.get('doi')):
                    output_txt_buffer.write(f"    DOI: {row['doi']}\n")
                    if pd.notna(row.get('doi_url')):
                        output_txt_buffer.write(f"    URL: {row['doi_url']}\n")
                
                if pd.notna(row.get('keywords_formatted')) and row['keywords_formatted']:
                    output_txt_buffer.write(f"    Matched Keywords: {row['keywords_formatted']}\n")
                
                output_txt_buffer.write("\n")
            
            output_txt_buffer.seek(0)
            return io.BytesIO(output_txt_buffer.getvalue().encode('utf-8'))
            
        except Exception as e:
            logger.error(f"Error creating recommendations TXT: {e}")
            return None
    
    @staticmethod
    def create_recommendations_csv(recommendations_df):
        """Создает CSV файл с рекомендациями"""
        try:
            if recommendations_df is None or recommendations_df.empty:
                return None
            
            # Подготавливаем DataFrame для экспорта
            export_df = recommendations_df.copy()
            
            # Выбираем нужные колонки
            columns_to_export = ['topic', 'title', 'authors_formatted', 'journal', 
                               'publication_year', 'cited_by_count', 'relevance_score',
                               'doi', 'doi_url', 'keywords_formatted', 'publication_date_formatted']
            
            # Фильтруем существующие колонки
            existing_columns = [col for col in columns_to_export if col in export_df.columns]
            
            output_csv_buffer = io.StringIO()
            export_df[existing_columns].to_csv(output_csv_buffer, index=False, encoding='utf-8')
            output_csv_buffer.seek(0)
            return io.BytesIO(output_csv_buffer.getvalue().encode('utf-8'))
            
        except Exception as e:
            logger.error(f"Error creating recommendations CSV: {e}")
            return None

class TopicSelectorUI:
    """UI компонент для выбора тем"""

    @staticmethod
    def render_topic_selection(topics_df, recommendations_df, container=None):
        """Отображает выбор тем и рекомендации"""
        if topics_df.empty or recommendations_df.empty:
            return None
        
        # Группируем рекомендации по темам
        topic_groups = {}
        for _, row in recommendations_df.iterrows():
            topic = row['topic']
            if topic not in topic_groups:
                topic_groups[topic] = []
            topic_groups[topic].append(row)
        
        # Создаем вкладки для каждой темы
        tab_labels = []
        
        # Отображаем вкладки
        if container:
            with container:
                # Берем топ-5 тем для вкладок
                topics_for_tabs = list(topic_groups.keys())[:5]
                tab_labels = [topic[:30] + "..." if len(topic) > 30 else topic 
                             for topic in topics_for_tabs]
                
                tabs = st.tabs(tab_labels)
                
                for i, tab in enumerate(tabs):
                    with tab:
                        topic_name = topics_for_tabs[i]
                        works = topic_groups[topic_name]
                        works.sort(key=lambda x: (x['relevance_score'], -x['cited_by_count']), reverse=True)
                        
                        # Заголовок вкладки
                        st.markdown(f"### 📚 Тема: {topic_name}")
                        st.markdown(f"*Найдено {len(works)} низкоцитируемых работ*")
                        
                        # Отображаем работы с использованием Streamlit компонентов
                        for j, work in enumerate(works[:20], 1):  # Топ-20 для каждой темы
                            with st.container():
                                col1, col2 = st.columns([3, 1])
                                
                                with col1:
                                    st.markdown(f"**#{j}. {work['title'][:120]}...**")
                                    
                                    authors = work.get('authors_formatted', '') or ', '.join(work.get('authors', []))
                                    if authors:
                                        st.markdown(f"👤 **Authors:** {authors}")
                                    
                                    if work.get('journal'):
                                        year_text = f" ({work['publication_year']})" if work.get('publication_year') else ""
                                        st.markdown(f"📰 **Journal:** {work['journal']}{year_text}")
                                    
                                    if work.get('matched_keywords'):
                                        keywords = ', '.join(work['matched_keywords'][:3])
                                        st.markdown(f"🔑 **Keywords:** {keywords}")
                                
                                with col2:
                                    citation_color = "🔴" if work['cited_by_count'] == 0 else "🟡"
                                    citation_text = "0 citations" if work['cited_by_count'] == 0 else f"{work['cited_by_count']} citations"
                                    st.markdown(f"{citation_color} **{citation_text}**")
                                    st.markdown(f"**Relevance:** {work['relevance_score']}/10")
                                
                                if work['doi']:
                                    doi_url = f"https://doi.org/{work['doi']}"
                                    st.markdown(f"[🔗 Open Article: {work['doi'][:30]}...]({doi_url})")
                                
                                st.divider()
        
        return tab_labels

    @staticmethod
    def _create_work_card(work, index):
        """Создает HTML карточку для работы"""
        citation_color = "🔴" if work['cited_by_count'] == 0 else "🟡"
        citation_text = f"0 цит" if work['cited_by_count'] == 0 else f"{work['cited_by_count']} цит"
        
        authors = work.get('authors_formatted', '') or ', '.join(work.get('authors', []))
        
        # Используем CSS классы из темы вместо инлайн-стилей
        html = f"""
        <div class="recommendation-item">
            <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 8px;">
                <div class="recommendation-title">#{index}. {work['title'][:120]}...</div>
                <div class="recommendation-score">
                    <span style="background-color: {'#FF6B6B' if work['cited_by_count'] == 0 else '#FFD166'}; 
                               color: white; padding: 3px 8px; border-radius: 12px; font-size: 0.9em;">
                        {citation_color} {citation_text}
                    </span>
                </div>
            </div>
            
            <div class="recommendation-meta">
        """
        
        if authors:
            html += f"<div><strong>👤 Authors:</strong> {authors}</div>"
        
        if work.get('journal'):
            year_text = f" ({work['publication_year']})" if work.get('publication_year') else ""
            html += f"<div><strong>📰 Journal:</strong> {work['journal']}{year_text}</div>"
        
        if work.get('matched_keywords'):
            keywords = ', '.join(work['matched_keywords'][:3])
            html += f"<div><strong>🔑 Keywords:</strong> {keywords}</div>"
        
        if work['doi']:
            doi_url = f"https://doi.org/{work['doi']}"
            html += f"""
            <div style="margin-top: 10px;">
                <a href="{doi_url}" target="_blank" 
                   style="text-decoration: none; color: var(--primary); font-weight: 500;">
                   🔗 Open Article DOI: {work['doi'][:30]}...
                </a>
            </div>
            """
        
        html += "</div></div>"
        return html
    
    @staticmethod
    def render_statistics(analysis_result, container=None):
        """Отображает статистику анализа"""
        if not analysis_result:
            return
        
        stats_html = """
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                   padding: 20px; border-radius: 10px; color: white; margin: 20px 0;">
            <h3 style="color: white; margin-top: 0;">📊 Analysis Statistics</h3>
            <div style="display: flex; justify-content: space-between; flex-wrap: wrap;">
        """
        
        stats = analysis_result.get('stats', {})
        
        stat_items = [
            ("Total DOI", analysis_result.get('total_works', 0), "📋"),
            ("Topics Found", len(analysis_result.get('topics', [])), "🎯"),
            ("Low-Citation Works", analysis_result.get('low_citation_count', 0), "📉"),
            ("Top Keywords", len(analysis_result.get('keywords', [])), "🔑")
        ]
        
        for label, value, icon in stat_items:
            stats_html += f"""
            <div style="text-align: center; padding: 10px; flex: 1;">
                <div style="font-size: 2em;">{icon}</div>
                <div style="font-size: 1.8em; font-weight: bold;">{value}</div>
                <div style="font-size: 0.9em; opacity: 0.9;">{label}</div>
            </div>
            """
        
        stats_html += "</div></div>"
        
        if container:
            with container:
                st.markdown(stats_html, unsafe_allow_html=True)
        else:
            st.markdown(stats_html, unsafe_allow_html=True)

# Document Generator with Recommendations
class DocumentGenerator:
    """Class for generating DOCX documents"""
    
    @staticmethod
    def add_hyperlink(paragraph, text, url):
        """Add hyperlink to paragraph"""
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        rPr.append(underline)
        
        new_run.append(rPr)
        new_text = OxmlElement('w:t')
        new_text.text = text
        new_run.append(new_text)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        
        return hyperlink
    
    @staticmethod
    def apply_yellow_background(run):
        """Apply yellow background to run"""
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'FFFF00')
        run._element.get_or_add_rPr().append(shd)
    
    @staticmethod
    def apply_blue_background(run):
        """Apply blue background to run"""
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'E6F3FF')
        run._element.get_or_add_rPr().append(shd)
    
    @staticmethod
    def apply_red_color(run):
        """Apply red color to run"""
        color = OxmlElement('w:color')
        color.set(qn('w:val'), 'FF0000')
        run._element.get_or_add_rPr().append(color)

    @staticmethod
    def apply_pink_background(run):
        """Apply pink background to run (for missing metadata warnings)"""
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'FFC0CB')  # Pink color
        run._element.get_or_add_rPr().append(shd)
    
    @staticmethod
    def generate_document(formatted_refs: List[Tuple[Any, bool, Any]], 
                         statistics: Dict[str, Any],
                         style_config: Dict[str, Any],
                         duplicates_info: Dict[int, int] = None,
                         missing_metadata_info: Dict[int, str] = None,
                         recommendations_df = None) -> io.BytesIO:
        """Generate DOCX document with references, statistics, and recommendations"""
        output_doc = Document()
        
        # Добавляем заголовок с логотипом
        try:
            if os.path.exists("logo.png"):
                # Создаем таблицу для выравнивания логотипа и текста
                title_table = output_doc.add_table(rows=1, cols=2)
                title_table.autofit = False
                
                # Ячейка для логотипа
                logo_cell = title_table.cell(0, 0)
                logo_cell.width = Pt(120)  # Ширина ячейки
                logo_paragraph = logo_cell.paragraphs[0]
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_run = logo_paragraph.add_run()
                logo_run.add_picture("logo.png", width=Pt(100))  # Размер логотипа
                
                # Ячейка для текста
                text_cell = title_table.cell(0, 1)
                text_paragraph = text_cell.paragraphs[0]
                title_run = text_paragraph.add_run('Citation Style Constructor')
                title_run.font.size = Pt(20)
                title_run.font.bold = True
                
                # Добавляем пустую строку
                output_doc.add_paragraph()
        except Exception as e:
            logger.warning(f"Не удалось добавить логотип в заголовок: {e}")
            # Если не удалось добавить логотип, просто добавляем обычный заголовок
            title_paragraph = output_doc.add_paragraph()
            title_run = title_paragraph.add_run('Citation Style Constructor')
            title_run.font.size = Pt(20)
            title_run.font.bold = True
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем разделительную линию
        sep_paragraph = output_doc.add_paragraph()
        sep_run = sep_paragraph.add_run("_" * 80)
        sep_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                         
        """Generate DOCX document with references, statistics, and recommendations"""
        output_doc = Document()
        output_doc.add_paragraph('Citation Style Construction / © IHTE, https://ihte.ru/ © CTA, https://chimicatechnoacta.ru / developed by daM©')
        output_doc.add_paragraph('See short stats after the References section')
        output_doc.add_heading('References', level=1)
        
        DocumentGenerator._add_formatted_references(output_doc, formatted_refs, style_config, duplicates_info, missing_metadata_info)
        DocumentGenerator._add_statistics_section(output_doc, statistics)
        
        if recommendations_df is not None and not recommendations_df.empty:
            DocumentGenerator._add_recommendations_section(output_doc, recommendations_df)
        
        output_doc_buffer = io.BytesIO()
        output_doc.save(output_doc_buffer)
        output_doc_buffer.seek(0)
        return output_doc_buffer

    @staticmethod
    def _add_formatted_references(doc: Document, 
                                formatted_refs: List[Tuple[Any, bool, Any]], 
                                style_config: Dict[str, Any],
                                duplicates_info: Dict[int, int] = None,
                                missing_metadata_info: Dict[int, str] = None):
        """Add formatted references to document"""
        for i, (elements, is_error, metadata) in enumerate(formatted_refs):
            numbering = style_config['numbering_style']
            
            if numbering == "No numbering":
                prefix = ""
            elif numbering == "1":
                prefix = f"{i + 1} "
            elif numbering == "1.":
                prefix = f"{i + 1}. "
            elif numbering == "1)":
                prefix = f"{i + 1}) "
            elif numbering == "(1)":
                prefix = f"({i + 1}) "
            elif numbering == "[1]":
                prefix = f"[{i + 1}] "
            else:
                prefix = f"{i + 1}. "
            
            para = doc.add_paragraph(prefix)
            
            if is_error:
                run = para.add_run(str(elements))
                DocumentGenerator.apply_yellow_background(run)
            elif duplicates_info and i in duplicates_info:
                original_index = duplicates_info[i] + 1
                duplicate_note = get_text('duplicate_reference').format(original_index)
                
                if isinstance(elements, str):
                    run = para.add_run(elements)
                    DocumentGenerator.apply_blue_background(run)
                    para.add_run(f" - {duplicate_note}").italic = True
                else:
                    for j, (value, italic, bold, separator, is_doi_hyperlink, doi_value) in enumerate(elements):
                        if is_doi_hyperlink and doi_value:
                            DocumentGenerator.add_hyperlink(para, value, f"https://doi.org/{doi_value}")
                        else:
                            run = para.add_run(value)
                            if italic:
                                run.font.italic = True
                            if bold:
                                run.font.bold = True
                            DocumentGenerator.apply_blue_background(run)
                        
                        if separator and j < len(elements) - 1:
                            para.add_run(separator)
                    
                    para.add_run(f" - {duplicate_note}").italic = True
            elif missing_metadata_info and i in missing_metadata_info:
                # Missing metadata warning
                warning_message = missing_metadata_info[i]
                
                if isinstance(elements, str):
                    run = para.add_run(elements)
                    DocumentGenerator.apply_pink_background(run)
                    para.add_run(f" - {warning_message}").italic = True
                else:
                    for j, (value, italic, bold, separator, is_doi_hyperlink, doi_value) in enumerate(elements):
                        if is_doi_hyperlink and doi_value:
                            DocumentGenerator.add_hyperlink(para, value, f"https://doi.org/{doi_value}")
                        else:
                            run = para.add_run(value)
                            if italic:
                                run.font.italic = True
                            if bold:
                                run.font.bold = True
                            DocumentGenerator.apply_pink_background(run)
                        
                        if separator and j < len(elements) - 1:
                            para.add_run(separator)
                    
                    para.add_run(f" - {warning_message}").italic = True
            else:
                if metadata is None:
                    run = para.add_run(str(elements))
                    run.font.italic = True
                else:
                    for j, (value, italic, bold, separator, is_doi_hyperlink, doi_value) in enumerate(elements):
                        if is_doi_hyperlink and doi_value:
                            DocumentGenerator.add_hyperlink(para, value, f"https://doi.org/{doi_value}")
                        else:
                            run = para.add_run(value)
                            if italic:
                                run.font.italic = True
                            if bold:
                                run.font.bold = True
                        
                        if separator and j < len(elements) - 1:
                            para.add_run(separator)
                    
                    if style_config['final_punctuation'] and not is_error:
                        para.add_run(".")
    
    @staticmethod
    def _add_statistics_section(doc: Document, statistics: Dict[str, Any]):
        """Add statistics section to document"""
        doc.add_heading('Stats', level=1)
        
        doc.add_heading('Journal Frequency', level=2)
        journal_table = doc.add_table(rows=1, cols=3)
        journal_table.style = 'Table Grid'
        
        hdr_cells = journal_table.rows[0].cells
        hdr_cells[0].text = 'Journal Name'
        hdr_cells[1].text = 'Count'
        hdr_cells[2].text = 'Percentage (%)'
        
        for journal_stat in statistics['journal_stats']:
            row_cells = journal_table.add_row().cells
            row_cells[0].text = journal_stat['journal']
            row_cells[1].text = str(journal_stat['count'])
            row_cells[2].text = str(journal_stat['percentage'])
        
        doc.add_paragraph()
        
        doc.add_heading('Year Distribution', level=2)
        
        if statistics['needs_more_recent_references']:
            warning_para = doc.add_paragraph()
            warning_run = warning_para.add_run("To improve the relevance and significance of the research, consider including more recent references published within the last 3-4 years")
            DocumentGenerator.apply_red_color(warning_run)
            doc.add_paragraph()
        
        year_table = doc.add_table(rows=1, cols=3)
        year_table.style = 'Table Grid'
        
        hdr_cells = year_table.rows[0].cells
        hdr_cells[0].text = 'Year'
        hdr_cells[1].text = 'Count'
        hdr_cells[2].text = 'Percentage (%)'
        
        for year_stat in statistics['year_stats']:
            row_cells = year_table.add_row().cells
            row_cells[0].text = str(year_stat['year'])
            row_cells[1].text = str(year_stat['count'])
            row_cells[2].text = str(year_stat['percentage'])
        
        doc.add_paragraph()
        
        doc.add_heading('Author Distribution', level=2)
        
        if statistics['has_frequent_author']:
            warning_para = doc.add_paragraph()
            warning_run = warning_para.add_run("The author(s) are referenced frequently. Either reduce the number of references to the author(s), or expand the reference list to include more sources")
            DocumentGenerator.apply_red_color(warning_run)
            doc.add_paragraph()
        
        author_table = doc.add_table(rows=1, cols=3)
        author_table.style = 'Table Grid'
        
        hdr_cells = author_table.rows[0].cells
        hdr_cells[0].text = 'Author'
        hdr_cells[1].text = 'Count'
        hdr_cells[2].text = 'Percentage (%)'
        
        for author_stat in statistics['author_stats']:
            row_cells = author_table.add_row().cells
            row_cells[0].text = author_stat['author']
            row_cells[1].text = str(author_stat['count'])
            row_cells[2].text = str(author_stat['percentage'])

    @staticmethod
    def _add_recommendations_section(doc: Document, recommendations_df):
        """Add recommendations section to document with citation counts"""
        if recommendations_df is None or recommendations_df.empty:
            return
        
        doc.add_page_break()
        doc.add_heading('Article Recommendations', level=1)
        
        current_year = datetime.now().year
        min_year = current_year - Config.RECOMMENDATION_YEARS_BACK
        
        intro_para = doc.add_paragraph()
        intro_para.add_run(f"Based on analysis of your reference list, here are {len(recommendations_df)} similar articles from the last {Config.RECOMMENDATION_YEARS_BACK} years (from {min_year} to {current_year}):").bold = True
        
        doc.add_paragraph()
        
        for idx, row in recommendations_df.iterrows():
            doc.add_heading(f"Recommendation {idx+1}: Score {row['score']:.3f}", level=2)
            
            title_para = doc.add_paragraph()
            title_para.add_run("Title: ").bold = True
            title_para.add_run(row['title'])
            
            authors_para = doc.add_paragraph()
            authors_para.add_run("Authors: ").bold = True
            authors_para.add_run(row['authors'])
            
            # Добавляем информацию о цитированиях
            info_para = doc.add_paragraph()
            info_para.add_run("Journal: ").bold = True
            info_para.add_run(f"{row['journal']}, ")
            info_para.add_run("Year: ").bold = True
            info_para.add_run(f"{row['year']}, ")
            info_para.add_run("Citations: ").bold = True
            info_para.add_run(f"{row.get('citation_count', 'N/A')}, ")
            info_para.add_run("Source: ").bold = True
            info_para.add_run(row.get('source', 'unknown'))
            
            # DOI с гиперссылкой
            if row['doi']:
                doi_para = doc.add_paragraph()
                doi_para.add_run("DOI: ").bold = True
                DocumentGenerator.add_hyperlink(doi_para, row['doi'], f"https://doi.org/{row['doi']}")
            
            if row.get('abstract'):
                abstract_para = doc.add_paragraph()
                abstract_para.add_run("Abstract: ").bold = True
                abstract_para.add_run(row['abstract'])
            
            # Общие термины
            if 'common_terms' in row and row['common_terms']:
                terms_para = doc.add_paragraph()
                terms_para.add_run("Common terms: ").bold = True
                terms_para.add_run(row['common_terms'])
            
            doc.add_paragraph()

# DOI Processor
class DOIProcessor:
    """Processor for working with DOI"""
    
    def __init__(self):
        self.cache = doi_cache
        self.works = works
    
    def find_doi_enhanced(self, reference: str) -> Optional[str]:
        """Enhanced DOI search using multiple strategies"""
        if self._is_section_header(reference):
            return None
        
        explicit_doi = self._find_explicit_doi(reference)
        if explicit_doi:
            logger.info(f"Found explicit DOI: {explicit_doi}")
            return explicit_doi
        
        bibliographic_doi = self._find_bibliographic_doi(reference)
        if bibliographic_doi:
            logger.info(f"Found bibliographic DOI: {bibliographic_doi}")
            return bibliographic_doi
        
        openalex_doi = self._find_openalex_doi(reference)
        if openalex_doi:
            logger.info(f"Found OpenAlex DOI: {openalex_doi}")
            return openalex_doi
        
        logger.warning(f"No DOI found for reference: {reference[:100]}...")
        return None
    
    def _is_section_header(self, text: str) -> bool:
        """Check if text is a section header"""
        text_upper = text.upper().strip()
        section_patterns = [
            r'^NOTES?\s+AND\s+REFERENCES?$',
            r'^REFERENCES?$',
            r'^BIBLIOGRAPHY$',
            r'^LITERATURE$',
            r'^WORKS?\s+CITED$',
            r'^SOURCES?$',
            r'^CHAPTER\s+\d+$',
            r'^SECTION\s+\d+$',
            r'^PART\s+\d+$'
        ]
        
        for pattern in section_patterns:
            if re.search(pattern, text_upper):
                return True
        return False
    
    def _find_explicit_doi(self, reference: str) -> Optional[str]:
        """Find explicit DOI in text"""
        doi_patterns = [
            r'https?://doi\.org/(10\.\d{4,9}/[-._;()/:A-Za-z0-9]+)',
            r'doi:\s*(10\.\d{4,9}/[-._;()/:A-Za-z0-9]+)',
            r'DOI:\s*(10\.\d{4,9}/[-._;()/:A-Za-z0-9]+)',
            r'\b(10\.\d{4,9}/[-._;()/:A-Za-z0-9]+)\b'
        ]
        
        for pattern in doi_patterns:
            match = re.search(pattern, reference, re.IGNORECASE)
            if match:
                doi = match.group(1).rstrip('.,;:')
                return doi
        
        clean_ref = reference.strip()
        if re.match(r'^(doi:|DOI:)?\s*10\.\d{4,9}/[-._;()/:A-Za-z0-9]+\s*$', clean_ref, re.IGNORECASE):
            doi_match = re.search(r'(10\.\d{4,9}/[-._;()/:A-Za-z0-9]+)', clean_ref)
            if doi_match:
                return doi_match.group(1).rstrip('.,;:')
        
        return None
    
    def _find_bibliographic_doi(self, reference: str) -> Optional[str]:
        """Find DOI by bibliographic data"""
        clean_ref = re.sub(r'\s*(https?://doi\.org/|doi:|DOI:)\s*[^\s,;]+', '', reference, flags=re.IGNORECASE)
        clean_ref = clean_ref.strip()
        
        if len(clean_ref) < 30:
            return None
        
        try:
            query = self.works.query(bibliographic=clean_ref).sort('relevance').order('desc')
            for result in query:
                if 'DOI' in result:
                    return result['DOI']
        except Exception as e:
            logger.error(f"Bibliographic search error for '{clean_ref}': {e}")
        
        return None
    
    def _find_openalex_doi(self, reference: str) -> Optional[str]:
        """Find DOI through OpenAlex API"""
        return None

    def extract_metadata_with_cache(self, doi: str) -> Optional[Dict]:
        """Extract metadata using cache"""
        cached_metadata = self.cache.get(doi)
        if cached_metadata:
            logger.info(f"Cache hit for DOI: {doi}")
            return cached_metadata
        
        logger.info(f"Cache miss for DOI: {doi}, fetching from API")
        metadata = self._extract_metadata_from_api(doi)
        
        if metadata:
            self.cache.set(doi, metadata)
        
        return metadata

    def _extract_metadata_from_api(self, doi: str) -> Optional[Dict]:
        """Extract metadata from Crossref API"""
        try:
            result = self.works.doi(doi)
            if not result:
                return None
            
            authors = result.get('author', [])
            author_list = []
            for author in authors:
                given_name = author.get('given', '')
                family_name = self._normalize_name(author.get('family', ''))
                author_list.append({
                    'given': given_name,
                    'family': family_name
                })
            
            title = ''
            if 'title' in result and result['title']:
                title = self._clean_text(result['title'][0])
                title = re.sub(r'</?sub>|</?i>|</?SUB>|</?I>', '', title, flags=re.IGNORECASE)
            
            journal = ''
            if 'container-title' in result and result['container-title']:
                journal = self._clean_text(result['container-title'][0])
            
            year = None
            
            if 'published-print' in result and 'date-parts' in result['published-print']:
                date_parts = result['published-print']['date-parts']
                if date_parts and date_parts[0] and len(date_parts[0]) > 0:
                    year = date_parts[0][0]
                    logger.info(f"Using published-print year {year} for DOI {doi}")
            
            if year is None and 'published' in result and 'date-parts' in result['published']:
                date_parts = result['published']['date-parts']
                if date_parts and date_parts[0] and len(date_parts[0]) > 0:
                    year = date_parts[0][0]
                    logger.info(f"Using published year {year} for DOI {doi}")
            
            if year is None:
                date_fields = ['issued', 'published-online', 'created']
                for field in date_fields:
                    if field in result and 'date-parts' in result[field]:
                        date_parts = result[field]['date-parts']
                        if date_parts and date_parts[0] and len(date_parts[0]) > 0:
                            year = date_parts[0][0]
                            logger.info(f"Using {field} year {year} for DOI {doi}")
                            break
            
            volume = result.get('volume', '')
            issue = result.get('issue', '')
            pages = result.get('page', '')
            article_number = result.get('article-number', '')
            
            abstract = ''
            if 'abstract' in result:
                abstract = self._clean_text(result['abstract'])
            
            metadata = {
                'authors': author_list,
                'title': title,
                'journal': journal,
                'year': year,
                'volume': volume,
                'issue': issue,
                'pages': pages,
                'article_number': article_number,
                'doi': doi,
                'original_doi': doi,
                'abstract': abstract
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting metadata for DOI {doi}: {e}")
            return None
    
    def _normalize_name(self, name: str) -> str:
        """Normalize author name"""
        if not name:
            return ''
        
        if '-' in name or "'" in name or '’' in name:
            parts = re.split(r'([-\'’])', name)
            normalized_parts = []
            
            for i, part in enumerate(parts):
                if part in ['-', "'", '’']:
                    normalized_parts.append(part)
                else:
                    if part:
                        normalized_parts.append(part[0].upper() + part[1:].lower() if len(part) > 1 else part.upper())
            
            return ''.join(normalized_parts)
        else:
            if len(name) > 1:
                return name[0].upper() + name[1:].lower()
            else:
                return name.upper()

    def _clean_text(self, text: str) -> str:
        """Clean text from HTML tags and entities"""
        if not text:
            return ""
        
        # Сохраняем важные теги <sub>, </sub>, <i>, </i>, <scp>, </scp>
        # Временно заменяем их на специальные маркеры
        text = text.replace('<sub>', '«SUB»').replace('</sub>', '«/SUB»')
        text = text.replace('<i>', '«I»').replace('</i>', '«/I»')
        text = text.replace('<scp>', '«SCP»').replace('</scp>', '«/SCP»')
        text = text.replace('<SUP>', '«SUP»').replace('</SUP>', '«/SUP»')
        text = text.replace('<sup>', '«SUP»').replace('</sup>', '«/SUP»')
        
        # Удаляем все остальные HTML-теги
        text = re.sub(r'<[^>]+>', '', text)
        
        # Восстанавливаем сохраненные теги
        text = text.replace('«SUB»', '<sub>').replace('«/SUB»', '</sub>')
        text = text.replace('«I»', '<i>').replace('«/I»', '</i>')
        text = text.replace('«SCP»', '').replace('«/SCP»', '')  # Убираем scp, так как это просто для капителизации
        text = text.replace('«SUP»', '<sup>').replace('«/SUP»', '</sup>')
        
        # Обрабатываем специальные символы и сущности
        text = html.unescape(text)
        
        # Заменяем оставшиеся сущности
        text = re.sub(r'&[^;]+;', '', text)
        
        # Нормализуем пробелы и тире
        text = text.replace('\n', ' ').replace('\r', '')
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        
        # Если весь текст в верхнем регистре или большинство букв заглавные, нормализуем регистр
        if text and (text.isupper() or sum(1 for c in text if c.isupper()) > len(text) * 0.7):
            text = text.title()
            corrections = {
                ' A ': ' a ',
                ' An ': ' an ',
                ' The ': ' the ',
                ' And ': ' and ',
                ' But ': ' but ',
                ' Or ': ' or ',
                ' For ': ' for ',
                ' Nor ': ' nor ',
                ' On ': ' on ',
                ' At ': ' at ',
                ' To ': ' to ',
                ' By ': ' by ',
                ' In ': ' in ',
                ' Of ': ' of ',
                ' With ': ' with ',
                ' As ': ' as ',
                ' Is ': ' is ',
                ' Via ': ' via ',
                ' Vs ': ' vs ',
                ' Vs. ': ' vs. ',
                ' Etc ': ' etc ',
                ' Etc. ': ' etc. ',
            }
            for wrong, correct in corrections.items():
                text = text.replace(wrong, correct)
                text = text.replace(wrong.upper(), correct)
        
        return text

# Reference Processor
class ReferenceProcessor:
    """Main processor for reference processing"""
    
    def __init__(self):
        self.doi_processor = DOIProcessor()
        self.progress_manager = ProgressManager()
        self.validator = StyleValidator()
    
    def process_references(self, references: List[str], style_config: Dict, 
                         progress_container, status_container) -> Tuple[List, io.BytesIO, io.BytesIO, int, int, Dict]:
        """Process list of references with progress display"""
        is_valid, validation_messages = self.validator.validate_references_count(references)
        for msg in validation_messages:
            if "error" in msg.lower():
                st.error(msg)
            else:
                st.warning(msg)
        
        if not is_valid:
            return [], io.BytesIO(), io.BytesIO(), 0, 0, {}
        
        doi_list = []
        formatted_refs = []
        formatted_texts = []
        doi_found_count = 0
        doi_not_found_count = 0
        
        valid_dois = []
        reference_doi_map = {}
        
        for i, ref in enumerate(references):
            if self.doi_processor._is_section_header(ref):
                doi_list.append(f"{ref} [SECTION HEADER - SKIPPED]")
                formatted_refs.append((ref, False, None))
                formatted_texts.append(ref)
                continue
                
            doi = self.doi_processor.find_doi_enhanced(ref)
            if doi:
                valid_dois.append(doi)
                reference_doi_map[i] = doi
                doi_list.append(doi)
            else:
                error_msg = self._create_error_message(ref, st.session_state.current_language)
                doi_list.append(error_msg)
                formatted_refs.append((error_msg, True, None))
                formatted_texts.append(error_msg)
                doi_not_found_count += 1
        
        if valid_dois:
            self._process_doi_batch(
                valid_dois, reference_doi_map, references, 
                formatted_refs, formatted_texts, doi_list, style_config,
                progress_container, status_container
            )
        
        doi_found_count = len([ref for ref in formatted_refs if not ref[1] and ref[2]])

        duplicates_info = self._find_duplicates(formatted_refs)
        missing_metadata_info = self._find_missing_metadata(formatted_refs)
        
        formatted_txt_buffer = self._create_formatted_txt_file(formatted_texts)
        original_txt_buffer = self._create_txt_file(doi_list)
        
        return formatted_refs, formatted_txt_buffer, original_txt_buffer, doi_found_count, doi_not_found_count, duplicates_info, missing_metadata_info
    
    def _process_doi_batch(self, valid_dois, reference_doi_map, references, 
                          formatted_refs, formatted_texts, doi_list, style_config,
                          progress_container, status_container):
        """Batch process DOI"""
        status_container.info(get_text('batch_processing'))
        
        total_to_process = len(valid_dois)
        self.progress_manager.start_processing(total_to_process)
        
        progress_bar = progress_container.progress(0)
        status_display = status_container.empty()
        
        metadata_results = self._extract_metadata_batch(valid_dois, progress_bar, status_display)
        
        doi_to_metadata = dict(zip(valid_dois, metadata_results))
        
        processed_count = 0
        found_count = 0
        error_count = 0
        
        for i, ref in enumerate(references):
            if i in reference_doi_map:
                doi = reference_doi_map[i]
                metadata = doi_to_metadata.get(doi)
                
                if metadata:
                    formatted_ref, is_error = self._format_reference(metadata, style_config)
                    formatted_text = self._format_reference_for_text(metadata, style_config)
                    
                    if doi in doi_list:
                        index = doi_list.index(doi)
                        doi_list[index] = formatted_text
                    
                    formatted_refs.append((formatted_ref, is_error, metadata))
                    formatted_texts.append(formatted_text)
                    found_count += 1
                else:
                    error_msg = self._create_error_message(ref, st.session_state.current_language)
                    if doi in doi_list:
                        index = doi_list.index(doi)
                        doi_list[index] = error_msg
                    formatted_refs.append((error_msg, True, None))
                    formatted_texts.append(error_msg)
                    error_count += 1
                
                processed_count += 1
                
                self.progress_manager.update_progress(processed_count, found_count, error_count, 'formatting')
                progress_ratio = processed_count / total_to_process if total_to_process > 0 else 0
                progress_bar.progress(progress_ratio)
                
                status_text = f"Processed: {processed_count}/{total_to_process} | Found: {found_count} | Errors: {error_count}"
                status_display.text(status_text)
        
        self.progress_manager.update_progress(total_to_process, found_count, error_count, 'complete')
        progress_bar.progress(1.0)

    def _extract_metadata_batch(self, doi_list, progress_bar, status_display) -> List:
        """Batch extract metadata with retry"""
        results = [None] * len(doi_list)
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=Config.CROSSREF_WORKERS) as executor:
            future_to_index = {
                executor.submit(self.doi_processor.extract_metadata_with_cache, doi): i 
                for i, doi in enumerate(doi_list)
            }
            
            completed = 0
            total = len(doi_list)
            for future in concurrent.futures.as_completed(future_to_index):
                index = future_to_index[future]
                try:
                    result = future.result(timeout=Config.REQUEST_TIMEOUT)
                    results[index] = result
                except Exception as e:
                    logger.error(f"Error processing DOI at index {index}: {e}")
                    results[index] = None
                
                completed += 1
                progress_ratio = completed / total if total > 0 else 0
                progress_bar.progress(progress_ratio)
                status_display.text(f"Fetching metadata: {completed}/{total}")
        
        failed_indices = [i for i, result in enumerate(results) if result is None]
        
        if failed_indices:
            logger.info(f"Retrying {len(failed_indices)} failed requests...")
            status_display.text(f"Retrying {len(failed_indices)} failed requests...")
            
            self._retry_failed_requests(failed_indices, doi_list, results, progress_bar, status_display)
        
        return results
    
    def _retry_failed_requests(self, failed_indices, doi_list, results, progress_bar, status_display):
        """Retry failed requests"""
        completed = len(doi_list) - len(failed_indices)
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=Config.CROSSREF_RETRY_WORKERS) as executor:
            retry_futures = {}
            for index in failed_indices:
                doi = doi_list[index]
                future = executor.submit(self.doi_processor.extract_metadata_with_cache, doi)
                retry_futures[future] = index
            
            for future in concurrent.futures.as_completed(retry_futures):
                index = retry_futures[future]
                try:
                    result = future.result(timeout=Config.REQUEST_TIMEOUT)
                    results[index] = result
                except Exception as e:
                    logger.error(f"Error in retry processing DOI at index {index}: {e}")
                    results[index] = None
                
                completed += 1
                self._update_progress_display(progress_bar, status_display, completed, len(doi_list), len(failed_indices))
    
    def _update_progress_display(self, progress_bar, status_display, completed, total, errors):
        """Update progress display"""
        progress_info = self.progress_manager.get_progress_info()
        progress_ratio = completed / total if total > 0 else 0
        progress_color = self.progress_manager.get_progress_color(progress_ratio)
        
        progress_bar.progress(progress_ratio)
        
        progress_bar.markdown(f"""
            <style>
                .stProgress > div > div > div > div {{
                    background-color: {progress_color};
                }}
            </style>
        """, unsafe_allow_html=True)
        
        status_text = f"Processed: {completed}/{total} | Errors: {errors}"
        if progress_info['time_remaining']:
            mins_remaining = int(progress_info['time_remaining'] / 60)
            status_text += f" | ETA: {mins_remaining} min"
        
        status_display.text(status_text)
    
    def _format_reference(self, metadata: Dict, style_config: Dict) -> Tuple[Any, bool]:
        """Format reference for DOCX"""
        formatter = CitationFormatterFactory.create_formatter(style_config)
        return formatter.format_reference(metadata, False)
    
    def _format_reference_for_text(self, metadata: Dict, style_config: Dict) -> str:
        """Format reference for TXT file"""
        formatter = CitationFormatterFactory.create_formatter(style_config)
        elements, _ = formatter.format_reference(metadata, False)
        
        if isinstance(elements, str):
            return elements
        
        ref_str = ""
        for i, (value, italic, bold, separator, is_doi_hyperlink, doi_value) in enumerate(elements):
            if italic and bold:
                formatted_value = f"***{value}***"
            elif italic:
                formatted_value = f"*{value}*"
            elif bold:
                formatted_value = f"**{value}**"
            else:
                formatted_value = value
            
            ref_str += formatted_value
            
            if separator and i < len(elements) - 1:
                ref_str += separator
        
        if style_config.get('final_punctuation') and not ref_str.endswith('.'):
            ref_str += "."
        
        return ref_str

    def _find_duplicates(self, formatted_refs: List) -> Dict[int, int]:
        """Find duplicate references"""
        seen_hashes = {}
        duplicates_info = {}
        
        for i, (elements, is_error, metadata) in enumerate(formatted_refs):
            if is_error or not metadata:
                continue
                
            ref_hash = self._generate_reference_hash(metadata)
            if not ref_hash:
                continue
                
            if ref_hash in seen_hashes:
                duplicates_info[i] = seen_hashes[ref_hash]
            else:
                seen_hashes[ref_hash] = i
        
        return duplicates_info
    
    def _find_missing_metadata(self, formatted_refs: List) -> Dict[int, str]:
        """Find references with missing important metadata (volume, pages/article number)"""
        missing_metadata_info = {}
        
        for i, (elements, is_error, metadata) in enumerate(formatted_refs):
            if is_error or not metadata:
                continue
            
            # Check for missing important metadata
            missing_fields = []
            
            # Check volume
            if not metadata.get('volume'):
                missing_fields.append('volume')
            
            # Check for pages OR article number
            if not metadata.get('pages') and not metadata.get('article_number'):
                missing_fields.append('pages/article number')
            
            # If we have missing fields, create warning message
            if missing_fields:
                missing_list = ' and '.join(missing_fields)
                warning_msg = f"⚠️ {missing_list.capitalize()} information is missing. "
                
                # Add specific guidance based on missing fields
                if 'volume' in missing_fields:
                    warning_msg += "This may indicate a non-journal source (book, chapter, or conference paper) or a journal article with incomplete issue assignment. "
                
                warning_msg += "Please verify the source type and consider manual correction."
                
                missing_metadata_info[i] = warning_msg
        
        return missing_metadata_info
    
    def _generate_reference_hash(self, metadata: Dict) -> Optional[str]:
        """Generate hash for identifying duplicates"""
        if not metadata:
            return None
        
        hash_string = ""
        
        if metadata.get('authors'):
            authors_hash = "|".join(sorted([author.get('family', '').lower() for author in metadata['authors']]))
            hash_string += authors_hash + "||"
        
        title = metadata.get('title', '')[:50].lower()
        hash_string += title + "||"
        
        hash_string += (metadata.get('journal', '') + "||").lower()
        hash_string += str(metadata.get('year', '')) + "||"
        hash_string += metadata.get('volume', '') + "||"
        hash_string += metadata.get('pages', '') + "||"
        hash_string += self._normalize_doi(metadata.get('doi', ''))
        
        return hashlib.md5(hash_string.encode('utf-8')).hexdigest()
    
    def _normalize_doi(self, doi: str) -> str:
        """Normalize DOI"""
        if not doi:
            return ""
        return re.sub(r'^(https?://doi\.org/|doi:|DOI:)', '', doi, flags=re.IGNORECASE).lower().strip()
    
    def _create_error_message(self, ref: str, language: str) -> str:
        """Create error message"""
        if language == 'ru':
            return f"{ref}\nПроверьте источник и добавьте DOI вручную."
        else:
            return f"{ref}\nPlease check this source and insert the DOI manually."
    
    def _create_missing_metadata_message(self, metadata: Dict, language: str) -> str:
        """Create missing metadata warning message"""
        missing_fields = []
        
        if not metadata.get('volume'):
            missing_fields.append('volume')
        
        if not metadata.get('pages') and not metadata.get('article_number'):
            missing_fields.append('pages/article number')
        
        missing_list = ' and '.join(missing_fields)
        
        if language == 'ru':
            return f"⚠️ Отсутствует информация о {missing_list}. Это может указывать на источник не из журнала (книга, глава или тезисы конференции). Проверьте тип источника."
        else:
            return f"⚠️ {missing_list.capitalize()} information is missing. This may indicate a non-journal source (book, chapter, or conference paper) or a journal article with incomplete issue assignment. Please verify the source type."
    
    def _create_formatted_txt_file(self, formatted_texts: List[str]) -> io.BytesIO:
        """Create TXT file with formatted references"""
        output_txt_buffer = io.StringIO()
        for text in formatted_texts:
            output_txt_buffer.write(f"{text}\n\n")
        output_txt_buffer.seek(0)
        return io.BytesIO(output_txt_buffer.getvalue().encode('utf-8'))
    
    def _create_txt_file(self, doi_list: List[str]) -> io.BytesIO:
        """Create TXT file with DOI list"""
        output_txt_buffer = io.StringIO()
        for doi in doi_list:
            output_txt_buffer.write(f"{doi}\n")
        output_txt_buffer.seek(0)
        return io.BytesIO(output_txt_buffer.getvalue().encode('utf-8'))

# Theme Manager
class ThemeManager:
    """Theme manager"""
    
    @staticmethod
    def get_theme_css(theme_name: str) -> str:
        """Get CSS styles for theme"""
        theme = Config.THEMES.get(theme_name, Config.THEMES['light'])
        
        button_styles = {
            'rounded': 'border-radius: 8px;',
            'classic': 'border-radius: 4px; border: 1px solid ' + theme['border'] + ';',
            'rounded-full': 'border-radius: 20px;',
            'square': 'border-radius: 0;'
        }
        
        button_style = button_styles.get(theme['buttonStyle'], 'border-radius: 8px;')
        
        return f"""
            <style>
            :root {{
                --primary: {theme['primary']};
                --secondary: {theme['secondary']};
                --accent: {theme['accent']};
                --background: {theme['background']};
                --secondaryBackground: {theme['secondaryBackground']};
                --text: {theme['text']};
                --font: {theme['font']};
                --border: {theme['border']};
                --cardBackground: {theme['cardBackground']};
                --shadow: {theme['shadow']};
            }}
            
            .main {{
                background-color: {theme['background']};
                color: {theme['text']};
                font-family: {theme['font']};
            }}
            
            .stage-container {{
                background-color: {theme['secondaryBackground']};
                border-radius: 10px;
                padding: 10px;
                margin-bottom: 15px;
                box-shadow: {theme['shadow']};
                border: 1px solid {theme['border']};
            }}
            
            .stage-active {{
                background-color: {theme['primary']};
                color: white;
                font-weight: bold;
                padding: 8px 15px;
                border-radius: 5px;
                margin: 0 5px;
                display: inline-block;
            }}
            
            .stage-inactive {{
                background-color: {theme['secondaryBackground']};
                color: {theme['text']};
                padding: 8px 15px;
                border-radius: 5px;
                margin: 0 5px;
                display: inline-block;
                opacity: 0.7;
                border: 1px solid {theme['border']};
            }}
            
            .stage-connector {{
                color: {theme['border']};
                margin: 0 5px;
                font-weight: bold;
            }}
            
            .stButton > button {{
                {button_style}
                background-color: {theme['primary']};
                color: white;
                font-family: {theme['font']};
                font-weight: 500;
                padding: 6px 12px;
                transition: all 0.2s ease;
                border: none;
            }}
            
            .stButton > button:hover {{
                background-color: {theme['secondary']};
                transform: translateY(-2px);
                box-shadow: 0 4px 8px rgba(0,0,0,0.2);
            }}
            
            .card {{
                background-color: {theme['cardBackground']};
                border-radius: 10px;
                padding: 20px;
                margin-bottom: 15px;
                box-shadow: {theme['shadow']};
                border: 1px solid {theme['border']};
            }}
            
            .card-title {{
                color: {theme['primary']};
                font-weight: bold;
                margin-bottom: 15px;
                font-size: 1.2rem;
            }}
            
            .style-item {{
                background-color: {theme['cardBackground']};
                border-radius: 6px;
                padding: 8px;
                margin-bottom: 10px;
                box-shadow: {theme['shadow']};
                border-left: 3px solid {theme['primary']};
                transition: all 0.2s ease;
            }}
            
            .style-item:hover {{
                transform: translateY(-2px);
                box-shadow: 0 6px 12px rgba(0,0,0,0.15);
            }}
            
            .style-preview {{
                background-color: {theme['secondaryBackground']};
                padding: 6px;
                border-radius: 4px;
                font-family: monospace;
                font-size: 0.85em;
                line-height: 1.3;
                margin-top: 6px;
                border: 1px solid {theme['border']};
            }}
            
            .stSelectbox, .stTextInput, .stNumberInput, .stCheckbox, .stRadio, .stFileUploader, .stTextArea {{
                margin-bottom: 10px;
            }}
            
            .stSelectbox > div > div {{
                background-color: {theme['secondaryBackground']};
                border: 1px solid {theme['border']};
                color: {theme['text']};
            }}
            
            .stTextArea textarea {{
                background-color: {theme['secondaryBackground']};
                color: {theme['text']};
                border: 1px solid {theme['border']};
                font-family: {theme['font']};
            }}
            
            h1, h2, h3 {{
                color: {theme['text']} !important;
                font-family: {theme['font']} !important;
            }}
            
            h1 {{
                color: {theme['primary']} !important;
                border-bottom: 2px solid {theme['primary']};
                padding-bottom: 10px;
                margin-bottom: 20px;
            }}
            
            .stat-card {{
                background-color: {theme['cardBackground']};
                border-left: 4px solid {theme['primary']};
                padding: 15px;
                margin-bottom: 10px;
                border-radius: 5px;
            }}
            
            .stat-value {{
                font-size: 1.5rem;
                font-weight: bold;
                color: {theme['primary']};
            }}
            
            .stat-label {{
                color: {theme['text']};
                opacity: 0.8;
                font-size: 0.9rem;
            }}
            
            .global-stats-container {{
                background-image: linear-gradient(135deg, {theme['primary']} 0%, {theme['secondary']} 100%);
                padding: 20px;
                border-radius: 15px;
                text-align: center;
                color: white;
                margin: 25px 0;
                box-shadow: 0 6px 15px rgba(0,0,0,0.2);
                position: relative;
                overflow: hidden;
            }}
            
            .global-stats-container::before {{
                content: '';
                position: absolute;
                top: -50%;
                left: -50%;
                width: 200%;
                height: 200%;
                background: radial-gradient(circle, rgba(255,255,255,0.1) 1px, transparent 1px);
                background-size: 20px 20px;
                opacity: 0.3;
                animation: moveBackground 20s linear infinite;
            }}
            
            @keyframes moveBackground {{
                0% {{ transform: translate(0, 0); }}
                100% {{ transform: translate(20px, 20px); }}
            }}
            
            .global-stats-title {{
                font-size: 0.95rem;
                opacity: 0.9;
                margin-bottom: 5px;
                position: relative;
                z-index: 1;
            }}
            
            .global-stats-value {{
                font-size: 2.2rem;
                font-weight: bold;
                margin: 10px 0;
                position: relative;
                z-index: 1;
                text-shadow: 0 2px 4px rgba(0,0,0,0.3);
            }}
            
            .global-stats-label {{
                font-size: 0.85rem;
                opacity: 0.9;
                position: relative;
                z-index: 1;
            }}
            
            .global-stats-date {{
                font-size: 0.75rem;
                opacity: 0.7;
                margin-top: 8px;
                position: relative;
                z-index: 1;
            }}
            
            .global-stats-info {{
                background-color: {theme['cardBackground']};
                padding: 12px;
                border-radius: 8px;
                margin-top: 10px;
                border-left: 3px solid {theme['accent']};
                font-size: 0.85rem;
                position: relative;
                z-index: 1;
            }}
            
            @keyframes fadeIn {{
                from {{ opacity: 0; transform: translateY(20px); }}
                to {{ opacity: 1; transform: translateY(0); }}
            }}
            
            .page-content {{
                animation: fadeIn 0.5s ease-out;
            }}
            
            .create-style-preview {{
                background-color: {theme['secondaryBackground']};
                padding: 15px;
                border-radius: 5px;
                border-left: 3px solid {theme['accent']};
                margin: 15px 0;
                font-family: {theme['font']};
                line-height: 1.5;
            }}
            
            .create-style-preview .formatted-text {{
                font-family: {theme['font']};
                line-height: 1.5;
            }}
            
            .create-style-preview .formatted-text-italic {{
                font-style: italic;
            }}
            
            .create-style-preview .formatted-text-bold {{
                font-weight: bold;
            }}
            
            .create-style-preview .formatted-text-italic-bold {{
                font-style: italic;
                font-weight: bold;
            }}
            
            .element-config-row {{
                background-color: {theme['secondaryBackground']};
                padding: 10px;
                margin: 5px 0;
                border-radius: 5px;
                border: 1px solid {theme['border']};
            }}
            
            .setting-item {{
                margin-bottom: 15px;
            }}
            
            .setting-label {{
                font-weight: 500;
                color: {theme['text']};
                margin-bottom: 5px;
                display: block;
            }}
            
            .result-item {{
                background-color: {theme['secondaryBackground']};
                padding: 10px;
                margin: 5px 0;
                border-radius: 5px;
                border-left: 3px solid {theme['primary']};
            }}
            
            .download-button {{
                background-color: {theme['accent']} !important;
            }}
            
            .download-button:hover {{
                background-color: {theme['secondary']} !important;
            }}
            
            .scrollable-results {{
                max-height: 400px;
                overflow-y: auto;
                padding: 10px;
                border: 1px solid {theme['border']};
                border-radius: 5px;
                background-color: {theme['secondaryBackground']};
                margin-top: 10px;
            }}
            
            .scrollable-results::-webkit-scrollbar {{
                width: 8px;
            }}
            
            .scrollable-results::-webkit-scrollbar-track {{
                background: {theme['background']};
                border-radius: 4px;
            }}
            
            .scrollable-results::-webkit-scrollbar-thumb {{
                background: {theme['primary']};
                border-radius: 4px;
            }}
            
            .scrollable-results::-webkit-scrollbar-thumb:hover {{
                background: {theme['secondary']};
            }}
            
            .formatted-text {{
                font-family: {theme['font']};
                line-height: 1.5;
                margin-bottom: 8px;
            }}
            
            .formatted-text-italic {{
                font-style: italic;
            }}
            
            .formatted-text-bold {{
                font-weight: bold;
            }}
            
            .formatted-text-italic-bold {{
                font-style: italic;
                font-weight: bold;
            }}
            
            .error-reference {{
                background-color: rgba(255, 204, 0, 0.1);
                border-left: 3px solid #ffcc00;
            }}
            
            .duplicate-reference {{
                background-color: rgba(78, 205, 196, 0.1);
                border-left: 3px solid #4ECDC4;
            }}

            .missing-metadata-reference {{
                background-color: rgba(255, 192, 203, 0.1);
                border-left: 3px solid #F7ADBA;
            }}
            
            .select-scroll-container {{
                max-height: 600px;
                overflow-y: auto;
                padding-right: 10px;
            }}
            
            .select-scroll-container::-webkit-scrollbar {{
                width: 8px;
            }}
            
            .select-scroll-container::-webkit-scrollbar-track {{
                background: {theme['background']};
                border-radius: 4px;
            }}
            
            .select-scroll-container::-webkit-scrollbar-thumb {{
                background: {theme['primary']};
                border-radius: 4px;
            }}
            
            .select-scroll-container::-webkit-scrollbar-thumb:hover {{
                background: {theme['secondary']};
            }}
            
            .compact-select-row {{
                display: flex;
                align-items: center;
                margin: 2px 0;
                padding: 2px 0;
                border-bottom: 1px solid {theme['border']};
                min-height: 30px;
            }}
            
            .compact-select-button {{
                width: 120px !important;
                min-width: 120px !important;
                max-width: 120px !important;
                margin: 0 !important;
                padding: 2px 5px !important;
                font-size: 0.8rem !important;
                height: 26px !important;
            }}
            
            .compact-select-preview {{
                font-family: 'Courier New', monospace;
                font-size: 0.75rem;
                line-height: 1.1;
                margin-left: 8px;
                flex-grow: 1;
                overflow-wrap: break-word;
                padding: 2px 4px;
                background-color: {theme['secondaryBackground']};
                border-radius: 3px;
                border-left: 2px solid {theme['primary']};
            }}
            
            .compact-select-name {{
                font-weight: bold;
                color: {theme['primary']};
                margin-right: 5px;
            }}
            
            .recommendation-item {{
                background-color: {theme['secondaryBackground']};
                padding: 15px;
                margin-bottom: 10px;
                border-radius: 5px;
                border-left: 4px solid {theme['accent']};
            }}
            
            .recommendation-score {{
                font-weight: bold;
                color: {theme['primary']};
                font-size: 1.1rem;
            }}
            
            .recommendation-title {{
                font-weight: bold;
                margin: 5px 0;
            }}
            
            .recommendation-meta {{
                color: {theme['text']};
                opacity: 0.8;
                font-size: 0.9rem;
                margin-bottom: 5px;
            }}
            
            .recommendation-abstract {{
                background-color: {theme['background']};
                padding: 10px;
                border-radius: 3px;
                margin-top: 5px;
                font-size: 0.9rem;
                line-height: 1.4;
            }}

            .recommendation-item {{
                background-color: var(--cardBackground);
                padding: 15px;
                margin-bottom: 10px;
                border-radius: 8px;
                border-left: 4px solid var(--accent);
                box-shadow: var(--shadow);
            }}
            
            .recommendation-score {{
                font-weight: bold;
                color: var(--primary);
                font-size: 1.1rem;
            }}
            
            .recommendation-title {{
                font-weight: bold;
                margin: 5px 0;
                color: var(--text);
            }}
            
            .recommendation-meta {{
                color: var(--text);
                opacity: 0.9;
                font-size: 0.95rem;
                margin-top: 8px;
            }}
            
            .recommendation-abstract {{
                background-color: var(--background);
                padding: 12px;
                border-radius: 6px;
                margin-top: 8px;
                font-size: 0.9rem;
                line-height: 1.5;
            }}
            
            .recommendation-progress {{
                margin: 10px 0;
            }}
            </style>
        """

    @staticmethod
    def apply_theme(theme_name: str):
        """Apply theme to application"""
        st.markdown(ThemeManager.get_theme_css(theme_name), unsafe_allow_html=True)

# Stage Manager
class StageManager:
    """Application stage manager"""
    
    @staticmethod
    def render_stage_indicator(current_stage: str):
        """Render stage indicator"""
        stages = list(Config.STAGES.keys())
        current_index = stages.index(current_stage)
        
        stage_html = '<div class="stage-container">'
        stage_html += '<div style="display: flex; align-items: center; justify-content: center; flex-wrap: wrap;">'
        
        for i, stage_key in enumerate(stages):
            stage_name = get_text(f'stage_{stage_key}')
            
            if i == current_index:
                stage_html += f'<div class="stage-active">{stage_name}</div>'
            else:
                stage_html += f'<div class="stage-inactive">{stage_name}</div>'
            
            if i < len(stages) - 1:
                stage_html += '<div class="stage-connector">→</div>'
        
        stage_html += '</div></div>'
        
        st.markdown(stage_html, unsafe_allow_html=True)
    
    @staticmethod
    def navigate_to(stage: str):
        """Navigate to specified stage"""
        if stage not in st.session_state.stage_history:
            st.session_state.stage_history.append(stage)
        st.session_state.current_stage = stage
        st.rerun()
    
    @staticmethod
    def go_back():
        """Go back to previous stage"""
        if len(st.session_state.stage_history) > 1:
            st.session_state.stage_history.pop()
            previous_stage = st.session_state.stage_history[-1]
            st.session_state.current_stage = previous_stage
            st.rerun()
    
    @staticmethod
    def clear_all():
        """Clear all data and return to start"""
        init_session_state()
        st.session_state.current_stage = 'start'
        st.session_state.stage_history = ['start']
        st.rerun()

# Start Page
class StartPage:
    """Start page"""
    
    @staticmethod
    def render():
        """Render start page"""
        st.markdown(f"<h1>{get_text('start_title')}</h1>", unsafe_allow_html=True)
        st.markdown(f"<p style='margin-bottom: 30px;'>{get_text('start_description')}</p>", unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button(get_text('start_ready_presets'), use_container_width=True, key="ready_presets_btn"):
                StageManager.navigate_to('select')
        
        with col2:
            if st.button(get_text('start_create_style'), use_container_width=True, key="create_style_btn"):
                StageManager.navigate_to('create')
        
        with col3:
            if st.button(get_text('start_load_style'), use_container_width=True, key="load_style_btn"):
                st.session_state.show_style_loader = True
        
        if st.session_state.get('show_style_loader', False):
            st.markdown("---")
            st.subheader(get_text('load_style'))
            
            uploaded_file = st.file_uploader(
                get_text('import_file'),
                type=['json'],
                help="Load style file in JSON format",
                key="style_loader"
            )
            
            if uploaded_file is not None:
                try:
                    content = uploaded_file.read().decode('utf-8')
                    imported_style = json.loads(content)
                    
                    if 'style_config' in imported_style:
                        style_config = imported_style['style_config']
                    else:
                        style_config = imported_style
                    
                    apply_imported_style(style_config)
                    st.session_state.style_config = style_config
                    
                    st.success(get_text('style_loaded'))
                    
                    if st.button(get_text('proceed_to_io'), type="primary"):
                        StageManager.navigate_to('io')
                        
                except Exception as e:
                    st.error(f"{get_text('import_error')}: {str(e)}")
            
            if st.button(get_text('back_button')):
                st.session_state.show_style_loader = False
                st.rerun()

# Select Page
class SelectPage:
    """Select page"""
    
    @staticmethod
    def _get_style_previews() -> List[Tuple[int, str, str]]:
        """Get previews for all styles"""
        previews = [
            (1, "ГОСТ", "Dreyer D.R., Park S., Bielawski C.W., Ruoff R.S. The chemistry of graphene oxide // Chemical Society Reviews. – 2010. – Vol. 39, № 1. – Р. 228-240. – https://doi.org/10.1039/B917103G"),
            (2, "ACS (MDPI)", "Dreyer, D.R.; Park, S.; Bielawski, C.W.; Ruoff, R.S. The chemistry of graphene oxide. *Chem. Soc. Rev.* **2010**, *39*, 228–240. https://doi.org/10.1039/B917103G"),
            (3, "RSC", "D.R. Dreyer, S. Park, C.W. Bielawski and R.S. Ruoff, *Chem. Soc. Rev.*, 2010, **39**, 228"),
            (4, "CTA", "Dreyer DR, Park S, Bielawski CW, Ruoff RS. The chemistry of graphene oxide. Chem Soc Rev. 2010;39(1):228–40. doi:10.1039/B917103G"),
            (5, "Style 5", "D.R. Dreyer, S. Park, C.W. Bielawski, R.S. Ruoff, The chemistry of graphene oxide, Chem. Soc. Rev. 39 (2010) 228–240. https://doi.org/10.1039/B917103G"),
            (6, "Style 6", "Dreyer, D.R., Park, S., Bielawski, C.W., Ruoff, R.S. (2010). The chemistry of graphene oxide. Chem. Soc. Rev. *39*, 228–240. https://doi.org/10.1039/B917103G."),
            (7, "Style 7", "Dreyer, D.R., Park, S., Bielawski, C.W. & Ruoff, R.S. (2010). The chemistry of graphene oxide. *Chemical Society Reviews* *39*(1), 228–240. https://doi.org/10.1039/B917103G."),
            (8, "Style 8", "D. R. Dreyer, S. Park, C. W. Bielawski, R. S. Ruoff, *Chem. Soc. Rev.* **2010**, *39*, 228"),
            (9, "RCR", "D.R.Dreyer, S.Park, C.W.Bielawski, R.S.Ruoff. *Chem. Soc. Rev.*, **39**, 228 (2010); https://doi.org/10.1039/B917103G"),
            (10, "Style 10", "Dreyer DR, Park S, Bielawski CW, Ruoff RS (2010) The chemistry of graphene oxide. Chem Soc Rev 39(1):228–240. https://doi.org/10.1039/B917103G")
        ]
        return previews
    
    @staticmethod
    def _apply_style_1():
        """Apply style 1 (GOST)"""
        st.session_state.num = "No numbering"
        st.session_state.auth = "Smith AA"
        st.session_state.sep = ", "
        st.session_state.etal = 0
        st.session_state.use_and_checkbox = False
        st.session_state.use_ampersand_checkbox = False
        st.session_state.doi = "https://doi.org/10.10/xxx"
        st.session_state.doilink = True
        st.session_state.page = "122-128"
        st.session_state.punct = ""
        st.session_state.journal_style = "{Full Journal Name}"
        
        for i in range(8):
            st.session_state[f"el{i}"] = ""
            st.session_state[f"it{i}"] = False
            st.session_state[f"bd{i}"] = False
            st.session_state[f"pr{i}"] = False
            st.session_state[f"sp{i}"] = ". "
        
        st.session_state.gost_style = True
        st.session_state.acs_style = False
        st.session_state.rsc_style = False
        st.session_state.cta_style = False
        st.session_state.style5 = False
        st.session_state.style6 = False
        st.session_state.style7 = False
        st.session_state.style8 = False
        st.session_state.style9 = False
        st.session_state.style10 = False
        st.session_state.custom_style_created = True
        
        st.session_state.style_config = {
            'author_format': st.session_state.auth,
            'author_separator': st.session_state.sep,
            'et_al_limit': st.session_state.etal if st.session_state.etal > 0 else None,
            'use_and_bool': st.session_state.use_and_checkbox,
            'use_ampersand_bool': st.session_state.use_ampersand_checkbox,
            'doi_format': st.session_state.doi,
            'doi_hyperlink': st.session_state.doilink,
            'page_format': st.session_state.page,
            'final_punctuation': st.session_state.punct,
            'numbering_style': st.session_state.num,
            'journal_style': st.session_state.journal_style,
            'elements': [],
            'gost_style': True,
            'acs_style': False,
            'rsc_style': False,
            'cta_style': False,
            'style5': False,
            'style6': False,
            'style7': False,
            'style8': False,
            'style9': False,
            'style10': False
        }
    
    @staticmethod
    def _apply_style_2():
        """Apply style 2 (ACS MDPI)"""
        st.session_state.num = "No numbering"
        st.session_state.auth = "Smith, A.A."
        st.session_state.sep = "; "
        st.session_state.etal = 0
        st.session_state.use_and_checkbox = False
        st.session_state.use_ampersand_checkbox = False
        st.session_state.doi = "10.10/xxx"
        st.session_state.doilink = True
        st.session_state.page = "122–128"
        st.session_state.punct = "."
        st.session_state.journal_style = "{J. Abbr.}"
        
        for i in range(8):
            st.session_state[f"el{i}"] = ""
            st.session_state[f"it{i}"] = False
            st.session_state[f"bd{i}"] = False
            st.session_state[f"pr{i}"] = False
            st.session_state[f"sp{i}"] = ". "
        
        st.session_state.gost_style = False
        st.session_state.acs_style = True
        st.session_state.rsc_style = False
        st.session_state.cta_style = False
        st.session_state.style5 = False
        st.session_state.style6 = False
        st.session_state.style7 = False
        st.session_state.style8 = False
        st.session_state.style9 = False
        st.session_state.style10 = False
        st.session_state.custom_style_created = True
        
        st.session_state.style_config = {
            'author_format': st.session_state.auth,
            'author_separator': st.session_state.sep,
            'et_al_limit': st.session_state.etal if st.session_state.etal > 0 else None,
            'use_and_bool': st.session_state.use_and_checkbox,
            'use_ampersand_bool': st.session_state.use_ampersand_checkbox,
            'doi_format': st.session_state.doi,
            'doi_hyperlink': st.session_state.doilink,
            'page_format': st.session_state.page,
            'final_punctuation': st.session_state.punct,
            'numbering_style': st.session_state.num,
            'journal_style': st.session_state.journal_style,
            'elements': [],
            'gost_style': False,
            'acs_style': True,
            'rsc_style': False,
            'cta_style': False,
            'style5': False,
            'style6': False,
            'style7': False,
            'style8': False,
            'style9': False,
            'style10': False
        }
    
    @staticmethod
    def _apply_style_3():
        """Apply style 3 (RSC)"""
        st.session_state.num = "No numbering"
        st.session_state.auth = "A.A. Smith"
        st.session_state.sep = ", "
        st.session_state.etal = 0
        st.session_state.use_and_checkbox = True
        st.session_state.use_ampersand_checkbox = False
        st.session_state.doi = "10.10/xxx"
        st.session_state.doilink = True
        st.session_state.page = "122"
        st.session_state.punct = "."
        st.session_state.journal_style = "{J. Abbr.}"
        
        for i in range(8):
            st.session_state[f"el{i}"] = ""
            st.session_state[f"it{i}"] = False
            st.session_state[f"bd{i}"] = False
            st.session_state[f"pr{i}"] = False
            st.session_state[f"sp{i}"] = ". "
        
        st.session_state.gost_style = False
        st.session_state.acs_style = False
        st.session_state.rsc_style = True
        st.session_state.cta_style = False
        st.session_state.style5 = False
        st.session_state.style6 = False
        st.session_state.style7 = False
        st.session_state.style8 = False
        st.session_state.style9 = False
        st.session_state.style10 = False
        st.session_state.custom_style_created = True
        
        st.session_state.style_config = {
            'author_format': st.session_state.auth,
            'author_separator': st.session_state.sep,
            'et_al_limit': st.session_state.etal if st.session_state.etal > 0 else None,
            'use_and_bool': st.session_state.use_and_checkbox,
            'use_ampersand_bool': st.session_state.use_ampersand_checkbox,
            'doi_format': st.session_state.doi,
            'doi_hyperlink': st.session_state.doilink,
            'page_format': st.session_state.page,
            'final_punctuation': st.session_state.punct,
            'numbering_style': st.session_state.num,
            'journal_style': st.session_state.journal_style,
            'elements': [],
            'gost_style': False,
            'acs_style': False,
            'rsc_style': True,
            'cta_style': False,
            'style5': False,
            'style6': False,
            'style7': False,
            'style8': False,
            'style9': False,
            'style10': False
        }
    
    @staticmethod
    def _apply_style_4():
        """Apply style 4 (CTA)"""
        st.session_state.num = "No numbering"
        st.session_state.auth = "Smith AA"
        st.session_state.sep = ", "
        st.session_state.etal = 0
        st.session_state.use_and_checkbox = False
        st.session_state.use_ampersand_checkbox = False
        st.session_state.doi = "doi:10.10/xxx"
        st.session_state.doilink = True
        st.session_state.page = "122–8"
        st.session_state.punct = ""
        st.session_state.journal_style = "{J Abbr}"
        
        for i in range(8):
            st.session_state[f"el{i}"] = ""
            st.session_state[f"it{i}"] = False
            st.session_state[f"bd{i}"] = False
            st.session_state[f"pr{i}"] = False
            st.session_state[f"sp{i}"] = ". "
        
        st.session_state.gost_style = False
        st.session_state.acs_style = False
        st.session_state.rsc_style = False
        st.session_state.cta_style = True
        st.session_state.style5 = False
        st.session_state.style6 = False
        st.session_state.style7 = False
        st.session_state.style8 = False
        st.session_state.style9 = False
        st.session_state.style10 = False
        st.session_state.custom_style_created = True
        
        st.session_state.style_config = {
            'author_format': st.session_state.auth,
            'author_separator': st.session_state.sep,
            'et_al_limit': st.session_state.etal if st.session_state.etal > 0 else None,
            'use_and_bool': st.session_state.use_and_checkbox,
            'use_ampersand_bool': st.session_state.use_ampersand_checkbox,
            'doi_format': st.session_state.doi,
            'doi_hyperlink': st.session_state.doilink,
            'page_format': st.session_state.page,
            'final_punctuation': st.session_state.punct,
            'numbering_style': st.session_state.num,
            'journal_style': st.session_state.journal_style,
            'elements': [],
            'gost_style': False,
            'acs_style': False,
            'rsc_style': False,
            'cta_style': True,
            'style5': False,
            'style6': False,
            'style7': False,
            'style8': False,
            'style9': False,
            'style10': False
        }
    
    @staticmethod
    def _apply_style_5():
        """Apply style 5"""
        st.session_state.num = "No numbering"
        st.session_state.auth = "A.A. Smith"
        st.session_state.sep = ", "
        st.session_state.etal = 0
        st.session_state.use_and_checkbox = False
        st.session_state.use_ampersand_checkbox = False
        st.session_state.doi = "https://doi.org/10.10/xxx"
        st.session_state.doilink = True
        st.session_state.page = "122–128"
        st.session_state.punct = "."
        st.session_state.journal_style = "{J. Abbr.}"
        
        for i in range(8):
            st.session_state[f"el{i}"] = ""
            st.session_state[f"it{i}"] = False
            st.session_state[f"bd{i}"] = False
            st.session_state[f"pr{i}"] = False
            st.session_state[f"sp{i}"] = ". "
        
        st.session_state.gost_style = False
        st.session_state.acs_style = False
        st.session_state.rsc_style = False
        st.session_state.cta_style = False
        st.session_state.style5 = True
        st.session_state.style6 = False
        st.session_state.style7 = False
        st.session_state.style8 = False
        st.session_state.style9 = False
        st.session_state.style10 = False
        st.session_state.custom_style_created = True
        
        st.session_state.style_config = {
            'author_format': st.session_state.auth,
            'author_separator': st.session_state.sep,
            'et_al_limit': st.session_state.etal if st.session_state.etal > 0 else None,
            'use_and_bool': st.session_state.use_and_checkbox,
            'use_ampersand_bool': st.session_state.use_ampersand_checkbox,
            'doi_format': st.session_state.doi,
            'doi_hyperlink': st.session_state.doilink,
            'page_format': st.session_state.page,
            'final_punctuation': st.session_state.punct,
            'numbering_style': st.session_state.num,
            'journal_style': st.session_state.journal_style,
            'elements': [],
            'gost_style': False,
            'acs_style': False,
            'rsc_style': False,
            'cta_style': False,
            'style5': True,
            'style6': False,
            'style7': False,
            'style8': False,
            'style9': False,
            'style10': False
        }
    
    @staticmethod
    def _apply_style_6():
        """Apply style 6"""
        st.session_state.num = "No numbering"
        st.session_state.auth = "Smith, A.A."
        st.session_state.sep = ", "
        st.session_state.etal = 0
        st.session_state.use_and_checkbox = False
        st.session_state.use_ampersand_checkbox = False
        st.session_state.doi = "https://doi.org/10.10/xxx"
        st.session_state.doilink = True
        st.session_state.page = "122–128"
        st.session_state.punct = "."
        st.session_state.journal_style = "{Full Journal Name}"
        
        for i in range(8):
            st.session_state[f"el{i}"] = ""
            st.session_state[f"it{i}"] = False
            st.session_state[f"bd{i}"] = False
            st.session_state[f"pr{i}"] = False
            st.session_state[f"sp{i}"] = ". "
        
        st.session_state.gost_style = False
        st.session_state.acs_style = False
        st.session_state.rsc_style = False
        st.session_state.cta_style = False
        st.session_state.style5 = False
        st.session_state.style6 = True
        st.session_state.style7 = False
        st.session_state.style8 = False
        st.session_state.style9 = False
        st.session_state.style10 = False
        st.session_state.custom_style_created = True
        
        st.session_state.style_config = {
            'author_format': st.session_state.auth,
            'author_separator': st.session_state.sep,
            'et_al_limit': st.session_state.etal if st.session_state.etal > 0 else None,
            'use_and_bool': st.session_state.use_and_checkbox,
            'use_ampersand_bool': st.session_state.use_ampersand_checkbox,
            'doi_format': st.session_state.doi,
            'doi_hyperlink': st.session_state.doilink,
            'page_format': st.session_state.page,
            'final_punctuation': st.session_state.punct,
            'numbering_style': st.session_state.num,
            'journal_style': st.session_state.journal_style,
            'elements': [],
            'gost_style': False,
            'acs_style': False,
            'rsc_style': False,
            'cta_style': False,
            'style5': False,
            'style6': True,
            'style7': False,
            'style8': False,
            'style9': False,
            'style10': False
        }
    
    @staticmethod
    def _apply_style_7():
        """Apply style 7"""
        st.session_state.num = "No numbering"
        st.session_state.auth = "Smith, A.A."
        st.session_state.sep = ", "
        st.session_state.etal = 0
        st.session_state.use_and_checkbox = False
        st.session_state.use_ampersand_checkbox = True
        st.session_state.doi = "https://doi.org/10.10/xxx"
        st.session_state.doilink = True
        st.session_state.page = "122–128"
        st.session_state.punct = "."
        st.session_state.journal_style = "{Full Journal Name}"
        
        for i in range(8):
            st.session_state[f"el{i}"] = ""
            st.session_state[f"it{i}"] = False
            st.session_state[f"bd{i}"] = False
            st.session_state[f"pr{i}"] = False
            st.session_state[f"sp{i}"] = ". "
        
        st.session_state.gost_style = False
        st.session_state.acs_style = False
        st.session_state.rsc_style = False
        st.session_state.cta_style = False
        st.session_state.style5 = False
        st.session_state.style6 = False
        st.session_state.style7 = True
        st.session_state.style8 = False
        st.session_state.style9 = False
        st.session_state.style10 = False
        st.session_state.custom_style_created = True
        
        st.session_state.style_config = {
            'author_format': st.session_state.auth,
            'author_separator': st.session_state.sep,
            'et_al_limit': st.session_state.etal if st.session_state.etal > 0 else None,
            'use_and_bool': st.session_state.use_and_checkbox,
            'use_ampersand_bool': st.session_state.use_ampersand_checkbox,
            'doi_format': st.session_state.doi,
            'doi_hyperlink': st.session_state.doilink,
            'page_format': st.session_state.page,
            'final_punctuation': st.session_state.punct,
            'numbering_style': st.session_state.num,
            'journal_style': st.session_state.journal_style,
            'elements': [],
            'gost_style': False,
            'acs_style': False,
            'rsc_style': False,
            'cta_style': False,
            'style5': False,
            'style6': False,
            'style7': True,
            'style8': False,
            'style9': False,
            'style10': False
        }
    
    @staticmethod
    def _apply_style_8():
        """Apply style 8"""
        st.session_state.num = "No numbering"
        st.session_state.auth = "A. A. Smith"
        st.session_state.sep = ", "
        st.session_state.etal = 0
        st.session_state.use_and_checkbox = False
        st.session_state.use_ampersand_checkbox = False
        st.session_state.doi = "10.10/xxx"
        st.session_state.doilink = False
        st.session_state.page = "122"
        st.session_state.punct = "."
        st.session_state.journal_style = "{J. Abbr.}"
        
        for i in range(8):
            st.session_state[f"el{i}"] = ""
            st.session_state[f"it{i}"] = False
            st.session_state[f"bd{i}"] = False
            st.session_state[f"pr{i}"] = False
            st.session_state[f"sp{i}"] = ". "
        
        st.session_state.gost_style = False
        st.session_state.acs_style = False
        st.session_state.rsc_style = False
        st.session_state.cta_style = False
        st.session_state.style5 = False
        st.session_state.style6 = False
        st.session_state.style7 = False
        st.session_state.style8 = True
        st.session_state.style9 = False
        st.session_state.style10 = False
        st.session_state.custom_style_created = True
        
        st.session_state.style_config = {
            'author_format': st.session_state.auth,
            'author_separator': st.session_state.sep,
            'et_al_limit': st.session_state.etal if st.session_state.etal > 0 else None,
            'use_and_bool': st.session_state.use_and_checkbox,
            'use_ampersand_bool': st.session_state.use_ampersand_checkbox,
            'doi_format': st.session_state.doi,
            'doi_hyperlink': st.session_state.doilink,
            'page_format': st.session_state.page,
            'final_punctuation': st.session_state.punct,
            'numbering_style': st.session_state.num,
            'journal_style': st.session_state.journal_style,
            'elements': [],
            'gost_style': False,
            'acs_style': False,
            'rsc_style': False,
            'cta_style': False,
            'style5': False,
            'style6': False,
            'style7': False,
            'style8': True,
            'style9': False,
            'style10': False
        }
    
    @staticmethod
    def _apply_style_9():
        """Apply style 9 (RCR)"""
        st.session_state.num = "No numbering"
        st.session_state.auth = "A.A.Smith"
        st.session_state.sep = ", "
        st.session_state.etal = 0
        st.session_state.use_and_checkbox = False
        st.session_state.use_ampersand_checkbox = False
        st.session_state.doi = "https://doi.org/10.10/xxx"
        st.session_state.doilink = True
        st.session_state.page = "122"
        st.session_state.punct = ""
        st.session_state.journal_style = "{J. Abbr.}"
        
        for i in range(8):
            st.session_state[f"el{i}"] = ""
            st.session_state[f"it{i}"] = False
            st.session_state[f"bd{i}"] = False
            st.session_state[f"pr{i}"] = False
            st.session_state[f"sp{i}"] = ". "
        
        st.session_state.gost_style = False
        st.session_state.acs_style = False
        st.session_state.rsc_style = False
        st.session_state.cta_style = False
        st.session_state.style5 = False
        st.session_state.style6 = False
        st.session_state.style7 = False
        st.session_state.style8 = False
        st.session_state.style9 = True
        st.session_state.style10 = False
        st.session_state.custom_style_created = True
        
        st.session_state.style_config = {
            'author_format': st.session_state.auth,
            'author_separator': st.session_state.sep,
            'et_al_limit': st.session_state.etal if st.session_state.etal > 0 else None,
            'use_and_bool': st.session_state.use_and_checkbox,
            'use_ampersand_bool': st.session_state.use_ampersand_checkbox,
            'doi_format': st.session_state.doi,
            'doi_hyperlink': st.session_state.doilink,
            'page_format': st.session_state.page,
            'final_punctuation': st.session_state.punct,
            'numbering_style': st.session_state.num,
            'journal_style': st.session_state.journal_style,
            'elements': [],
            'gost_style': False,
            'acs_style': False,
            'rsc_style': False,
            'cta_style': False,
            'style5': False,
            'style6': False,
            'style7': False,
            'style8': False,
            'style9': True,
            'style10': False
        }
    
    @staticmethod
    def _apply_style_10():
        """Apply style 10"""
        st.session_state.num = "No numbering"
        st.session_state.auth = "Smith AA"
        st.session_state.sep = " "
        st.session_state.etal = 0
        st.session_state.use_and_checkbox = False
        st.session_state.use_ampersand_checkbox = False
        st.session_state.doi = "https://doi.org/10.10/xxx"
        st.session_state.doilink = True
        st.session_state.page = "122–128"
        st.session_state.punct = ""
        st.session_state.journal_style = "{J Abbr}"
        
        for i in range(8):
            st.session_state[f"el{i}"] = ""
            st.session_state[f"it{i}"] = False
            st.session_state[f"bd{i}"] = False
            st.session_state[f"pr{i}"] = False
            st.session_state[f"sp{i}"] = ". "
        
        st.session_state.gost_style = False
        st.session_state.acs_style = False
        st.session_state.rsc_style = False
        st.session_state.cta_style = False
        st.session_state.style5 = False
        st.session_state.style6 = False
        st.session_state.style7 = False
        st.session_state.style8 = False
        st.session_state.style9 = False
        st.session_state.style10 = True
        st.session_state.custom_style_created = True
        
        st.session_state.style_config = {
            'author_format': st.session_state.auth,
            'author_separator': st.session_state.sep,
            'et_al_limit': st.session_state.etal if st.session_state.etal > 0 else None,
            'use_and_bool': st.session_state.use_and_checkbox,
            'use_ampersand_bool': st.session_state.use_ampersand_checkbox,
            'doi_format': st.session_state.doi,
            'doi_hyperlink': st.session_state.doilink,
            'page_format': st.session_state.page,
            'final_punctuation': st.session_state.punct,
            'numbering_style': st.session_state.num,
            'journal_style': st.session_state.journal_style,
            'elements': [],
            'gost_style': False,
            'acs_style': False,
            'rsc_style': False,
            'cta_style': False,
            'style5': False,
            'style6': False,
            'style7': False,
            'style8': False,
            'style9': False,
            'style10': True
        }
    
    @staticmethod
    def _apply_style_by_number(style_num: int):
        """Apply style by number"""
        style_apply_functions = {
            1: SelectPage._apply_style_1,
            2: SelectPage._apply_style_2,
            3: SelectPage._apply_style_3,
            4: SelectPage._apply_style_4,
            5: SelectPage._apply_style_5,
            6: SelectPage._apply_style_6,
            7: SelectPage._apply_style_7,
            8: SelectPage._apply_style_8,
            9: SelectPage._apply_style_9,
            10: SelectPage._apply_style_10
        }
        
        if style_num in style_apply_functions:
            style_apply_functions[style_num]()

    @staticmethod
    def _render_compact_style_row(style_num: int, style_name: str, preview_text: str):
        """Compact render style row with button and preview"""
        col_btn, col_preview = st.columns([1, 9])
        
        with col_btn:
            btn_key = f"select_style_{style_num}_{hash(preview_text)}"
            if st.button(f"Style {style_num}", 
                        key=btn_key,
                        use_container_width=True,
                        type="primary" if style_num <= 4 else "secondary"):
                SelectPage._apply_style_by_number(style_num)
                StageManager.navigate_to('io')
        
        with col_preview:
            preview_clean = preview_text.replace('\n', ' ')
            display_text = preview_clean
            formatted_html = display_text
            formatted_html = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', formatted_html)
            formatted_html = re.sub(r'\*(?!\*)(.*?)(?<!\*)\*', r'<em>\1</em>', formatted_html)
            formatted_html = re.sub(r'_(.*?)_', r'<em>\1</em>', formatted_html)
            
            html_content = f"""
            <div style="font-family: 'Courier New', monospace; font-size: 0.8rem; 
                        line-height: 1.2; padding: 3px; background-color: var(--secondaryBackground); 
                        border-radius: 3px; border-left: 2px solid var(--primary); 
                        margin: 2px 0;">
                <span style="font-weight: bold; color: var(--primary);">{style_name}:</span> {formatted_html}
            </div>
            """
            st.markdown(html_content, unsafe_allow_html=True)
    
    @staticmethod
    def render():
        """Compact render select page"""
        st.markdown(f"<h1 style='margin-bottom: 5px; font-size: 1.4rem;'>{get_text('select_title')}</h1>", unsafe_allow_html=True)
        st.markdown(f"<p style='margin-bottom: 10px; font-size: 0.85rem;'>{get_text('select_description')}</p>", unsafe_allow_html=True)
        
        style_previews = SelectPage._get_style_previews()
        
        for style_num, style_name, preview_text in style_previews:
            SelectPage._render_compact_style_row(style_num, style_name, preview_text)
        
        st.markdown("<div style='margin-top: 15px; padding-top: 10px; border-top: 1px solid var(--border);'>", unsafe_allow_html=True)
        col_back, col_custom = st.columns([1, 1])
        
        with col_back:
            if st.button(get_text('back_to_start'), use_container_width=True, key="back_from_select", 
                        help="Return to start page"):
                StageManager.navigate_to('start')
        
        with col_custom:
            if st.button("Create Custom Style", use_container_width=True, key="go_to_custom",
                        help="Go to custom style creation"):
                StageManager.navigate_to('create')
        
        st.markdown("</div>", unsafe_allow_html=True)

# Create Page
class CreatePage:
    """Create page"""
    
    @staticmethod
    def render():
        """Render create page"""
        st.markdown(f"<h1>{get_text('create_title')}</h1>", unsafe_allow_html=True)
        st.markdown(f"<p style='margin-bottom: 30px;'>{get_text('create_description')}</p>", unsafe_allow_html=True)
        
        with st.container():
            CreatePage._render_general_settings()
            CreatePage._render_element_configuration()
            CreatePage._render_style_preview()
            CreatePage._render_action_buttons()
    
    @staticmethod
    def _render_general_settings():
        """Render general settings in 3x4 format"""
        st.markdown(f"<div class='card' style='margin-bottom: 5px; padding: 10px;'><div class='card-title' style='margin-bottom: 10px;'>{get_text('general_settings')}</div>", unsafe_allow_html=True)

        st.markdown("""
        <style>
        div[data-testid="column"] {
            padding-top: 0px !important;
            padding-bottom: 0px !important;
        }
        .stSelectbox, .stTextInput, .stNumberInput, .stCheckbox, .stRadio {
            margin-bottom: 5px !important;
        }
        </style>
        """, unsafe_allow_html=True)
        
        col1, col2, col3, col4 = st.columns(4)

        with col1:
            num_value = st.session_state.num
            num_index = 0
            if num_value in Config.NUMBERING_STYLES:
                num_index = Config.NUMBERING_STYLES.index(num_value)
            
            st.selectbox(
                get_text('numbering_style'),
                Config.NUMBERING_STYLES,
                key="num",
                index=num_index
            )

        with col2:
            # Получаем обновленный список форматов
            available_formats = [
                "AA Smith", "A.A. Smith", "Smith AA", "Smith A.A", "Smith, A.A.",
                "A A Smith", "A. A. Smith", "Smith A A", "Smith A. A.", "Smith A. A", "Smith, A. A."
            ]
            
            # Обновляем значение в session_state, если оно невалидно
            if st.session_state.auth not in available_formats:
                st.session_state.auth = available_formats[0]
            
            # Создаем selectbox
            st.selectbox(
                get_text('author_format'),
                available_formats,
                key="auth"
            )
            
        with col3:
            st.selectbox(
                get_text('author_separator'),
                [", ", "; "],
                key="sep",
                index=[", ", "; "].index(st.session_state.sep)
            )
        
        with col4:
            st.number_input(
                get_text('et_al_limit'),
                min_value=0,
                step=1,
                key="etal",
                value=st.session_state.etal
            )
        
        col5, col6, col7, col8 = st.columns(4)
        
        with col5:
            st.checkbox(
                get_text('use_and'),
                key="use_and_checkbox",
                value=st.session_state.use_and_checkbox
            )
        
        with col6:
            st.checkbox(
                get_text('use_ampersand'),
                key="use_ampersand_checkbox",
                value=st.session_state.use_ampersand_checkbox
            )
        
        with col7:
            journal_style_value = st.session_state.journal_style
            journal_index = 0
            if journal_style_value in Config.JOURNAL_STYLES:
                journal_index = Config.JOURNAL_STYLES.index(journal_style_value)
            
            st.selectbox(
                get_text('journal_style'),
                Config.JOURNAL_STYLES,
                key="journal_style",
                index=journal_index,
                format_func=lambda x: {
                    "{Full Journal Name}": get_text('full_journal_name'),
                    "{J. Abbr.}": get_text('journal_abbr_with_dots'),
                    "{J Abbr}": get_text('journal_abbr_no_dots')
                }[x]
            )
        
        with col8:
            current_page = st.session_state.page
            page_index = 0
            if current_page in Config.PAGE_FORMATS:
                page_index = Config.PAGE_FORMATS.index(current_page)
            
            st.selectbox(
                get_text('page_format'),
                Config.PAGE_FORMATS,
                key="page",
                index=page_index
            )
        
        col9, col10, col11, col12 = st.columns(4)

        with col9:
            doi_value = st.session_state.doi
            doi_index = 0
            if doi_value in Config.DOI_FORMATS:
                doi_index = Config.DOI_FORMATS.index(doi_value)
            
            st.selectbox(
                get_text('doi_format'),
                Config.DOI_FORMATS,
                key="doi",
                index=doi_index
            )
                
        with col10:
            st.checkbox(
                get_text('doi_hyperlink'),
                key="doilink",
                value=st.session_state.doilink
            )
        
        with col11:
            st.selectbox(
                get_text('final_punctuation'),
                ["", "."],
                key="punct",
                index=["", "."].index(st.session_state.punct)
            )
        
        with col12:
            st.write("")
        
        st.markdown("</div>", unsafe_allow_html=True)
    
    @staticmethod
    def _render_element_configuration():
        """Render element configuration in 5 columns"""
        st.markdown(f"<div class='card' style='margin-bottom: 5px; padding: 10px;'><div class='card-title' style='margin-bottom: 10px;'>{get_text('element_config')}</div>", unsafe_allow_html=True)
        
        cols = st.columns([2, 1, 1, 1, 2])
        with cols[0]:
            st.markdown(f"<small><b>{get_text('element')}</b></small>", unsafe_allow_html=True)
        with cols[1]:
            st.markdown(f"<small><b>{get_text('italic')}</b></small>", unsafe_allow_html=True)
        with cols[2]:
            st.markdown(f"<small><b>{get_text('bold')}</b></small>", unsafe_allow_html=True)
        with cols[3]:
            st.markdown(f"<small><b>{get_text('parentheses')}</b></small>", unsafe_allow_html=True)
        with cols[4]:
            st.markdown(f"<small><b>{get_text('separator')}</b></small>", unsafe_allow_html=True)
        
        for i in range(8):
            cols = st.columns([2, 1, 1, 1, 2])
            
            with cols[0]:
                el_value = st.session_state[f"el{i}"]
                el_index = 0
                if el_value in Config.AVAILABLE_ELEMENTS:
                    el_index = Config.AVAILABLE_ELEMENTS.index(el_value)
                
                st.selectbox(
                    "",
                    Config.AVAILABLE_ELEMENTS,
                    key=f"el{i}",
                    label_visibility="collapsed",
                    index=el_index
                )
            
            with cols[1]:
                st.checkbox(
                    "",
                    key=f"it{i}",
                    label_visibility="collapsed"
                )
            
            with cols[2]:
                st.checkbox(
                    "",
                    key=f"bd{i}",
                    label_visibility="collapsed"
                )
            
            with cols[3]:
                st.checkbox(
                    "",
                    key=f"pr{i}",
                    label_visibility="collapsed"
                )
            
            with cols[4]:
                st.text_input(
                    "",
                    value=st.session_state[f"sp{i}"],
                    key=f"sp{i}",
                    label_visibility="collapsed"
                )
        
        st.markdown("</div>", unsafe_allow_html=True)
        
    @staticmethod              
    def _render_style_preview():
        """Render style preview"""
        style_config = CreatePage._get_style_config()
        
        # Всегда показываем превью, даже если элементы не настроены
        # Это даст пользователю визуальную обратную связь
        preview_metadata = CreatePage._get_preview_metadata(style_config)
        
        if preview_metadata:
            elements, _ = format_reference(preview_metadata, style_config, for_preview=False)
            
            st.markdown(f"<div class='card' style='margin-bottom: 5px; padding: 10px;'><div class='card-title' style='margin-bottom: 10px;'>{get_text('style_preview')}</div>", unsafe_allow_html=True)
            
            st.markdown(f"<small><b>{get_text('example')}</b></small>", unsafe_allow_html=True)
            
            if isinstance(elements, str):
                preview_with_numbering = CreatePage._add_numbering_to_elements(elements, style_config)
                display_html = f'<div class="formatted-text">{preview_with_numbering}</div>'
            else:
                html_parts = []
                
                numbering = style_config.get('numbering_style', 'No numbering')
                prefix = ""
                if numbering != "No numbering":
                    if numbering == "1":
                        prefix = f"<span>1 </span>"
                    elif numbering == "1.":
                        prefix = f"<span>1. </span>"
                    elif numbering == "1)":
                        prefix = f"<span>1) </span>"
                    elif numbering == "(1)":
                        prefix = f"<span>(1) </span>"
                    elif numbering == "[1]":
                        prefix = f"<span>[1] </span>"
                    else:
                        prefix = f"<span>1. </span>"
                
                html_parts.append(prefix)
                
                for j, element_data in enumerate(elements):
                    value, italic, bold, separator, is_doi_hyperlink, doi_value = element_data
                    
                    format_classes = []
                    if italic and bold:
                        format_classes.append("formatted-text-italic-bold")
                    elif italic:
                        format_classes.append("formatted-text-italic")
                    elif bold:
                        format_classes.append("formatted-text-bold")
                    
                    format_class = " ".join(format_classes) if format_classes else ""
                    
                    if format_class:
                        value_html = f'<span class="{format_class}">{value}</span>'
                    else:
                        value_html = f'<span>{value}</span>'
                    
                    html_parts.append(value_html)
                    
                    if separator and j < len(elements) - 1:
                        html_parts.append(f'<span>{separator}</span>')
                
                if style_config.get('final_punctuation'):
                    if html_parts and html_parts[-1].endswith('.'):
                        html_parts[-1] = html_parts[-1][:-1]
                    html_parts.append('<span>.</span>')
                
                full_html = "".join(html_parts)
                display_html = f'<div class="formatted-text">{full_html}</div>'
            
            st.markdown(f'<div class="style-preview">{display_html}</div>', unsafe_allow_html=True)
            
            # Добавим информационное сообщение, если элементы не настроены
            if not style_config['elements'] and not any([
                style_config.get('gost_style', False),
                style_config.get('acs_style', False),
                style_config.get('rsc_style', False),
                style_config.get('cta_style', False),
                style_config.get('style5', False),
                style_config.get('style6', False),
                style_config.get('style7', False),
                style_config.get('style8', False),
                style_config.get('style9', False),
                style_config.get('style10', False)
            ]):
                st.info("💡 Настройте элементы в разделе выше, чтобы увидеть изменения в превью")
            
            st.markdown("</div>", unsafe_allow_html=True)
        else:
            # Если нет данных для превью, покажем информационное сообщение
            st.markdown(f"<div class='card' style='margin-bottom: 5px; padding: 10px;'><div class='card-title' style='margin-bottom: 10px;'>{get_text('style_preview')}</div>", unsafe_allow_html=True)
            st.info("ℹ️ Выберите стиль или настройте элементы, чтобы увидеть превью")
            st.markdown("</div>", unsafe_allow_html=True)
    
    @staticmethod
    def _add_numbering_to_elements(elements, style_config):
        """Add numbering to elements"""
        if isinstance(elements, str):
            numbering = style_config.get('numbering_style', 'No numbering')
            if numbering == "No numbering":
                return elements
            elif numbering == "1":
                return f"1 {elements}"
            elif numbering == "1.":
                return f"1. {elements}"
            elif numbering == "1)":
                return f"1) {elements}"
            elif numbering == "(1)":
                return f"(1) {elements}"
            elif numbering == "[1]":
                return f"[1] {elements}"
            else:
                return f"1. {elements}"
        return elements
    
    @staticmethod
    def _get_style_config() -> Dict:
        """Get style configuration"""
        element_configs = []
        used_elements = set()
        
        for i in range(8):
            element = st.session_state.get(f"el{i}", "")
            if element and element not in used_elements:
                element_configs.append((
                    element,
                    {
                        'italic': st.session_state.get(f"it{i}", False),
                        'bold': st.session_state.get(f"bd{i}", False),
                        'parentheses': st.session_state.get(f"pr{i}", False),
                        'separator': st.session_state.get(f"sp{i}", ". ")
                    }
                ))
                used_elements.add(element)
        
        return {
            'author_format': st.session_state.get('auth', "AA Smith"),
            'author_separator': st.session_state.get('sep', ", "),
            'et_al_limit': st.session_state.get('etal', 0) if st.session_state.get('etal', 0) > 0 else None,
            'use_and_bool': st.session_state.get('use_and_checkbox', False),
            'use_ampersand_bool': st.session_state.get('use_ampersand_checkbox', False),
            'doi_format': st.session_state.get('doi', "10.10/xxx"),
            'doi_hyperlink': st.session_state.get('doilink', True),
            'page_format': st.session_state.get('page', "122-128"),
            'final_punctuation': st.session_state.get('punct', ""),
            'numbering_style': st.session_state.get('num', "No numbering"),
            'journal_style': st.session_state.get('journal_style', '{Full Journal Name}'),
            'elements': element_configs,
            'gost_style': st.session_state.get('gost_style', False),
            'acs_style': st.session_state.get('acs_style', False),
            'rsc_style': st.session_state.get('rsc_style', False),
            'cta_style': st.session_state.get('cta_style', False),
            'style5': st.session_state.get('style5', False),
            'style6': st.session_state.get('style6', False),
            'style7': st.session_state.get('style7', False),
            'style8': st.session_state.get('style8', False),
            'style9': st.session_state.get('style9', False),
            'style10': st.session_state.get('style10', False)
        }

    @staticmethod
    def _get_preview_metadata(style_config: Dict) -> Optional[Dict]:
        """Get metadata for preview - всегда возвращаем данные для превью"""
        # Всегда возвращаем данные для превью, даже если стиль не выбран
        # Это даст пользователю визуальную обратную связь
        return {
            'authors': [{'given': 'John A.', 'family': 'Smith'}, {'given': 'Alice B.', 'family': 'Doe'}],
            'title': 'Advanced Research in Materials Science',
            'journal': 'Journal of Materials Chemistry A',
            'year': 2023,
            'volume': '11',
            'issue': '15',
            'pages': '102-115',
            'article_number': 'e2301234',
            'doi': '10.1000/abc123'
        }
    
    @staticmethod
    def _add_numbering(preview_ref: str, style_config: Dict) -> str:
        """Add numbering to preview"""
        numbering = style_config['numbering_style']
        if numbering == "No numbering":
            return preview_ref
        elif numbering == "1":
            return f"1 {preview_ref}"
        elif numbering == "1.":
            return f"1. {preview_ref}"
        elif numbering == "1)":
            return f"1) {preview_ref}"
        elif numbering == "(1)":
            return f"(1) {preview_ref}"
        elif numbering == "[1]":
            return f"[1] {preview_ref}"
        else:
            return f"1. {preview_ref}"
    
    @staticmethod
    def _render_action_buttons():
        """Render action buttons"""
        col1, col2, col3 = st.columns([1, 1, 1])
        
        with col1:
            if st.button(get_text('back_to_start'), use_container_width=True, key="back_from_create"):
                StageManager.navigate_to('start')
        
        with col2:
            if st.button(get_text('export_style_button'), use_container_width=True, key="export_style_create"):
                style_config = CreatePage._get_style_config()
                export_data = CreatePage._export_style(style_config)
                if export_data:
                    st.download_button(
                        label=get_text('export_style'),
                        data=export_data,
                        file_name=f"{st.session_state.style_export_name}.json",
                        mime="application/json",
                        use_container_width=True,
                        key="download_exported_style"
                    )
        
        with col3:
            if st.button(get_text('proceed_to_io'), use_container_width=True, key="proceed_from_create"):
                style_config = CreatePage._get_style_config()
                st.session_state.style_config = style_config
                st.session_state.custom_style_created = True
                StageManager.navigate_to('io')
    
    @staticmethod
    def _export_style(style_config: Dict) -> Optional[bytes]:
        """Export style"""
        try:
            export_data = {
                'version': '1.0',
                'export_date': str(datetime.now()),
                'style_config': style_config
            }
            json_data = json.dumps(export_data, indent=2, ensure_ascii=False)
            return json_data.encode('utf-8')
        except Exception as e:
            st.error(f"Export error: {str(e)}")
            return None

# Input/Output Page
class InputOutputPage:
    """Input/Output page"""
    
    @staticmethod
    def render():
        """Render Input/Output page"""
        st.markdown(f"<h1>{get_text('io_title')}</h1>", unsafe_allow_html=True)
        st.markdown(f"<p style='margin-bottom: 30px;'>{get_text('io_description')}</p>", unsafe_allow_html=True)
        
        if not hasattr(st.session_state, 'style_config') or not st.session_state.style_config:
            st.warning(get_text('validation_error_no_elements'))
            if st.button(get_text('back_to_start'), use_container_width=True):
                StageManager.navigate_to('start')
            return
        
        st.markdown(f"<div class='card'><div class='card-title'>{get_text('data_input')}</div>", unsafe_allow_html=True)
        
        input_method = st.radio(
            get_text('input_method'),
            ['DOCX', 'Text' if st.session_state.current_language == 'en' else 'Текст'],
            horizontal=True,
            key="input_method"
        )
        
        if input_method == 'DOCX':
            uploaded_file = st.file_uploader(
                get_text('select_docx'),
                type=['docx'],
                label_visibility="collapsed",
                key="docx_uploader_io"
            )
            st.session_state.uploaded_file = uploaded_file
        else:
            text_input = st.text_area(
                get_text('references'),
                placeholder=get_text('enter_references'),
                height=150,
                label_visibility="collapsed",
                key="text_input_io"
            )
            st.session_state.text_input = text_input
        
        st.markdown("</div>", unsafe_allow_html=True)
        
        st.markdown(f"<div class='card'><div class='card-title'>{get_text('data_output')}</div>", unsafe_allow_html=True)
        
        output_method = st.radio(
            get_text('output_method'),
            ['DOCX', 'Text' if st.session_state.current_language == 'en' else 'Текст'],
            horizontal=True,
            key="output_method_io"
        )
        
        st.markdown("</div>", unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns([1, 2, 1])
        
        with col1:
            if st.button(get_text('back_button'), use_container_width=True, key="back_from_io"):
                StageManager.navigate_to('start')
        
        with col2:
            if st.button(get_text('process_references'), use_container_width=True, key="process_io"):
                InputOutputPage._process_data()
        
        with col3:
            if st.button(get_text('clear_all'), use_container_width=True, key="clear_io"):
                StageManager.clear_all()
    
    @staticmethod
    def _process_data():
        """Process data"""
        if not hasattr(st.session_state, 'style_config') or not st.session_state.style_config:
            st.error(get_text('validation_error_no_elements'))
            return
        
        if st.session_state.input_method == 'DOCX':
            if not st.session_state.uploaded_file:
                st.error(get_text('upload_file'))
                return
            references = InputOutputPage._extract_references_from_docx(st.session_state.uploaded_file)
        else:
            if not st.session_state.text_input.strip():
                st.error(get_text('enter_references_error'))
                return
            references = [ref.strip() for ref in st.session_state.text_input.split('\n') if ref.strip()]
        
        processor = ReferenceProcessor()
        progress_container = st.empty()
        status_container = st.empty()

        with st.spinner(get_text('processing')):
            formatted_refs, formatted_txt_buffer, original_txt_buffer, doi_found_count, doi_not_found_count, duplicates_info, missing_metadata_info = processor.process_references(
                references, st.session_state.style_config, progress_container, status_container
            )
            
            statistics = generate_statistics(formatted_refs)
            
            recommendations_df = None
            if len(formatted_refs) >= Config.MIN_REFERENCES_FOR_RECOMMENDATIONS:
                st.info(f"Found {len(formatted_refs)} references. Recommendations will be available on the Results page.")
            
            docx_buffer = DocumentGenerator.generate_document(
                formatted_refs, statistics, st.session_state.style_config, duplicates_info, missing_metadata_info, recommendations_df
            )
            
            st.session_state.formatted_refs = formatted_refs
            st.session_state.txt_buffer = formatted_txt_buffer
            st.session_state.formatted_txt_buffer = formatted_txt_buffer
            st.session_state.original_txt_buffer = original_txt_buffer
            st.session_state.docx_buffer = docx_buffer
            st.session_state.doi_found_count = doi_found_count
            st.session_state.doi_not_found_count = doi_not_found_count
            st.session_state.duplicates_info = duplicates_info
            st.session_state.missing_metadata_info = missing_metadata_info
            st.session_state.processing_complete = True
            st.session_state.processing_start_time = time.time()
            st.session_state.recommendations_generated = False
            st.session_state.recommendations = None
            
            StageManager.navigate_to('results')
    
    @staticmethod
    def _extract_references_from_docx(uploaded_file) -> List[str]:
        """Extract references from DOCX file"""
        doc = Document(uploaded_file)
        return [para.text.strip() for para in doc.paragraphs if para.text.strip()]

# Results Page with Recommendations
class ResultsPage:
    """Results page with recommendations"""
    
    @staticmethod
    def render():
        """Render results page"""
        st.markdown(f"<h1>{get_text('results_title')}</h1>", unsafe_allow_html=True)
        st.markdown(f"<p style='margin-bottom: 30px;'>{get_text('results_description')}</p>", unsafe_allow_html=True)
        
        if not st.session_state.processing_complete:
            st.warning(get_text('processing'))
            if st.button(get_text('back_to_start'), use_container_width=True):
                StageManager.navigate_to('start')
            return
        
        st.markdown(f"<div class='card'><div class='card-title'>{get_text('statistics_title')}</div>", unsafe_allow_html=True)
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.markdown(f"<div class='stat-card'><div class='stat-value'>{len(st.session_state.formatted_refs)}</div><div class='stat-label'>{get_text('total_references')}</div></div>", unsafe_allow_html=True)
        
        with col2:
            st.markdown(f"<div class='stat-card'><div class='stat-value'>{st.session_state.doi_found_count}</div><div class='stat-label'>{get_text('doi_found')}</div></div>", unsafe_allow_html=True)
        
        with col3:
            st.markdown(f"<div class='stat-card'><div class='stat-value'>{st.session_state.doi_not_found_count}</div><div class='stat-label'>{get_text('doi_not_found')}</div></div>", unsafe_allow_html=True)
        
        with col4:
            duplicates_count = len(st.session_state.duplicates_info)
            st.markdown(f"<div class='stat-card'><div class='stat-value'>{duplicates_count}</div><div class='stat-label'>{get_text('duplicates_found')}</div></div>", unsafe_allow_html=True)
        
        st.markdown("</div>", unsafe_allow_html=True)
        
        st.markdown(f"<div class='card'><div class='card-title'>Preview of Results ({len(st.session_state.formatted_refs)} references)</div>", unsafe_allow_html=True)
        
        st.markdown('<div class="scrollable-results">', unsafe_allow_html=True)
        
        for i, (elements, is_error, metadata) in enumerate(st.session_state.formatted_refs):
            css_class = "formatted-text"
            if is_error:
                css_class += " error-reference"
            elif i in st.session_state.duplicates_info:
                css_class += " duplicate-reference"
            elif i in st.session_state.missing_metadata_info:
                css_class += " missing-metadata-reference"
            
            if is_error:
                formatted_text = str(elements)
                display_html = f'<div class="{css_class}">{formatted_text}</div>'
            else:
                if isinstance(elements, str):
                    formatted_text = elements
                    display_html = f'<div class="{css_class}">{formatted_text}</div>'
                else:
                    html_parts = []
                    for j, element_data in enumerate(elements):
                        value, italic, bold, separator, is_doi_hyperlink, doi_value = element_data
                        
                        format_classes = []
                        if italic and bold:
                            format_classes.append("formatted-text-italic-bold")
                        elif italic:
                            format_classes.append("formatted-text-italic")
                        elif bold:
                            format_classes.append("formatted-text-bold")
                        
                        format_class = " ".join(format_classes) if format_classes else ""
                        
                        if format_class:
                            value_html = f'<span class="{format_class}">{value}</span>'
                        else:
                            value_html = value
                        
                        html_parts.append(value_html)
                        
                        if separator and j < len(elements) - 1:
                            html_parts.append(separator)
                    
                    if i in st.session_state.duplicates_info:
                        original_index = st.session_state.duplicates_info[i] + 1
                        duplicate_note = get_text('duplicate_reference').format(original_index)
                        html_parts.append(f' - <em>{duplicate_note}</em>')
                    elif i in st.session_state.missing_metadata_info:
                        missing_metadata_note = st.session_state.missing_metadata_info[i]
                        html_parts.append(f' - <em>{missing_metadata_note}</em>')
                    
                    if st.session_state.style_config.get('final_punctuation') and not is_error:
                        if html_parts and html_parts[-1].endswith('.'):
                            html_parts[-1] = html_parts[-1][:-1]
                        html_parts.append('.')
                    
                    numbering = st.session_state.style_config.get('numbering_style', 'No numbering')
                    prefix = ""
                    if numbering != "No numbering":
                        if numbering == "1":
                            prefix = f"{i + 1} "
                        elif numbering == "1.":
                            prefix = f"{i + 1}. "
                        elif numbering == "1)":
                            prefix = f"{i + 1}) "
                        elif numbering == "(1)":
                            prefix = f"({i + 1}) "
                        elif numbering == "[1]":
                            prefix = f"[{i + 1}] "
                        else:
                            prefix = f"{i + 1}. "
                    
                    full_html = prefix + "".join(html_parts)
                    display_html = f'<div class="{css_class}">{full_html}</div>'
            
            st.markdown(display_html, unsafe_allow_html=True)
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        st.markdown("</div>", unsafe_allow_html=True)
        
        st.markdown(f"<div class='card'><div class='card-title'>{get_text('download_results')}</div>", unsafe_allow_html=True)
        
        col_download1, col_download2 = st.columns(2)
        
        with col_download1:
            if st.session_state.txt_buffer:
                st.download_button(
                    label=get_text('download_txt'),
                    data=st.session_state.txt_buffer.getvalue(),
                    file_name='formatted_references.txt',
                    mime='text/plain',
                    use_container_width=True,
                    key="download_txt_results",
                    help="Download formatted references as plain text"
                )
        
        with col_download2:
            if st.session_state.docx_buffer:
                st.download_button(
                    label=get_text('download_docx'),
                    data=st.session_state.docx_buffer.getvalue(),
                    file_name='Reformatted references.docx',
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    use_container_width=True,
                    key="download_docx_results",
                    help="Download formatted references as DOCX document with full formatting"
                )
        
        st.markdown("</div>", unsafe_allow_html=True)
        
        if len(st.session_state.formatted_refs) >= Config.MIN_REFERENCES_FOR_RECOMMENDATIONS:
            ResultsPage._render_recommendations_section()
        
        col_nav1, col_nav2, col_nav3 = st.columns([1, 1, 1])
        
        with col_nav1:
            if st.button(get_text('back_button'), use_container_width=True, key="back_from_results"):
                StageManager.navigate_to('io')
        
        with col_nav2:
            if st.button(get_text('try_again'), use_container_width=True, key="try_again_results"):
                StageManager.navigate_to('io')
        
        with col_nav3:
            if st.button(get_text('new_session'), use_container_width=True, key="new_session_results"):
                StageManager.clear_all()

    @staticmethod
    def _render_recommendations_section():
        """Render recommendations section with topic-based analysis"""
        st.markdown(f"<div class='card'><div class='card-title'>{get_text('recommendations_title')}</div>", unsafe_allow_html=True)
        
        current_year = datetime.now().year
        min_year = current_year - 5

        st.markdown(f"<p>{get_text('recommendations_description')}</p>", unsafe_allow_html=True)
        # Можно добавить дополнительное пояснение
        current_year = datetime.now().year
        st.markdown(f"<p style='font-size: 0.9em; color: #666;'>Recommendations are based on new or low-citation articles published within the last 3 years (since {current_year - 3})</p>", unsafe_allow_html=True)
        
        # Проверяем достаточно ли ссылок
        if len(st.session_state.formatted_refs) < Config.MIN_REFERENCES_FOR_RECOMMENDATIONS:
            st.warning(get_text('recommendations_not_enough').format(Config.MIN_REFERENCES_FOR_RECOMMENDATIONS))
            st.markdown("</div>", unsafe_allow_html=True)
            return
        
        # Кнопка генерации рекомендаций
        if not st.session_state.get('recommendations_generated', False):
            col_rec1, col_rec2 = st.columns([3, 1])
            
            with col_rec1:
                st.info(f"📚 Found {len(st.session_state.formatted_refs)} references. Click the button to generate recommendations based on new or low-citation articles published within the last 3 years.")
            
            with col_rec2:
                if st.button(get_text('recommend_similar_articles'), 
                            use_container_width=True, 
                            key="generate_recommendations_new",
                            type="primary"):
                    st.session_state.recommendations_loading = True
                    st.rerun()
        
        # Генерация рекомендаций
        if st.session_state.get('recommendations_loading', False):
            progress_container = st.empty()
            status_container = st.empty()
            progress_bar = progress_container.progress(0)
            
            def update_progress(progress_value, message):
                """Обновление прогресс-бара"""
                try:
                    progress_bar.progress(progress_value / 100.0)
                    status_container.text(f"{message}")
                except:
                    pass
            
            try:
                # Генерируем рекомендации
                recommendations_df = ArticleRecommender.generate_recommendations(
                    st.session_state.formatted_refs,
                    progress_callback=update_progress
                )
                
                if recommendations_df is not None and not recommendations_df.empty:
                    # Сохраняем результаты
                    st.session_state.recommendations = recommendations_df
                    st.session_state.recommendations_generated = True
                    
                    # Создаем файлы для скачивания
                    recommendations_txt = ArticleRecommender.create_recommendations_txt(recommendations_df)
                    recommendations_csv = ArticleRecommender.create_recommendations_csv(recommendations_df)
                    
                    if recommendations_txt:
                        st.session_state.recommendations_txt_buffer = recommendations_txt
                    if recommendations_csv:
                        st.session_state.recommendations_csv_buffer = recommendations_csv
                    
                    update_progress(100, "✅ Analysis complete!")
                    time.sleep(1)
                    
                    st.success(f"✅ Found {len(recommendations_df)} new/low-citation article recommendations across {recommendations_df['topic'].nunique()} topics (last 3 years)")
                    st.rerun()
                    
                else:
                    update_progress(100, "⚠️ No recommendations found")
                    st.warning(get_text('recommendations_no_results'))
                    
            except Exception as e:
                update_progress(100, f"❌ Error occurred")
                logger.error(f"Recommendation generation error: {e}", exc_info=True)
                st.error(f"{get_text('recommendations_error')}: {str(e)[:200]}")
            
            finally:
                st.session_state.recommendations_loading = False
        
        # Отображение результатов
        if st.session_state.get('recommendations_generated', False) and st.session_state.recommendations is not None:
            recommendations_df = st.session_state.recommendations
            
            # Статистика
            topics_count = recommendations_df['topic'].nunique()
            avg_citations = recommendations_df['cited_by_count'].mean()
            zero_citation_count = (recommendations_df['cited_by_count'] == 0).sum()
            
            st.markdown(f"""
            <div class='stat-card' style='margin: 20px 0;'>
                <div class='stat-value'>{len(recommendations_df)}</div>
                <div class='stat-label'>New/low-citation articles</div>
                <div style='font-size: 0.8rem; margin-top: 5px;'>
                    📚 Topics: {topics_count} | 
                    📉 Avg citations: {avg_citations:.1f} | 
                    🔴 Zero citations: {zero_citation_count}
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            # Отображение рекомендаций по темам
            st.markdown("### 📋 Recommendations by Topic")
            
            # Создаем вкладки для каждой темы
            topic_names = recommendations_df['topic'].unique()[:5]  # Берем топ-5 тем
            
            if len(topic_names) > 0:
                tabs = st.tabs([f"📚 {topic[:25]}..." if len(topic) > 25 else f"📚 {topic}" 
                               for topic in topic_names])
                
                for i, tab in enumerate(tabs):
                    with tab:
                        topic_name = topic_names[i]
                        topic_works = recommendations_df[recommendations_df['topic'] == topic_name]
                        
                        st.markdown(f"### 📚 {topic_name}")
                        st.markdown(f"*Найдено {len(topic_works)} низкоцитируемых статей*")
                        
                        # Отображаем работы для этой темы с использованием expander
                        for idx, row in topic_works.head(30).iterrows():
                            # Создаем заголовок для expander с краткой информацией
                            expander_title = f"#{idx+1}: {row['title'][:180]}... (📊 {row['cited_by_count']} цит.)"
                            
                            # Используем expander с кастомным стилем
                            with st.expander(expander_title, expanded=False):
                                # Создаем карточку внутри expander с применением стилей темы
                                st.markdown(
                                    f"""
                                    <div class='recommendation-item' style='border: 1px solid var(--border); 
                                         background-color: var(--cardBackground); padding: 15px; border-radius: 8px;'>
                                    """,
                                    unsafe_allow_html=True
                                )
                                
                                col1, col2 = st.columns([3, 1])
                                
                                with col1:
                                    # Полное название
                                    st.markdown(f"**📝 Полное название:**")
                                    st.markdown(f"{row['title']}")
                                    
                                    # Авторы
                                    if pd.notna(row.get('authors_formatted')) and row['authors_formatted'] != "Unknown":
                                        st.markdown(f"**👤 Авторы:**")
                                        st.markdown(f"{row['authors_formatted']}")
                                    
                                    # Журнал и год
                                    if pd.notna(row.get('journal')):
                                        journal_text = f"**📰 Журнал:** {row['journal']}"
                                        if pd.notna(row.get('publication_year')):
                                            journal_text += f", **Год:** {row['publication_year']}"
                                        st.markdown(journal_text)
                                    
                                    # Ключевые слова
                                    if pd.notna(row.get('keywords_formatted')) and row['keywords_formatted']:
                                        st.markdown(f"**🔑 Ключевые слова:**")
                                        st.markdown(f"{row['keywords_formatted']}")
                                
                                with col2:
                                    # Статистика цитирований
                                    citation_color = "🔴" if row['cited_by_count'] == 0 else "🟡"
                                    citation_status = "0 цитирований" if row['cited_by_count'] == 0 else f"{row['cited_by_count']} цитирований"
                                    st.markdown(f"**📊 Цитирования:**")
                                    st.markdown(f"{citation_color} {citation_status}")
                                    
                                    # Релевантность
                                    st.markdown(f"**🎯 Релевантность:**")
                                    st.markdown(f"{row['relevance_score']}/10 баллов")
                                
                                # DOI ссылка (в отдельной строке)
                                if pd.notna(row.get('doi')) and row['doi']:
                                    doi_url = f"https://doi.org/{row['doi']}"
                                    st.markdown(f"**🔗 DOI:** [{row['doi']}]({doi_url})")
                                
                                st.markdown("</div>", unsafe_allow_html=True)
            
            # Кнопки скачивания
            st.markdown(f"<div class='card' style='margin-top: 20px;'><div class='card-title'>{get_text('recommendation_download')}</div>", unsafe_allow_html=True)
            
            col_rec_download1, col_rec_download2 = st.columns(2)
            
            with col_rec_download1:
                if hasattr(st.session_state, 'recommendations_txt_buffer') and st.session_state.recommendations_txt_buffer:
                    st.download_button(
                        label=get_text('recommendation_download_txt'),
                        data=st.session_state.recommendations_txt_buffer.getvalue(),
                        file_name='low_citation_recommendations.txt',
                        mime='text/plain',
                        use_container_width=True,
                        key="download_recommendations_txt_new"
                    )
            
            with col_rec_download2:
                if hasattr(st.session_state, 'recommendations_csv_buffer') and st.session_state.recommendations_csv_buffer:
                    st.download_button(
                        label=get_text('recommendation_download_csv'),
                        data=st.session_state.recommendations_csv_buffer.getvalue(),
                        file_name='low_citation_recommendations.csv',
                        mime='text/csv',
                        use_container_width=True,
                        key="download_recommendations_csv_new"
                    )
            
            st.markdown("</div>", unsafe_allow_html=True)
        
        st.markdown("</div>", unsafe_allow_html=True)

# Main Application Class
class CitationStyleApp:
    """Main application class"""
    
    def __init__(self):
        self.user_prefs = UserPreferencesManager()
        init_session_state()
    
    def run(self):
        """Run application"""
        st.set_page_config(
            page_title="Citation Style Constructor",
            page_icon="🎨",
            layout="wide"
        )
        
        self._load_user_preferences()
        
        ThemeManager.apply_theme(st.session_state.current_theme)
        
        self._render_header()
        
        StageManager.render_stage_indicator(st.session_state.current_stage)
        
        self._render_current_page()

    def _load_user_preferences(self):
        """Load user preferences"""
        if not st.session_state.user_prefs_loaded:
            ip = self.user_prefs.get_user_ip()
            prefs = self.user_prefs.get_preferences(ip)
            
            if 'current_language' not in st.session_state or not st.session_state.current_language:
                st.session_state.current_language = prefs['language']
            
            if 'current_theme' not in st.session_state or not st.session_state.current_theme:
                st.session_state.current_theme = prefs['theme']
            
            st.session_state.user_prefs_loaded = True

    def _render_header(self):
        """Render header and controls"""
        col_title, col_lang, col_theme = st.columns([2, 1, 1])
        
        with col_title:
            st.image("logo.png", width=250)
            st.title("Citation Style Constructor")
        
        with col_lang:
            languages = [('Русский', 'ru'), ('English', 'en')]
            selected_language = st.selectbox(
                get_text('choose_language'),
                languages,
                format_func=lambda x: x[0],
                index=0 if st.session_state.current_language == 'ru' else 1,
                key="language_selector_header"
            )
            
            if selected_language[1] != st.session_state.current_language:
                self.user_prefs.save_preferences(
                    self.user_prefs.get_user_ip(),
                    {
                        'language': selected_language[1],
                        'theme': st.session_state.current_theme
                    }
                )
                st.session_state.current_language = selected_language[1]
                st.rerun()
        
        with col_theme:
            themes = [
                (get_text('light_theme'), 'light'),
                (get_text('dark_theme'), 'dark'),
                (get_text('library_theme'), 'library'),
                (get_text('barbie_theme'), 'barbie'),
                (get_text('newspaper_theme'), 'newspaper')
            ]
            
            current_theme_index = 0
            for i, (_, theme_id) in enumerate(themes):
                if theme_id == st.session_state.current_theme:
                    current_theme_index = i
                    break
            
            selected_theme = st.selectbox(
                get_text('choose_theme'),
                themes,
                format_func=lambda x: x[0],
                index=current_theme_index,
                key="theme_selector_header"
            )
            
            if selected_theme[1] != st.session_state.current_theme:
                st.session_state.current_theme = selected_theme[1]
                self.user_prefs.save_preferences(
                    self.user_prefs.get_user_ip(),
                    {
                        'language': st.session_state.current_language,
                        'theme': st.session_state.current_theme
                    }
                )
                st.rerun()
    
    def _render_current_page(self):
        """Render current page"""
        current_stage = st.session_state.current_stage
        
        if current_stage == 'start':
            StartPage.render()
        elif current_stage == 'select':
            SelectPage.render()
        elif current_stage == 'create':
            CreatePage.render()
        elif current_stage == 'io':
            InputOutputPage.render()
        elif current_stage == 'results':
            ResultsPage.render()
        else:
            StartPage.render()

# Compatibility functions
def clean_text(text):
    return DOIProcessor()._clean_text(text)

def normalize_name(name):
    return DOIProcessor()._normalize_name(name)

def is_section_header(text):
    return DOIProcessor()._is_section_header(text)

def find_doi(reference):
    processor = DOIProcessor()
    return processor.find_doi_enhanced(reference)

def normalize_doi(doi):
    processor = ReferenceProcessor()
    return processor._normalize_doi(doi)

def generate_reference_hash(metadata):
    processor = ReferenceProcessor()
    return processor._generate_reference_hash(metadata)

def extract_metadata_batch(doi_list, progress_callback=None):
    processor = ReferenceProcessor()
    return [processor.doi_processor.extract_metadata_with_cache(doi) for doi in doi_list]

def extract_metadata_sync(doi):
    processor = ReferenceProcessor()
    return processor.doi_processor.extract_metadata_with_cache(doi)

def format_reference(metadata, style_config, for_preview=False):
    formatter = CitationFormatterFactory.create_formatter(style_config)
    return formatter.format_reference(metadata, for_preview)

def find_duplicate_references(formatted_refs):
    processor = ReferenceProcessor()
    return processor._find_duplicates(formatted_refs)

def generate_statistics(formatted_refs):
    journals = []
    years = []
    authors = []
    
    current_year = datetime.now().year
    
    for _, _, metadata in formatted_refs:
        if not metadata:
            continue
            
        if metadata.get('journal'):
            journals.append(metadata['journal'])
        
        if metadata.get('year'):
            years.append(metadata['year'])
        
        if metadata.get('authors'):
            for author in metadata['authors']:
                given = author.get('given', '')
                family = author.get('family', '')
                if family:
                    first_initial = given[0] if given else ''
                    author_formatted = f"{family} {first_initial}." if first_initial else family
                    authors.append(author_formatted)
    
    unique_dois = set()
    for _, _, metadata in formatted_refs:
        if metadata and metadata.get('doi'):
            unique_dois.add(metadata['doi'])
    
    total_unique_dois = len(unique_dois)
    
    journal_counter = Counter(journals)
    journal_stats = []
    for journal, count in journal_counter.most_common(20):
        percentage = (count / total_unique_dois) * 100 if total_unique_dois > 0 else 0
        journal_stats.append({
            'journal': journal,
            'count': count,
            'percentage': round(percentage, 2)
        })
    
    year_counter = Counter(years)
    year_stats = []
    for year in range(current_year, 2009, -1):
        if year in year_counter:
            count = year_counter[year]
            percentage = (count / total_unique_dois) * 100 if total_unique_dois > 0 else 0
            year_stats.append({
                'year': year,
                'count': count,
                'percentage': round(percentage, 2)
            })
    
    recent_years = [current_year - i for i in range(4)]
    recent_count = sum(year_counter.get(year, 0) for year in recent_years)
    recent_percentage = (recent_count / total_unique_dois) * 100 if total_unique_dois > 0 else 0
    needs_more_recent_references = recent_percentage < 20
    
    author_counter = Counter(authors)
    author_stats = []
    for author, count in author_counter.most_common(20):
        percentage = (count / total_unique_dois) * 100 if total_unique_dois > 0 else 0
        author_stats.append({
            'author': author,
            'count': count,
            'percentage': round(percentage, 2)
        })
    
    has_frequent_author = any(stats['percentage'] > 30 for stats in author_stats)
    
    return {
        'journal_stats': journal_stats,
        'year_stats': year_stats,
        'author_stats': author_stats,
        'total_unique_dois': total_unique_dois,
        'needs_more_recent_references': needs_more_recent_references,
        'has_frequent_author': has_frequent_author
    }

def process_references_with_progress(references, style_config, progress_container, status_container):
    processor = ReferenceProcessor()
    return processor.process_references(references, style_config, progress_container, status_container)

def process_docx(input_file, style_config, progress_container, status_container):
    processor = ReferenceProcessor()
    doc = Document(input_file)
    references = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    return processor.process_references(references, style_config, progress_container, status_container)

def export_style(style_config, file_name):
    try:
        export_data = {
            'version': '1.0',
            'export_date': str(datetime.now()),
            'style_config': style_config
        }
        json_data = json.dumps(export_data, indent=2, ensure_ascii=False)
        return json_data.encode('utf-8')
    except Exception as e:
        return None

def import_style(uploaded_file):
    try:
        content = uploaded_file.read().decode('utf-8')
        import_data = json.loads(content)
        
        if 'style_config' in import_data:
            return import_data['style_config']
        else:
            return import_data
    except Exception as e:
        return None

def apply_imported_style(imported_style):
    """Apply imported style"""
    if not imported_style:
        return
    
    if 'author_format' in imported_style:
        st.session_state.auth = imported_style['author_format']
    if 'author_separator' in imported_style:
        st.session_state.sep = imported_style['author_separator']
    if 'et_al_limit' in imported_style:
        st.session_state.etal = imported_style['et_al_limit'] or 0
    if 'use_and_bool' in imported_style:
        st.session_state.use_and_checkbox = imported_style['use_and_bool']
    if 'use_ampersand_bool' in imported_style:
        st.session_state.use_ampersand_checkbox = imported_style['use_ampersand_bool']
    if 'doi_format' in imported_style:
        st.session_state.doi = imported_style['doi_format']
    if 'doi_hyperlink' in imported_style:
        st.session_state.doilink = imported_style['doi_hyperlink']
    if 'page_format' in imported_style:
        st.session_state.page = imported_style['page_format']
    if 'final_punctuation' in imported_style:
        st.session_state.punct = imported_style['final_punctuation']
    if 'journal_style' in imported_style:
        st.session_state.journal_style = imported_style['journal_style']
    if 'numbering_style' in imported_style:
        st.session_state.num = imported_style['numbering_style']
    
    st.session_state.gost_style = imported_style.get('gost_style', False)
    st.session_state.acs_style = imported_style.get('acs_style', False)
    st.session_state.rsc_style = imported_style.get('rsc_style', False)
    st.session_state.cta_style = imported_style.get('cta_style', False)
    st.session_state.style5 = imported_style.get('style5', False)
    st.session_state.style6 = imported_style.get('style6', False)
    st.session_state.style7 = imported_style.get('style7', False)
    st.session_state.style8 = imported_style.get('style8', False)
    st.session_state.style9 = imported_style.get('style9', False)
    st.session_state.style10 = imported_style.get('style10', False)
    
    for i in range(8):
        st.session_state[f"el{i}"] = ""
        st.session_state[f"it{i}"] = False
        st.session_state[f"bd{i}"] = False
        st.session_state[f"pr{i}"] = False
        st.session_state[f"sp{i}"] = ". "
    
    elements = imported_style.get('elements', [])
    for i, (element, config) in enumerate(elements):
        if i < 8:
            st.session_state[f"el{i}"] = element
            st.session_state[f"it{i}"] = config.get('italic', False)
            st.session_state[f"bd{i}"] = config.get('bold', False)
            st.session_state[f"pr{i}"] = config.get('parentheses', False)
            st.session_state[f"sp{i}"] = config.get('separator', ". ")
    
    st.session_state.style_config = imported_style
    st.session_state.custom_style_created = True

def main():
    """Main function"""
    app = CitationStyleApp()
    app.run()

if __name__ == "__main__":
    main()














